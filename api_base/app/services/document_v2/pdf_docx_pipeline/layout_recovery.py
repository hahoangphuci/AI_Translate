"""Recover DOCX layout after PDF->DOCX translation.

Keeps original spacing/indent/alignment as closely as possible:
- Title block: copy from source (centered author/affiliation)
- Abstract + Keywords: inset column, justify, merge pdf2docx line fragments
- Section 1 body: body template from section 2 reference
- Section 2+: preserve full paragraph properties from source (spacing, firstLine)
"""

from __future__ import annotations

import copy
import os
import re
from typing import Any, Dict, Iterable, Iterator, List, Optional, Tuple

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

_ABSTRACT_EXTRA_INDENT_TWIPS = 566

_ABSTRACT_START_RE = re.compile(
    r"^(abstract|abstrakt|abstrak|opsomming|p[eë]rmbledhje|"
    r"t[oóô]m\s*t[aắ]t|t[oóô]mt[aắ]t|zusammenfassung|r[eé]sum[eé])\b",
    re.IGNORECASE,
)
_KEYWORDS_LINE_RE = re.compile(
    r"^(keywords?|keyword|t[uừ]\s*kh[oó]a|t[uừ]kh[oó]a|"
    r"sleutelwoorde?|fjal[eë]t\s*ky[çc]e|schl[uü]sselw[oö]rter)\b",
    re.IGNORECASE,
)
_SECTION1_HEAD_RE = re.compile(
    r"^1\s+(introduction|inleiding|hyrje|gi[oớ]i\s*th[iệ]u|gi[oớ]ith[iệ]u|einleitung)\b",
    re.IGNORECASE,
)
_FALSE_SECTION2_RE = re.compile(r"^2\s+F\.\s", re.IGNORECASE)
_RUNNING_HEADER_RE = re.compile(
    r"^\d+\s+[A-ZÀ-Ỹ]\.\s+(?:Author|Skrywer)\b",
    re.IGNORECASE,
)
_PAGE_NO_ONLY_RE = re.compile(r"^\d{1,4}$")
_PDF_ARTIFACT_RE = re.compile(
    # Match pdf2docx internal word tokens (__word_123__) OR
    # lines composed entirely of separator characters with NO digits (digits could be formula content).
    r"^(?:__?\s*[\wÀ-Ỹ]{2,10}\s*[_\s]*\d+\s*__?|(?!.*\d)[\s_,.|/-]{5,})$",
    re.IGNORECASE,
)
_AFFILIATION_LINE_RE = re.compile(
    r"^\d+\s+(faculty|department|university|springer|institute|fakulteit|khoa|tr[uườ]ng)\b",
    re.IGNORECASE,
)
_SECTION2_HEAD_RE = re.compile(r"^2\s+\S", re.IGNORECASE)
_PDF_SUBSET_FONT_RE = re.compile(r"^[A-Z]{6}\+(.+)$")
_SYMBOL_FONT_HINT_RE = re.compile(r"(symbol|wingdings|webdings|dingbats?|zapf|mt\s*extra)", re.IGNORECASE)
_UNSAFE_FONT_HINT_RE = re.compile(
    r"(symbol|wingdings|webdings|dingbats?|zapf|mt\s*extra|marlett|itc\s*zapf|books\s*symbol|"
    r"script|cursive|handwriting|comic\s*sans|brush\s*script|freestyle|"
    r"segoe\s*script|lucida\s*hand|bradley|kristen|jokerman|"
    r"papyrus|chalkduster|mistral|monotype\s*corsiva|curlz|ravie|showcard|snap\s*itc|"
    r"calibri[-\s]*light|cambriamath|ms\s*reference|architects\s*daughter|"
    r"Segoe\s*Print|Segoe\s*UI\s*Emoji)",
    re.IGNORECASE,
)
_PS_FONT_CANONICAL = (
    (re.compile(r"timesnewroman|times[-\s]*roman", re.I), "Times New Roman"),
    (re.compile(r"^arial(?!.*narrow)", re.I), "Arial"),
    (re.compile(r"helvetica", re.I), "Arial"),
    (re.compile(r"calibri", re.I), "Calibri"),
    (re.compile(r"cambria(?!math)", re.I), "Cambria"),
    (re.compile(r"courier(?:\s*new)?", re.I), "Courier New"),
    (re.compile(r"georgia", re.I), "Georgia"),
    (re.compile(r"verdana", re.I), "Verdana"),
    (re.compile(r"tahoma", re.I), "Tahoma"),
    (re.compile(r"garamond", re.I), "Garamond"),
    (re.compile(r"trebuchet", re.I), "Trebuchet MS"),
    (re.compile(r"noto\s*sans", re.I), "Noto Sans"),
)
_KNOWN_WINDOWS_FONTS = {
    "times new roman": "Times New Roman",
    "arial": "Arial",
    "calibri": "Calibri",
    "cambria": "Cambria",
    "courier new": "Courier New",
    "georgia": "Georgia",
    "verdana": "Verdana",
    "tahoma": "Tahoma",
    "garamond": "Garamond",
    "trebuchet ms": "Trebuchet MS",
    "noto sans": "Noto Sans",
    "segoe ui": "Segoe UI",
    "book antiqua": "Book Antiqua",
    "palatino linotype": "Palatino Linotype",
}


def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in ("1", "true", "yes", "on")


def _env_str(name: str, default: str = "") -> str:
    raw = os.getenv(name)
    if raw is None:
        return str(default)
    return str(raw).strip()


def normalize_bilingual_mode(mode: Optional[str]) -> str:
    """Map frontend/API bilingual_mode values to none|inline|newline."""
    bi = (str(mode or "").strip().lower() or "none")
    if bi in ("preserve_layout", "inline", "lien_ke", "adjacent", "side_by_side"):
        return "inline"
    if bi in ("line_by_line", "newline", "xuong_dong", "stacked"):
        return "newline"
    if bi in ("none", "off", "0", "false"):
        return "none"
    return "none"


def _env_float(name: str, default: float) -> float:
    raw = os.getenv(name)
    if raw is None:
        return default
    try:
        return float(str(raw).strip())
    except Exception:
        return default


def _paragraph_plain(paragraph) -> str:
    try:
        return (paragraph.text or "").replace("\r", "").replace("\n", " ").strip()
    except Exception:
        return ""


def _paragraph_word_count(paragraph) -> int:
    return len(_paragraph_plain(paragraph).split())


def _looks_like_body_prose(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if len(text) < 40:
        return False
    if _is_section_heading_line(paragraph):
        return False
    return _paragraph_word_count(paragraph) >= 8


def _is_short_label_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    return 0 < len(text) <= 90 and _paragraph_word_count(paragraph) <= 14


def _is_section_heading_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if not text:
        return False
    if _is_running_header_line(paragraph):
        return False
    if _SECTION1_HEAD_RE.match(text):
        return True
    if _SECTION2_HEAD_RE.match(text) and not _FALSE_SECTION2_RE.match(text):
        return True
    if re.match(r"^\d+\s+[A-ZÀ-Ỹ]", text) and _paragraph_word_count(paragraph) <= 8:
        return True
    return False


def _is_running_header_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if not text:
        return False
    if _RUNNING_HEADER_RE.match(text):
        return True
    if _FALSE_SECTION2_RE.match(text):
        return True
    return False


def _is_in_table_cell(paragraph) -> bool:
    try:
        el = paragraph._element
        while el is not None:
            tag = getattr(el, "tag", "") or ""
            if tag.endswith("}tc"):
                return True
            el = el.getparent()
    except Exception:
        pass
    return False


def iter_all_paragraphs(doc: docx.Document) -> Iterator:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _strip_pdf_subset_font_prefix(name: Optional[str]) -> Optional[str]:
    if name is None:
        return None
    out = str(name).strip()
    if not out:
        return out
    while True:
        m = _PDF_SUBSET_FONT_RE.match(out)
        if not m:
            break
        nxt = str(m.group(1) or "").strip()
        if not nxt or nxt == out:
            break
        out = nxt
    return out


def _normalize_rfonts_attrs(rfonts_el) -> int:
    if rfonts_el is None:
        return 0
    changed = 0
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        try:
            current = rfonts_el.get(qn(attr))
        except Exception:
            current = None
        if not current:
            continue
        cleaned = _strip_pdf_subset_font_prefix(current)
        if cleaned and cleaned != current:
            try:
                rfonts_el.set(qn(attr), cleaned)
                changed += 1
            except Exception:
                pass
    return changed


def _normalize_run_font_names(run) -> int:
    changed = 0
    try:
        current = getattr(getattr(run, "font", None), "name", None)
        cleaned = _strip_pdf_subset_font_prefix(current)
        if cleaned and cleaned != current:
            run.font.name = cleaned
            changed += 1
    except Exception:
        pass
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            changed += _normalize_rfonts_attrs(rfonts)
    except Exception:
        pass
    return changed


def _iter_table_paragraphs_recursive(table) -> Iterator:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs_recursive(nested_table)


def _iter_all_paragraphs_for_font_fix(doc: docx.Document) -> Iterator:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_table_paragraphs_recursive(table)

    seen_header_footer = set()
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            try:
                key = id(header_footer._element)
            except Exception:
                key = id(header_footer)
            if key in seen_header_footer:
                continue
            seen_header_footer.add(key)

            for paragraph in header_footer.paragraphs:
                yield paragraph
            for table in header_footer.tables:
                yield from _iter_table_paragraphs_recursive(table)


def _normalize_style_font_names(doc: docx.Document) -> int:
    changed = 0
    try:
        for style in doc.styles:
            try:
                current = getattr(getattr(style, "font", None), "name", None)
                cleaned = _strip_pdf_subset_font_prefix(current)
                if cleaned and cleaned != current:
                    style.font.name = cleaned
                    changed += 1
            except Exception:
                pass

            try:
                style_el = getattr(style, "element", None) or getattr(style, "_element", None)
                if style_el is None:
                    continue
                rpr = style_el.find(qn("w:rPr"))
                if rpr is None:
                    continue
                rfonts = rpr.find(qn("w:rFonts"))
                changed += _normalize_rfonts_attrs(rfonts)
            except Exception:
                pass
    except Exception:
        pass

    try:
        styles_el = getattr(doc.styles, "element", None)
        if styles_el is not None:
            doc_defaults = styles_el.find(qn("w:docDefaults"))
            if doc_defaults is not None:
                rpr_default = doc_defaults.find(qn("w:rPrDefault"))
                if rpr_default is not None:
                    rpr = rpr_default.find(qn("w:rPr"))
                    if rpr is not None:
                        rfonts = rpr.find(qn("w:rFonts"))
                        changed += _normalize_rfonts_attrs(rfonts)
    except Exception:
        pass

    return changed


def _normalize_subset_font_names_in_doc(doc: docx.Document) -> int:
    if not _env_bool("PDF_DOCX_NORMALIZE_SUBSET_FONTS", True):
        return 0
    changed = 0
    for paragraph in _iter_all_paragraphs_for_font_fix(doc):
        for run in paragraph.runs:
            changed += _normalize_run_font_names(run)
    changed += _normalize_style_font_names(doc)
    return changed


def _pdf_docx_text_font_fallback() -> str:
    return _env_str("PDF_DOCX_TEXT_FONT_FALLBACK", "Times New Roman") or "Times New Roman"


def _run_font_name_candidates(run) -> List[str]:
    names: List[str] = []
    seen = set()
    try:
        current = getattr(getattr(run, "font", None), "name", None)
        if current:
            cleaned = _strip_pdf_subset_font_prefix(current) or ""
            if cleaned and cleaned.lower() not in seen:
                seen.add(cleaned.lower())
                names.append(cleaned)
    except Exception:
        pass
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    val = rfonts.get(qn(attr))
                    if not val:
                        continue
                    cleaned = _strip_pdf_subset_font_prefix(val) or ""
                    if cleaned and cleaned.lower() not in seen:
                        seen.add(cleaned.lower())
                        names.append(cleaned)
    except Exception:
        pass
    return names


def _canonical_document_font_name(name: Optional[str]) -> str:
    fallback = _pdf_docx_text_font_fallback()
    cleaned = _strip_pdf_subset_font_prefix(name) or ""
    cleaned = re.sub(r"\s+", " ", str(cleaned)).strip()
    if not cleaned:
        return fallback
    if _UNSAFE_FONT_HINT_RE.search(cleaned) or _SYMBOL_FONT_HINT_RE.search(cleaned):
        return fallback
    low = cleaned.lower()
    if low in _KNOWN_WINDOWS_FONTS:
        return _KNOWN_WINDOWS_FONTS[low]
    if re.search(r"PSMT|PS-Bold|PS-Italic|BoldItalicMT|,-?Bold(?:MT)?$", cleaned, re.I):
        for pattern, canonical in _PS_FONT_CANONICAL:
            if pattern.search(cleaned):
                return canonical
        return fallback
    for pattern, canonical in _PS_FONT_CANONICAL:
        if pattern.search(cleaned):
            return canonical
    if _env_bool("PDF_DOCX_CANONICALIZE_UNKNOWN_FONTS", True):
        return fallback
    return cleaned


def _set_run_font_name(run, font_name: str) -> None:
    target = (font_name or "").strip()
    if not target:
        return
    try:
        run.font.name = target
    except Exception:
        pass
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attr), target)
    except Exception:
        pass


def _canonicalize_rfonts_in_rpr(rpr, text: str = "") -> bool:
    if rpr is None:
        return False
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return False
    _normalize_rfonts_attrs(rfonts)
    has_text_tokens = bool(re.search(r"[A-Za-zÀ-ỹ0-9]", text or "", flags=re.UNICODE))
    has_list_marker = bool(re.search(r"[+\-•*]", text or ""))
    if not (has_text_tokens or has_list_marker):
        return False
    current = None
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        val = rfonts.get(qn(attr))
        if val:
            current = val
            break
    canonical = _canonical_document_font_name(current)
    if not canonical:
        return False
    changed = False
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        cur = rfonts.get(qn(attr))
        if cur:
            nxt = _canonical_document_font_name(cur)
            if nxt and nxt != cur:
                rfonts.set(qn(attr), nxt)
                changed = True
        elif canonical:
            rfonts.set(qn(attr), canonical)
            changed = True
    return changed


def _sanitize_run_text_font(run, text: str = "") -> int:
    raw = text if text is not None else (run.text or "")
    if not re.search(r"[A-Za-zÀ-ỹ0-9]", raw or "", flags=re.UNICODE):
        return 0
    changed = _normalize_run_font_names(run)
    candidates = _run_font_name_candidates(run)
    source_name = candidates[0] if candidates else ""
    canonical = _canonical_document_font_name(source_name)
    if not canonical:
        return changed
    if source_name and source_name != canonical:
        _set_run_font_name(run, canonical)
        changed += 1
    try:
        rpr = run._element.find(qn("w:rPr"))
        if _canonicalize_rfonts_in_rpr(rpr, raw):
            changed += 1
    except Exception:
        pass
    return changed


def sanitize_document_text_fonts(doc: docx.Document) -> int:
    """Map pdf2docx/PS/script fonts to stable Windows-safe fonts before PDF export."""
    if not _env_bool("PDF_DOCX_SANITIZE_FONTS", True):
        return 0
    changed = 0
    for paragraph in _iter_all_paragraphs_for_font_fix(doc):
        for run in paragraph.runs:
            try:
                changed += _sanitize_run_text_font(run, run.text or "")
            except Exception:
                continue
    changed += _normalize_style_font_names(doc)
    return changed


def _dominant_font_pt(paragraph) -> Optional[float]:
    sizes: List[float] = []
    for run in paragraph.runs:
        try:
            if run.font.size is not None:
                sizes.append(float(run.font.size.pt))
        except Exception:
            pass
    if not sizes:
        return None
    sizes.sort()
    return sizes[len(sizes) // 2]


def _find_abstract_start(paras: List) -> Optional[int]:
    for i, p in enumerate(paras):
        if _ABSTRACT_START_RE.match(_paragraph_plain(p)):
            return i
    return None


def _find_keywords_paragraph_index(paras: List, start: int = 0) -> Optional[int]:
    for i in range(start, len(paras)):
        if _KEYWORDS_LINE_RE.match(_paragraph_plain(paras[i])):
            return i
    return None


def _find_section_one_start(paras: List) -> Optional[int]:
    for i, p in enumerate(paras):
        text = _paragraph_plain(p)
        if _SECTION1_HEAD_RE.match(text):
            return i
        if re.match(r"^1\s+[A-ZÀ-Ỹ]", text) and _paragraph_word_count(p) <= 8:
            if not _AFFILIATION_LINE_RE.match(text):
                return i
    return None


def _find_section_two_start(paras: List, *, min_index: int = 0) -> Optional[int]:
    sec1 = _find_section_one_start(paras)
    start_at = max(int(min_index or 0), (sec1 + 1) if sec1 is not None else 0)
    for i, p in enumerate(paras):
        if i < start_at:
            continue
        text = _paragraph_plain(p)
        if _FALSE_SECTION2_RE.match(text):
            continue
        if _is_running_header_line(p):
            continue
        if _AFFILIATION_LINE_RE.match(text):
            continue
        if re.search(r"\bskrywer\b", text, re.IGNORECASE) and _paragraph_word_count(p) <= 12:
            continue
        if _SECTION2_HEAD_RE.match(text) and _paragraph_word_count(p) <= 8:
            return i
    return None


def _analysis_is_academic_like(analysis: Optional[Dict[str, Any]]) -> bool:
    if not isinstance(analysis, dict):
        return False
    if bool(analysis.get("is_academic_like")):
        return True
    try:
        return int(analysis.get("academic_score") or 0) >= 5
    except Exception:
        return False


def _preserve_layout_mode() -> str:
    mode = _env_str("PDF_DOCX_PRESERVE_LAYOUT_MODE", "force").lower()
    if mode in ("force", "strict", "mirror", "on", "1", "true", "yes"):
        return "force"
    if mode in ("off", "disable", "disabled", "0", "false", "no"):
        return "off"
    return "auto"


def _doc_looks_academic_like(doc_obj: Optional[docx.Document]) -> bool:
    if doc_obj is None:
        return False
    try:
        paras = list(doc_obj.paragraphs)
    except Exception:
        return False
    if len(paras) < 10:
        return False

    abs_idx = _find_abstract_start(paras)
    kw_idx = _find_keywords_paragraph_index(paras, abs_idx if abs_idx is not None else 0)
    sec1 = _find_section_one_start(paras)
    sec2 = _find_section_two_start(paras, min_index=(sec1 + 1) if sec1 is not None else 0)

    refs_hits = 0
    numbered_hits = 0
    for paragraph in paras[: min(len(paras), 180)]:
        text = _paragraph_plain(paragraph).lower()
        if not text:
            continue
        if refs_hits < 2 and re.search(
            r"\b(references|reference|bibliography|tài liệu tham khảo|tai lieu tham khao)\b",
            text,
            re.IGNORECASE,
        ):
            refs_hits += 1
        if numbered_hits < 6 and re.match(r"^\d{1,2}(?:\.\d+)?\s+[a-zà-ỹ]", text):
            numbered_hits += 1

    score = 0
    if abs_idx is not None:
        score += 2
    if kw_idx is not None:
        score += 1
    if sec1 is not None:
        score += 1
    if sec2 is not None:
        score += 1
    if refs_hits:
        score += 2
    if numbered_hits >= 2:
        score += 1
    if abs_idx is not None and sec1 is not None and abs_idx < sec1:
        score += 1

    return bool(score >= 5 or (score >= 4 and refs_hits > 0))


def _should_apply_regional_layout(
    *,
    doc: Optional[docx.Document] = None,
    src_doc: Optional[docx.Document] = None,
    analysis: Optional[Dict[str, Any]] = None,
) -> bool:
    preserve_mode = _preserve_layout_mode()
    mode = _env_str("PDF_DOCX_REGIONAL_MODE", "auto").lower()

    # Explicit regional force works unless strict source-layout preservation is forced.
    if mode in ("force", "on", "1", "true", "yes", "regional"):
        return preserve_mode != "force"
    if mode in ("off", "0", "false", "no", "source", "safe", "disable", "disabled"):
        return False

    # Keep strict source-layout by default to avoid damaging normal PDFs.
    if preserve_mode == "force":
        return False

    # Backward-compatible hard off switch.
    if not _env_bool("PDF_DOCX_REGIONAL_LAYOUT", True):
        return False

    if _analysis_is_academic_like(analysis):
        return True
    if _doc_looks_academic_like(src_doc):
        return True
    if _doc_looks_academic_like(doc):
        return True
    return False


def _should_preserve_source_layout(
    *,
    doc: Optional[docx.Document] = None,
    src_doc: Optional[docx.Document] = None,
    analysis: Optional[Dict[str, Any]] = None,
) -> bool:
    mode = _preserve_layout_mode()
    if mode == "force":
        return True
    if mode == "off":
        return False

    # auto: preserve source layout for non-academic docs; allow regional only when academic-like.
    if _analysis_is_academic_like(analysis):
        return False
    if _doc_looks_academic_like(src_doc):
        return False
    if _doc_looks_academic_like(doc):
        return False
    return True


def is_running_header_text(text: str) -> bool:
    t = (text or "").replace("\r", " ").replace("\n", " ").strip()
    if not t:
        return False
    if _PAGE_NO_ONLY_RE.match(t):
        return True
    if _RUNNING_HEADER_RE.match(t):
        return True
    if _FALSE_SECTION2_RE.match(t):
        return True
    if re.match(r"^\d+\s+[A-ZÀ-Ỹ]\.\s+\S+\s+(?:en|and)\s+[A-ZÀ-Ỹ]\.\s+(?:Author|Skrywer)\b", t, re.IGNORECASE):
        return True
    return False


def is_pdf_artifact_text(text: str) -> bool:
    t = (text or "").replace("\r", " ").replace("\n", " ").strip()
    if not t:
        return False
    if _PDF_ARTIFACT_RE.match(t):
        return True
    if re.search(r"__\s*[\wÀ-Ỹ]{2,10}\s*_\s*\d+\s*__", t, re.IGNORECASE):
        return True
    if re.match(r"^__?[\wÀ-Ỹ]{2,10}_?\d*__?\s*[,.\s]*$", t, re.IGNORECASE):
        return True
    if re.search(r"__?\s*HOU[\s_\d]*__?", t, re.IGNORECASE):
        return True
    # Pure separator/whitespace lines — only if they contain NO digits and NO math symbols.
    if re.match(r"^[,.\s|_\-]+$", t) and not re.search(r"[\d\u0391-\u03C9\u2200-\u22FF\u2190-\u21FF]", t):
        return True
    # Strings dominated by separator chars — exclude any with digits, Unicode math, or operators.
    # These could be formula expressions rendered as plain text from PDF.
    has_math = bool(re.search(
        r"[\d=+*/^(){}\[\]<>≤≥≠∑∫∏√∞±∓∂∇"
        r"\u0391-\u03C9"       # Greek
        r"\u2200-\u22FF"       # Math operators
        r"\u2190-\u21FF"       # Arrows
        r"\u00B2\u00B3\u00B9"  # ² ³ ¹
        r"\u2070-\u2079"       # Superscript digits
        r"\u2080-\u2089"       # Subscript digits
        r"]", t
    ))
    if has_math:
        return False
    letters = len(re.findall(r"[A-Za-zÀ-ỹ]", t))
    if len(t) <= 48 and letters <= 2 and re.search(r"[_|,.\s-]{3,}", t):
        return True
    return False


def _para_has_drawing(paragraph) -> bool:
    """Return True if the paragraph contains any drawing, image, or VML picture element.

    Such paragraphs must never be removed or merged, as they may hold formula images,
    charts, or other embedded objects that cannot be reconstructed from text.
    """
    try:
        return bool(
            paragraph._element.xpath(
                './/*[local-name()="drawing" or local-name()="pict"'
                ' or local-name()="object" or local-name()="OLEObject"]'
            )
        )
    except Exception:
        return False


def _para_has_omath(paragraph) -> bool:
    """Return True if the paragraph contains Office Math (OMML) markup."""
    try:
        return bool(paragraph._element.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]'))
    except Exception:
        return False


def _para_has_protected_content(paragraph) -> bool:
    """True if a paragraph contains drawings, OMML math, or bookmarks that must be preserved."""
    return _para_has_drawing(paragraph) or _para_has_omath(paragraph)


def _is_noise_paragraph(paragraph) -> bool:
    # Never remove paragraphs that carry drawings, formulas, or embedded objects.
    if _para_has_protected_content(paragraph):
        return False
    text = _paragraph_plain(paragraph)
    return is_running_header_text(text) or is_pdf_artifact_text(text)


def strip_noise_paragraphs(doc: docx.Document) -> Dict[str, int]:
    """Remove running headers/footers and pdf2docx artifacts injected into body flow."""
    stats = {"noise_removed": 0, "inline_artifacts_stripped": 0}
    if not _env_bool("PDF_DOCX_STRIP_NOISE", True):
        return stats
    stats["inline_artifacts_stripped"] = _strip_inline_pdf_artifacts(doc)
    paras = list(doc.paragraphs)
    idx = 0
    while idx < len(paras):
        if _is_noise_paragraph(paras[idx]):
            try:
                paras[idx]._element.getparent().remove(paras[idx]._element)
                paras.pop(idx)
                stats["noise_removed"] += 1
                continue
            except Exception:
                pass
        idx += 1
    return stats


_INLINE_ARTIFACT_RE = re.compile(
    r"__?\s*[\wÀ-Ỹ]{2,10}\s*_\s*\d+\s*__?\s*[,.\s]*",
    re.IGNORECASE,
)


def _run_has_drawing(run) -> bool:
    """Return True if the run contains a drawing or VML picture child element."""
    try:
        return bool(
            run._element.xpath(
                './*[local-name()="drawing" or local-name()="pict"'
                ' or local-name()="object"]'
            )
        )
    except Exception:
        return False


def _strip_inline_pdf_artifacts(doc: docx.Document) -> int:
    """Remove pdf2docx placeholder tokens left inside paragraph runs.

    Skips runs that contain drawings or embedded objects to avoid destroying
    formula images or embedded content when clearing run text.
    """
    stripped = 0
    for para in doc.paragraphs:
        # Skip the entire paragraph if it contains any drawing/formula element.
        if _para_has_protected_content(para):
            continue
        changed = False
        for run in para.runs:
            # Skip runs that hold drawing elements — setting run.text would remove them.
            if _run_has_drawing(run):
                continue
            raw = run.text or ""
            if not raw:
                continue
            # Skip pure-whitespace runs — they may carry <w:br/> soft-break elements
            # that must not be destroyed (setting run.text calls clear_content() which
            # removes ALL child XML including <w:br/>).
            if not raw.strip():
                continue
            cleaned = _INLINE_ARTIFACT_RE.sub("", raw)
            cleaned = re.sub(r"^\s*[,.\s]+\s*", "", cleaned)
            if cleaned != raw:
                run.text = cleaned
                changed = True
        if changed:
            stripped += 1
    return stripped


def _paragraph_jc_val(paragraph) -> Optional[str]:
    ppr = paragraph._element.find(qn("w:pPr"))
    return _read_ppr_jc_val(ppr)


def _detect_title_block_end(paras: List) -> int:
    """Detect title/author block; body formatting starts after this index."""
    abs_i = _find_abstract_start(paras)
    if abs_i is not None:
        return abs_i
    sec1 = _find_section_one_start(paras)
    if sec1 is not None:
        return sec1
    end = 0
    for i, p in enumerate(paras[:24]):
        text = _paragraph_plain(p)
        if not text:
            end = max(end, i + 1)
            continue
        jc = _paragraph_jc_val(p) or ""
        if jc == "center" and len(text) <= 140 and _paragraph_word_count(p) <= 22:
            end = i + 1
            continue
        if i < 8 and len(text) <= 90:
            end = i + 1
            continue
        break
    return end


def _paragraph_is_centered(paragraph) -> bool:
    jc = _paragraph_jc_val(paragraph)
    if jc == "center":
        return True
    try:
        return paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        return False


def _detect_layout_regions(paras: List) -> Tuple[int, int, int, int, Optional[int]]:
    """Return title_end, abstract_start, abstract_end, body_start, sec2_start."""
    title_end = _detect_title_block_end(paras)
    abs_start = _find_abstract_start(paras)
    if abs_start is None:
        abs_start = title_end
    kw_idx = _find_keywords_paragraph_index(paras, title_end)
    sec1 = _find_section_one_start(paras)
    sec2 = _find_section_two_start(paras, min_index=(sec1 + 1) if sec1 is not None else 0)

    if kw_idx is not None:
        abstract_end = kw_idx + 1
    elif sec1 is not None:
        abstract_end = sec1
    elif sec2 is not None:
        abstract_end = sec2
    else:
        abstract_end = min(abs_start + 24, len(paras))

    body_start = sec1 if sec1 is not None else abstract_end
    return title_end, abs_start, abstract_end, body_start, sec2


def _emu_to_twips(value_emu: int) -> int:
    return int(int(value_emu) * 1440 / 914400)


def _get_content_width_twips(doc: docx.Document) -> int:
    try:
        sec = doc.sections[0]
        content_emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
        return max(3600, _emu_to_twips(content_emu))
    except Exception:
        return 9024


def _get_fallback_body_profile(doc: docx.Document) -> Dict[str, int]:
    left = _env_int("PDF_DOCX_BODY_INDENT_LEFT", 1054)
    right = _env_int("PDF_DOCX_BODY_INDENT_RIGHT", 1008)
    try:
        sec = doc.sections[0]
        lm = _emu_to_twips(int(sec.left_margin))
        rm = _emu_to_twips(int(sec.right_margin))
        if 400 <= lm <= 1800:
            left = lm
        if 400 <= rm <= 1800:
            right = rm
    except Exception:
        pass
    return {"left": left, "right": right}


def _is_abstract_style_indent(left: int, right: int) -> bool:
    return left >= 1350 and right >= 1350


def _is_narrow_column_indent(left: int, right: int, content_twips: int) -> bool:
    total = int(left) + int(right)
    if total <= 0:
        return False
    if _is_abstract_style_indent(left, right):
        return True
    return total > max(2200, int(content_twips * 0.28))


def _derive_layout_profiles_from_pdf(
    pdf_path: str,
    doc: docx.Document,
) -> Optional[Tuple[Dict[str, int], Dict[str, int]]]:
    if not pdf_path or not os.path.isfile(pdf_path):
        return None
    if not _env_bool("PDF_DOCX_PDF_GEOMETRY_PROFILE", True):
        return None

    formats = extract_pdf_paragraph_formats(pdf_path)
    if not formats:
        return None

    content_twips = _get_content_width_twips(doc)
    try:
        sec = doc.sections[0]
        lm_ratio = float(int(sec.left_margin)) / float(int(sec.page_width))
        rm_ratio = float(int(sec.right_margin)) / float(int(sec.page_width))
    except Exception:
        lm_ratio, rm_ratio = 0.121, 0.121
    content_ratio = max(0.45, 1.0 - lm_ratio - rm_ratio)

    body_indents: List[Tuple[int, int, float]] = []
    abstract_indents: List[Tuple[int, int, float]] = []

    for fmt in formats:
        pw = float(fmt.get("page_width") or 0)
        x0 = fmt.get("x0")
        x1 = fmt.get("x1")
        if pw <= 0 or x0 is None or x1 is None:
            continue

        left_pt = max(0.0, float(x0) - pw * lm_ratio)
        right_pt = max(0.0, pw * (1.0 - rm_ratio) - float(x1))
        width_pt = max(0.0, float(x1) - float(x0))
        width_ratio = width_pt / max(1.0, pw * content_ratio)

        left_twips = int(max(0.0, left_pt / max(1.0, pw * content_ratio)) * content_twips)
        right_twips = int(max(0.0, right_pt / max(1.0, pw * content_ratio)) * content_twips)

        if width_ratio >= 0.74 and left_twips <= 1800:
            body_indents.append((left_twips, right_twips, width_ratio))
        elif 0.52 <= width_ratio <= 0.78 and abs(left_twips - right_twips) <= 400:
            abstract_indents.append((left_twips, right_twips, width_ratio))

    if len(body_indents) < 3:
        return None

    body_indents.sort(key=lambda item: item[0] + item[1])
    wider = body_indents[: max(3, len(body_indents) // 2)]
    lefts = [item[0] for item in wider]
    rights = [item[1] for item in wider]
    body_profile = {
        "left": _median_int(lefts, 1054),
        "right": _median_int(rights, 1008),
    }

    if abstract_indents:
        abstract_indents.sort(key=lambda item: item[0] + item[1])
        mid = abstract_indents[len(abstract_indents) // 2]
        abs_left = int(mid[0])
        abs_right = int(mid[1])
        inset = max(abs_left, abs_right, int(body_profile["left"]) + _ABSTRACT_EXTRA_INDENT_TWIPS)
        abstract_profile = {"left": inset, "right": inset}
    else:
        abstract_profile = {
            "left": body_profile["left"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
            "right": body_profile["right"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        }
    return body_profile, abstract_profile


def _compute_indent_profile(
    paras: List,
    start: int,
    end: int,
    *,
    doc: Optional[docx.Document] = None,
    default_left: int = 1054,
    default_right: int = 1008,
    for_body: bool = True,
) -> Dict[str, int]:
    candidates: List[Tuple[int, int, int]] = []
    for p in paras[start:end]:
        text = _paragraph_plain(p)
        if len(text) < 25:
            continue
        if _is_section_heading_line(p) or _is_running_header_line(p):
            continue
        if _paragraph_is_centered(p):
            continue
        ppr = p._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))
        right = int(ind.get("right", 0))
        if left <= 0 and right <= 0:
            continue
        if for_body and _is_abstract_style_indent(left, right):
            continue
        if left + right > 4200:
            continue
        candidates.append((left + right, left, right))

    if not candidates:
        if doc is not None:
            return _get_fallback_body_profile(doc)
        return {"left": default_left, "right": default_right}

    candidates.sort(key=lambda item: item[0])
    wider = candidates[: max(1, len(candidates) // 2 + 1)]
    lefts = [item[1] for item in wider]
    rights = [item[2] for item in wider]
    return {
        "left": _median_int(lefts, default_left),
        "right": _median_int(rights, default_right),
    }


def _paragraph_needs_indent_fix(
    paragraph,
    *,
    region: str,
    body_profile: Dict[str, int],
    abstract_profile: Dict[str, int],
    content_twips: int,
) -> bool:
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    left = int(ind.get("left", 0))
    right = int(ind.get("right", 0))
    text = _paragraph_plain(paragraph)
    if len(text) < 15:
        return False

    if _paragraph_is_centered(paragraph) and len(text) >= 15:
        return True

    if region == "abstract":
        target_l = int(abstract_profile.get("left", 0))
        target_r = int(abstract_profile.get("right", 0))
        if left + right == 0 and len(text) >= 20:
            return True
        if abs(left - target_l) > 700 or abs(right - target_r) > 700:
            return left + right > 0
        return False

    if region in ("body", "section1"):
        if _is_abstract_style_indent(left, right):
            return True
        if _is_narrow_column_indent(left, right, content_twips):
            return True
        if left + right == 0 and len(text) >= 20 and not _is_section_heading_line(paragraph):
            return True
        if _is_indent_outlier(left, body_profile):
            return True
        if right == 0 and left > int(body_profile.get("left", 0)) + 900:
            return True
    return False


def _resolve_layout_profiles(
    doc: docx.Document,
    paras: List,
    *,
    pdf_path: Optional[str] = None,
) -> Tuple[Dict[str, int], Dict[str, int]]:
    pdf_profiles = _derive_layout_profiles_from_pdf(pdf_path, doc) if pdf_path else None
    _, abs_start, abstract_end, body_start, sec2 = _detect_layout_regions(paras)
    body_end = sec2 if sec2 is not None else len(paras)
    first_line_profile = _compute_first_line_profile(paras, body_start, body_end)

    if pdf_profiles:
        pdf_body, pdf_abstract = pdf_profiles
        if int(pdf_body.get("left", 0)) + int(pdf_body.get("right", 0)) <= 2400:
            pdf_body = dict(pdf_body)
            pdf_body["firstLine"] = first_line_profile
            return pdf_body, pdf_abstract

    body_profile = _compute_indent_profile(
        paras,
        body_start,
        body_end,
        doc=doc,
        for_body=True,
    )
    abstract_profile = _compute_indent_profile(
        paras,
        abs_start,
        abstract_end,
        doc=doc,
        default_left=body_profile["left"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        default_right=body_profile["right"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        for_body=False,
    )
    if abstract_profile["left"] <= body_profile["left"]:
        abstract_profile = {
            "left": body_profile["left"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
            "right": body_profile["right"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        }
    body_profile["firstLine"] = first_line_profile
    return body_profile, abstract_profile


def _read_first_line_twips(paragraph) -> int:
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    return int(ind.get("firstLine", 0))


def _paragraph_has_hanging_indent(paragraph) -> bool:
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    return int(ind.get("hanging", 0)) > 0


def _is_first_paragraph_after_heading(paras: List, idx: int, body_start: int) -> bool:
    if idx <= body_start:
        return True
    for j in range(idx - 1, max(body_start - 1, -1), -1):
        prev = paras[j]
        text = _paragraph_plain(prev)
        if not text:
            continue
        if _is_section_heading_line(prev) and _paragraph_word_count(prev) <= 10:
            return True
        return False
    return True


def _compute_first_line_profile(paras: List, start: int, end: int) -> int:
    values: List[int] = []
    first_after_heading = True
    for p in paras[start:end]:
        text = _paragraph_plain(p)
        if not text:
            continue
        if _is_section_heading_line(p) and _paragraph_word_count(p) <= 10:
            first_after_heading = True
            continue
        if first_after_heading:
            first_after_heading = False
            continue
        fl = _read_first_line_twips(p)
        if fl > 0:
            values.append(fl)
    return _median_int(values, 0)


def _resolve_body_first_line_twips(
    dst_para,
    src_para,
    body_profile: Dict[str, int],
    *,
    is_first_after_heading: bool,
) -> Optional[int]:
    """Resolve firstLine twips; None means keep existing hanging/firstLine untouched."""
    del body_profile
    if is_first_after_heading:
        return 0
    for para in (dst_para, src_para):
        if para is None:
            continue
        if _paragraph_has_hanging_indent(para):
            return None
        fl = _read_first_line_twips(para)
        if fl > 0:
            return fl
    return 0


def _apply_regional_layout_profiles(
    doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
    src_doc: Optional[docx.Document] = None,
) -> Dict[str, int]:
    """Apply title/abstract/body layout profiles to every paragraph (global PDF fix)."""
    stats = {"regional_layout_applied": 0}
    if not _env_bool("PDF_DOCX_REGIONAL_LAYOUT", True):
        return stats

    paras = list(doc.paragraphs)
    if not paras:
        return stats

    title_end, abs_start, abstract_end, body_start, _sec2 = _detect_layout_regions(paras)
    body_profile, abstract_profile = _resolve_layout_profiles(doc, paras, pdf_path=pdf_path)
    src_paras = list(src_doc.paragraphs) if src_doc is not None else []

    for i, p in enumerate(paras):
        if _is_in_table_cell(p):
            continue
        text = _paragraph_plain(p)
        if not text:
            continue
        src_p = src_paras[i] if i < len(src_paras) else None

        if i < title_end:
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.CENTER)
            _set_paragraph_indents(p, left=0, right=0, first_line=0)
            stats["regional_layout_applied"] += 1
            continue

        if _is_running_header_line(p):
            _compact_running_header(p, profile=body_profile)
            stats["regional_layout_applied"] += 1
            continue

        if _is_section_heading_line(p) and _paragraph_word_count(p) <= 10:
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.LEFT)
            _set_paragraph_indents(
                p,
                left=int(body_profile["left"]),
                right=0,
                first_line=0,
            )
            stats["regional_layout_applied"] += 1
            continue

        if abs_start <= i < abstract_end or (title_end <= i < body_start):
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            _set_paragraph_indents(
                p,
                left=int(abstract_profile["left"]),
                right=int(abstract_profile["right"]),
                first_line=0,
            )
            _clear_first_line_indent(p)
            stats["regional_layout_applied"] += 1
            continue

        if i >= body_start:
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            first_line = _resolve_body_first_line_twips(
                p,
                src_p,
                body_profile,
                is_first_after_heading=_is_first_paragraph_after_heading(paras, i, body_start),
            )
            indent_kwargs: Dict[str, int] = {
                "left": int(body_profile["left"]),
                "right": int(body_profile["right"]),
            }
            if first_line is not None:
                indent_kwargs["first_line"] = int(first_line)
            _set_paragraph_indents(p, **indent_kwargs)
            stats["regional_layout_applied"] += 1

    return stats


def normalize_converted_docx_layout_in_doc(
    doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
    src_doc: Optional[docx.Document] = None,
    analysis: Optional[Dict[str, Any]] = None,
) -> Dict[str, int]:
    """Fix pdf2docx center alignment and inconsistent indents before/after translation."""
    stats: Dict[str, int] = {"alignment_normalized": 0, "indents_normalized": 0}
    if not _env_bool("PDF_DOCX_NORMALIZE_ALIGN", True):
        return stats

    paras = list(doc.paragraphs)
    if not paras:
        return stats

    # Normal PDFs: keep source layout instead of heuristic normalization.
    if _should_preserve_source_layout(doc=doc, src_doc=src_doc, analysis=analysis):
        stats["source_layout_preserved"] = len(paras)
        return stats

    if _should_apply_regional_layout(doc=doc, src_doc=src_doc, analysis=analysis):
        regional = _apply_regional_layout_profiles(
            doc,
            pdf_path=pdf_path,
            src_doc=src_doc or doc,
        )
        stats["regional_layout_applied"] = int(regional.get("regional_layout_applied", 0))
        stats["indents_normalized"] = stats["regional_layout_applied"]
        stats["alignment_normalized"] = stats["regional_layout_applied"]
        return stats

    title_end, abs_start, abstract_end, body_start, sec2 = _detect_layout_regions(paras)
    force_justify = _env_bool("PDF_DOCX_FORCE_BODY_JUSTIFY", False)
    content_twips = _get_content_width_twips(doc)
    body_profile, abstract_profile = _resolve_layout_profiles(doc, paras, pdf_path=pdf_path)

    body_end = sec2 if sec2 is not None else len(paras)
    narrow_body_count = 0
    body_para_count = 0
    for i in range(body_start, body_end):
        text = _paragraph_plain(paras[i])
        if len(text) < 20:
            continue
        body_para_count += 1
        ppr = paras[i]._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))
        right = int(ind.get("right", 0))
        if _paragraph_is_centered(paras[i]) or _is_narrow_column_indent(left, right, content_twips):
            narrow_body_count += 1
    force_body_width = (
        body_para_count >= 2
        and narrow_body_count / max(1, body_para_count) >= 0.35
    )

    for i, p in enumerate(paras):
        if i < title_end:
            continue
        text = _paragraph_plain(p)
        if len(text) < 15:
            continue
        if _is_running_header_line(p):
            continue
        if _is_section_heading_line(p) and _paragraph_word_count(p) <= 10:
            continue

        if abs_start <= i < abstract_end:
            region = "abstract"
        elif i >= body_start:
            region = "body"
        else:
            region = "abstract"

        centered = _paragraph_is_centered(p)
        jc = _paragraph_jc_val(p) or ""
        if centered or (force_justify and jc in ("left", "start", "") and len(text) >= 30):
            if not (_is_section_heading_line(p) and _paragraph_word_count(p) <= 10):
                _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
                stats["alignment_normalized"] += 1

        if not _env_bool("PDF_DOCX_NORMALIZE_INDENTS", True):
            continue

        needs_fix = _paragraph_needs_indent_fix(
            p,
            region=region,
            body_profile=body_profile,
            abstract_profile=abstract_profile,
            content_twips=content_twips,
        )
        if force_body_width and region == "body" and len(text) >= 20:
            if not _is_section_heading_line(p) or _paragraph_word_count(p) > 10:
                needs_fix = True

        if not needs_fix:
            continue

        target = abstract_profile if region == "abstract" else body_profile
        ppr = p._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        first_line = 0 if region == "abstract" else int(ind.get("firstLine", 0))
        _set_paragraph_indents(
            p,
            left=int(target["left"]),
            right=int(target["right"]),
            first_line=first_line,
        )
        if region == "abstract":
            _clear_first_line_indent(p)
        stats["indents_normalized"] += 1

    return stats


def sanitize_converted_docx(
    docx_path: str,
    *,
    pdf_path: Optional[str] = None,
    analysis: Optional[Dict[str, Any]] = None,
) -> Dict[str, int]:
    """Clean pdf2docx noise + normalize body alignment/indents before translation."""
    if not os.path.isfile(docx_path):
        return {
            "noise_removed": 0,
            "alignment_normalized": 0,
            "indents_normalized": 0,
            "font_names_normalized": 0,
        }
    doc = docx.Document(docx_path)
    stats = strip_noise_paragraphs(doc)
    stats["font_names_normalized"] = _normalize_subset_font_names_in_doc(doc)
    stats["fonts_sanitized"] = sanitize_document_text_fonts(doc)
    if _env_bool("PDF_DOCX_SANITIZE_MERGE_FRAGMENTS", False):
        paras = list(doc.paragraphs)
        sec2 = _find_section_two_start(paras) or len(paras)
        stats["fragments_merged"] = _collapse_pre_section2_fragments(doc, paras, sec2)
    stats.update(normalize_converted_docx_layout_in_doc(doc, pdf_path=pdf_path, analysis=analysis))
    if any(int(v or 0) for v in stats.values()):
        doc.save(docx_path)
    return stats


def _find_title_block_end(paras: List) -> int:
    idx = _find_abstract_start(paras)
    return idx if idx is not None else 0


def _resolve_layout_region(
    idx: int,
    *,
    title_end: int,
    sec1_start: Optional[int],
    sec2_start: Optional[int],
) -> str:
    if idx < title_end:
        return "title"
    if sec2_start is not None and idx >= sec2_start:
        return "preserve"
    if sec1_start is not None and idx >= sec1_start:
        return "section1"
    return "abstract"


def _read_ppr_jc_val(ppr) -> Optional[str]:
    if ppr is None:
        return None
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        return None
    return jc.get(qn("w:val"))


def _read_ppr_ind_twips(ppr) -> Dict[str, int]:
    out: Dict[str, int] = {}
    if ppr is None:
        return out
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return out
    for key in ("left", "right", "firstLine", "hanging"):
        raw = ind.get(qn(f"w:{key}"))
        if raw is not None:
            try:
                out[key] = int(raw)
            except Exception:
                pass
    return out


def _read_paragraph_alignment(paragraph) -> Optional[WD_PARAGRAPH_ALIGNMENT]:
    try:
        return paragraph.alignment
    except Exception:
        return None


def _set_paragraph_alignment(paragraph, align: WD_PARAGRAPH_ALIGNMENT) -> None:
    try:
        paragraph.alignment = align
    except Exception:
        pass
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    val = {
        WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
        WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
        WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
        WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "both",
        WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: "distribute",
    }.get(align, "left")
    jc.set(qn("w:val"), val)


def _copy_paragraph_properties(src_ppr, dst_ppr) -> None:
    if src_ppr is None:
        return
    new_ppr = copy.deepcopy(src_ppr)
    parent = dst_ppr.getparent()
    if parent is None:
        return
    if dst_ppr is not None:
        parent.remove(dst_ppr)
    parent.insert(0, new_ppr)


def _copy_paragraph_layout_template(src_para, dst_para, *, include_spacing: bool = True) -> None:
    src_ppr = src_para._element.find(qn("w:pPr"))
    if src_ppr is None:
        return
    cloned = copy.deepcopy(src_ppr)
    if not include_spacing:
        spacing = cloned.find(qn("w:spacing"))
        if spacing is not None:
            cloned.remove(spacing)
    dst_ppr = dst_para._element.find(qn("w:pPr"))
    _copy_paragraph_properties(cloned, dst_ppr)


def _sync_paragraph_layout_like_inline(
    src_para,
    dst_para,
    *,
    include_spacing: bool = True,
) -> None:
    """Copy source paragraph layout, then downgrade pdf2docx false center on body lines.

    Inline bilingual uses this in recover_docx_layout; newline reuses the same rule.
    """
    _copy_paragraph_layout_template(src_para, dst_para, include_spacing=include_spacing)
    if _paragraph_is_centered(dst_para) and not _is_section_heading_line(dst_para):
        _set_paragraph_alignment(dst_para, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)


def _set_paragraph_indents(
    paragraph,
    *,
    left: Optional[int] = None,
    right: Optional[int] = None,
    first_line: Optional[int] = None,
    hanging: Optional[int] = None,
) -> None:
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if left is not None:
        ind.set(qn("w:left"), str(int(left)))
    if right is not None:
        ind.set(qn("w:right"), str(int(right)))
    if hanging is not None and int(hanging) > 0:
        ind.set(qn("w:hanging"), str(int(hanging)))
        ind.attrib.pop(qn("w:firstLine"), None)
    elif first_line is not None:
        if int(first_line) > 0:
            ind.set(qn("w:firstLine"), str(int(first_line)))
            ind.attrib.pop(qn("w:hanging"), None)
        else:
            ind.attrib.pop(qn("w:firstLine"), None)
            if hanging is None:
                ind.attrib.pop(qn("w:hanging"), None)


def _clear_first_line_indent(paragraph) -> None:
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return
    for key in (qn("w:firstLine"), qn("w:hanging")):
        if ind.get(key) is not None:
            ind.attrib.pop(key, None)


def _apply_font_pt(paragraph, pt: float) -> None:
    if pt <= 0:
        return
    for run in paragraph.runs:
        try:
            run.font.size = Pt(pt)
        except Exception:
            pass


def _score_body_reference_paragraph(paragraph) -> int:
    text = _paragraph_plain(paragraph)
    if not text or len(text) < 30 or _is_in_table_cell(paragraph):
        return -1
    if _is_section_heading_line(paragraph) or _is_running_header_line(paragraph):
        return -1
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    if ind.get("left", 0) <= 0:
        return -1
    return 1


def _find_body_reference_paragraph(paras: List, sec2_start: int) -> Optional[int]:
    for i in range(sec2_start + 1, min(sec2_start + 12, len(paras))):
        if _score_body_reference_paragraph(paras[i]) > 0:
            return i
    return None


def _find_abstract_reference_paragraph(paras: List, start: int, end: int) -> Optional[int]:
    for i in range(start, end):
        ppr = paras[i]._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        if ind.get("left", 0) >= 1500:
            return i
    return start


def _body_mode_font_pt(paras: List, sec2_start: int) -> Optional[float]:
    ref = _find_body_reference_paragraph(paras, sec2_start)
    if ref is None:
        return None
    return _dominant_font_pt(paras[ref])


def _set_paragraph_spacing(
    paragraph,
    *,
    line: Optional[int] = None,
    line_rule: Optional[str] = None,
    before: Optional[int] = None,
    after: Optional[int] = None,
) -> None:
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    if line is not None:
        spacing.set(qn("w:line"), str(int(line)))
    if line_rule is not None:
        spacing.set(qn("w:lineRule"), line_rule)
    if before is not None:
        spacing.set(qn("w:before"), str(int(before)))
    if after is not None:
        spacing.set(qn("w:after"), str(int(after)))


def _read_spacing_twips(paragraph) -> Dict[str, int]:
    out: Dict[str, int] = {}
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return out
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        return out
    for key in ("line", "before", "after"):
        raw = spacing.get(qn(f"w:{key}"))
        if raw is not None:
            try:
                out[key] = int(raw)
            except Exception:
                pass
    rule = spacing.get(qn("w:lineRule"))
    if rule:
        out["lineRule"] = rule  # type: ignore[assignment]
    return out


def _env_int(name: str, default: int) -> int:
    raw = os.getenv(name)
    if raw is None:
        return int(default)
    try:
        return int(str(raw).strip())
    except Exception:
        return int(default)


def _median_int(values: List[int], default: int = 0) -> int:
    if not values:
        return default
    ordered = sorted(int(v) for v in values)
    return ordered[len(ordered) // 2]


def _compute_body_layout_profile(paras: List) -> Dict[str, int]:
    """Median indent/spacing from body-like paragraphs (dynamic per document)."""
    lefts: List[int] = []
    rights: List[int] = []
    befores: List[int] = []
    afters: List[int] = []

    for paragraph in paras:
        text = _paragraph_plain(paragraph)
        if len(text) < 20:
            continue
        if _is_section_heading_line(paragraph) and _paragraph_word_count(paragraph) <= 8:
            continue
        ppr = paragraph._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        sp = _read_spacing_twips(paragraph)
        if ind.get("left") is not None:
            lefts.append(int(ind["left"]))
        if ind.get("right") is not None:
            rights.append(int(ind["right"]))
        if sp.get("before") is not None:
            befores.append(int(sp["before"]))
        if sp.get("after") is not None:
            afters.append(int(sp["after"]))

    return {
        "left": _median_int(lefts),
        "right": _median_int(rights),
        "before": _median_int(befores),
        "after": _median_int(afters),
    }


def _is_indent_outlier(value: int, profile: Dict[str, int]) -> bool:
    if value <= 0:
        return False
    hard_max = _env_int("PDF_DOCX_INDENT_OUTLIER_MAX", 2200)
    if value > hard_max:
        return True
    med = int(profile.get("left", 0))
    if med <= 0:
        return value > 1600
    return value > med + max(500, int(med * 0.65))


def _is_confident_academic_structure(paras: List) -> bool:
    """Only apply abstract/section heuristics when structure is clearly academic."""
    if not _env_bool("PDF_DOCX_ACADEMIC_LAYOUT", False):
        return False
    abs_idx = _find_abstract_start(paras)
    sec1 = _find_section_one_start(paras)
    sec2 = _find_section_two_start(paras, min_index=(sec1 + 1) if sec1 is not None else 0)
    if abs_idx is None or sec1 is None or sec2 is None:
        return False
    if not (abs_idx < sec1 < sec2):
        return False
    abs_text = _paragraph_plain(paras[abs_idx]) if abs_idx < len(paras) else ""
    if not _ABSTRACT_START_RE.match(abs_text):
        return False
    body_ref_idx = _find_body_reference_paragraph(paras, sec2)
    if body_ref_idx is None:
        return False
    body_ind = _read_ppr_ind_twips(paras[body_ref_idx]._element.find(qn("w:pPr")))
    if _is_indent_outlier(int(body_ind.get("left", 0)), {"left": 900}):
        return False
    return True


def _apply_abstract_layout(
    paragraph,
    *,
    body_ref,
    abstract_ref,
    spacing_before: Optional[int] = None,
) -> None:
    _copy_paragraph_layout_template(abstract_ref, paragraph, include_spacing=False)

    body_ppr = body_ref._element.find(qn("w:pPr"))
    abs_ppr = abstract_ref._element.find(qn("w:pPr"))
    body_ind = _read_ppr_ind_twips(body_ppr)
    abs_ind = _read_ppr_ind_twips(abs_ppr)

    body_left = int(body_ind.get("left", 0))
    body_right = int(body_ind.get("right", 0))
    abs_left = abs_ind.get("left")
    if abs_left is None:
        abs_left = body_left + _ABSTRACT_EXTRA_INDENT_TWIPS
    else:
        abs_left = int(abs_left)
    max_abs_left = body_left + max(_ABSTRACT_EXTRA_INDENT_TWIPS, 800)
    if abs_left > max_abs_left or _is_indent_outlier(abs_left, {"left": body_left}):
        abs_left = body_left + _ABSTRACT_EXTRA_INDENT_TWIPS
    abs_right = abs_ind.get("right")
    if abs_right is None:
        abs_right = body_right + (_ABSTRACT_EXTRA_INDENT_TWIPS if body_right else 0)
    else:
        abs_right = int(abs_right)

    abs_sp = _read_spacing_twips(abstract_ref)
    _set_paragraph_spacing(
        paragraph,
        line=abs_sp.get("line", 242),
        line_rule=str(abs_sp.get("lineRule", "exact")),
        before=spacing_before if spacing_before is not None else abs_sp.get("before", 0),
        after=abs_sp.get("after", 0),
    )
    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    _set_paragraph_indents(
        paragraph,
        left=abs_left,
        right=abs_right,
        first_line=0,
    )
    _clear_first_line_indent(paragraph)


def _apply_section1_body_layout(paragraph, body_ref, *, first_para: bool = False) -> None:
    _copy_paragraph_layout_template(body_ref, paragraph, include_spacing=False)
    body_ppr = body_ref._element.find(qn("w:pPr"))
    body_ind = _read_ppr_ind_twips(body_ppr)
    body_sp = _read_spacing_twips(body_ref)

    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    _set_paragraph_indents(
        paragraph,
        left=int(body_ind.get("left", 0)),
        right=int(body_ind.get("right", 0)),
        first_line=0 if first_para else int(body_ind.get("firstLine", 0)),
    )
    _set_paragraph_spacing(
        paragraph,
        line=body_sp.get("line", 240),
        line_rule=str(body_sp.get("lineRule", "exact")),
        before=body_sp.get("before", 0) if first_para else min(int(body_sp.get("before", 0) or 0), 120),
        after=0,
    )


def _compact_running_header(paragraph, profile: Optional[Dict[str, int]] = None) -> None:
    profile = profile or {}
    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT)
    _set_paragraph_spacing(paragraph, line=240, line_rule="exact", before=0, after=0)
    _set_paragraph_indents(
        paragraph,
        left=int(profile.get("left", 0)),
        right=int(profile.get("right", 0)),
        first_line=0,
    )


_FLAG_ITALIC = 1 << 1
_FLAG_BOLD = 1 << 4


def _normalize_match_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\r", " ").replace("\n", " ").strip().lower())


def _copy_run_rpr(src_run, dst_run) -> None:
    try:
        src_rpr = src_run._element.find(qn("w:rPr"))
        if src_rpr is None:
            return
        new_rpr = copy.deepcopy(src_rpr)

        # Normalize subset font prefixes and avoid applying symbol fonts to regular text runs.
        rfonts = new_rpr.find(qn("w:rFonts"))
        _normalize_rfonts_attrs(rfonts)
        if rfonts is not None:
            is_symbol_font = False
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                try:
                    val = rfonts.get(qn(attr))
                except Exception:
                    val = None
                if val and _SYMBOL_FONT_HINT_RE.search(str(val)):
                    is_symbol_font = True
                    break

            dst_text = ""
            try:
                dst_text = dst_run.text or ""
            except Exception:
                dst_text = ""

            has_text_tokens = bool(re.search(r"[A-Za-zÀ-ỹ0-9]", dst_text, flags=re.UNICODE))
            has_list_marker = bool(re.search(r"[+\-•*]", dst_text))

            # pdf2docx often stores bullets/markers in symbol fonts; when those
            # formats are copied to translated runs, markers can render as tofu boxes.
            # Force text fallback not only for alphanumeric content, but also for
            # common list markers.
            if is_symbol_font and (has_text_tokens or has_list_marker):
                fallback = _pdf_docx_text_font_fallback()
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rfonts.set(qn(attr), fallback)

        _canonicalize_rfonts_in_rpr(new_rpr, dst_text)

        dst_rpr = dst_run._element.find(qn("w:rPr"))
        if dst_rpr is not None:
            dst_run._element.remove(dst_rpr)
        dst_run._element.insert(0, new_rpr)
    except Exception:
        pass


def _dominant_source_run(runs: List) -> Optional[Any]:
    if not runs:
        return None
    best = None
    best_score = -10**9
    for run in runs:
        text = run.text or ""
        score = len(re.findall(r"[A-Za-zÀ-ỹ0-9]", text, flags=re.UNICODE)) * 10 + len(text.strip())
        try:
            rpr = run._element.find(qn("w:rPr"))
            rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
            if rfonts is not None:
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    val = rfonts.get(qn(attr))
                    if val and _SYMBOL_FONT_HINT_RE.search(str(val)):
                        score -= 10000
                        break
        except Exception:
            pass

        if best is None or score > best_score:
            best = run
            best_score = score
    return best


def _preserve_run_formats(src_para, dst_para) -> None:
    src_runs = list(src_para.runs)
    dst_runs = list(dst_para.runs)
    if not src_runs or not dst_runs:
        return
    if len(src_runs) == len(dst_runs):
        for src_run, dst_run in zip(src_runs, dst_runs):
            _copy_run_rpr(src_run, dst_run)
        return
    dominant = _dominant_source_run(src_runs)
    if dominant is None:
        return
    for dst_run in dst_runs:
        _copy_run_rpr(dominant, dst_run)


def _preserve_paragraph_style(src_para, dst_para) -> None:
    try:
        src_style = getattr(getattr(src_para, "style", None), "name", None)
        if src_style and src_style != "Normal":
            dst_para.style = src_style
    except Exception:
        pass


def _pdf_hint_raw_align_key(hint: Optional[Dict[str, Any]]) -> str:
    if not hint:
        return ""
    return str(hint.get("align") or "left").lower()


def _pdf_line_is_geometrically_centered(
    hint: Optional[Dict[str, Any]],
    *,
    max_width_ratio: float = 0.72,
) -> bool:
    """True when PDF text block is short and sits near page horizontal center."""
    if not hint or _pdf_hint_raw_align_key(hint) != "center":
        return False
    pw = float(hint.get("page_width") or 0)
    x0 = hint.get("x0")
    x1 = hint.get("x1")
    if x0 is None or x1 is None or pw <= 0:
        return True
    width = max(0.0, float(x1) - float(x0))
    if width >= pw * max_width_ratio:
        return False
    mid = (float(x0) + float(x1)) / 2.0
    return abs(mid - pw / 2.0) <= pw * 0.12


def _should_preserve_center_alignment(
    paragraph,
    para_index: int,
    title_end: int,
    pdf_hint: Optional[Dict[str, Any]],
) -> bool:
    """Keep center only for lines that are genuinely centered in the source PDF."""
    text = _paragraph_plain(paragraph)
    if not text.strip():
        return False
    if _is_section_heading_line(paragraph) and re.match(r"^\d+\.", text.strip()):
        return False
    if len(text) >= 45 or _paragraph_word_count(paragraph) >= 8:
        return False

    raw_align = _pdf_hint_raw_align_key(pdf_hint)
    if raw_align in ("left", "right", "justify", "both"):
        return False

    geo_center = _pdf_line_is_geometrically_centered(pdf_hint)
    if geo_center:
        if para_index < title_end:
            return len(text) <= 160 and _paragraph_word_count(paragraph) <= 22
        return len(text) <= 100 and _paragraph_word_count(paragraph) <= 15

    if para_index < title_end:
        return bool(
            _paragraph_is_centered(paragraph)
            and len(text) <= 160
            and _paragraph_word_count(paragraph) <= 22
            and raw_align != "left"
        )
    return False


def _lookup_pdf_hint_for_text(
    text: str,
    src_paras: List,
    pdf_hints: List[Optional[Dict[str, Any]]],
    *,
    bilingual_delimiter: str = "|",
) -> Optional[Dict[str, Any]]:
    """Match a dst paragraph back to a src index for PDF alignment hints."""
    raw = (text or "").strip()
    if not raw:
        return None
    d = (bilingual_delimiter or "|").strip()
    if d and d in raw:
        raw = raw.split(d, 1)[0].strip()
    norm = _normalize_match_text(raw)
    if not norm:
        return None
    best_i: Optional[int] = None
    best_score = 0.0
    for i, src_p in enumerate(src_paras):
        src_norm = _normalize_match_text(_paragraph_plain(src_p))
        if not src_norm:
            continue
        if norm == src_norm:
            return pdf_hints[i] if i < len(pdf_hints) else None
        if norm in src_norm or src_norm in norm:
            score = min(len(norm), len(src_norm)) / max(len(norm), len(src_norm), 1)
        else:
            a = set(norm.split())
            b = set(src_norm.split())
            if not a or not b:
                continue
            score = len(a & b) / max(len(a), len(b))
        if score > best_score:
            best_score = score
            best_i = i
    if best_i is not None and best_score >= 0.45:
        return pdf_hints[best_i] if best_i < len(pdf_hints) else None
    return None


def _fix_spurious_center_alignment_in_doc(
    doc: docx.Document,
    src_doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
    bilingual_delimiter: str = "|",
) -> int:
    """Force left/justify on body paragraphs; pdf2docx often marks them centered."""
    src_paras = list(src_doc.paragraphs)
    title_end, _, _, _, _ = _detect_layout_regions(src_paras)
    pdf_hints: List[Optional[Dict[str, Any]]] = []
    if pdf_path and os.path.isfile(pdf_path):
        pdf_formats = extract_pdf_paragraph_formats(pdf_path)
        src_texts = [_paragraph_plain(p) for p in src_paras]
        pdf_hints = _match_pdf_formats_to_paragraphs(src_texts, pdf_formats)

    fixed = 0
    src_cursor = 0
    for i, para in enumerate(doc.paragraphs):
        if _is_in_table_cell(para):
            continue
        text = _paragraph_plain(para)
        if not text.strip() or not _paragraph_is_centered(para):
            continue
        hint = _lookup_pdf_hint_for_text(
            text,
            src_paras,
            pdf_hints,
            bilingual_delimiter=bilingual_delimiter,
        )
        para_index = min(src_cursor, max(0, len(src_paras) - 1))
        if _should_preserve_center_alignment(para, para_index, title_end, hint):
            src_cursor = min(src_cursor + 1, len(src_paras) - 1)
            continue

        align = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        pdf_align = _alignment_enum_from_pdf_hint(hint)
        if pdf_align and pdf_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
            align = pdf_align
        elif _is_section_heading_line(para) or re.match(r"^\d+\.", text.strip()):
            align = WD_PARAGRAPH_ALIGNMENT.LEFT

        ppr = para._element.find(qn("w:pPr"))
        jc = ppr.find(qn("w:jc")) if ppr is not None else None
        if jc is not None:
            ppr.remove(jc)
        _clear_paragraph_indents_and_tabs(para)
        _set_paragraph_alignment(para, align)
        if align != WD_PARAGRAPH_ALIGNMENT.CENTER:
            _set_paragraph_indents(para, left=0, right=0, first_line=0, hanging=0)
        fixed += 1
        src_cursor = min(src_cursor + 1, len(src_paras) - 1)
    return fixed


def _clear_paragraph_indents_and_tabs(paragraph) -> None:
    """Remove paragraph indents/tab stops only (keep alignment jc intact)."""
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return
    for tag in ("w:ind", "w:tabs"):
        el = ppr.find(qn(tag))
        if el is not None:
            ppr.remove(el)


def _paragraph_has_false_pdf2docx_body_offset(
    paragraph,
    *,
    content_twips: int,
    text: Optional[str] = None,
) -> bool:
    """Detect margin-as-indent offsets on body lines (not genuine title center)."""
    plain = (text or _paragraph_plain(paragraph) or "").strip()
    if not plain:
        return False
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    left = int(ind.get("left", 0))
    right = int(ind.get("right", 0))
    first_line = int(ind.get("firstLine", 0))
    hanging = int(ind.get("hanging", 0))
    if first_line >= 360 or hanging >= 360:
        return len(plain) >= 20
    if left >= 360 or right >= 360:
        if len(plain) < 20 and _paragraph_is_centered(paragraph):
            return True
        return len(plain) >= 20
    if _is_abstract_style_indent(left, right) or _is_narrow_column_indent(left, right, content_twips):
        return len(plain) >= 20
    # pdf2docx stores page margin (~1054 twips) as paragraph left indent on body text.
    if 400 <= left <= 1800 and right <= 360 and len(plain) >= 30:
        return True
    return False


def _apply_standard_newline_paragraph_layout(
    paragraph,
    *,
    src_i: int = 0,
    title_end: int = 0,
    abs_start: int = 0,
    abstract_end: int = 0,
    body_start: int = 0,
    is_heading: bool = False,
    pdf_hint: Optional[Dict[str, Any]] = None,
    content_twips: int = 9024,
    src_para=None,
) -> bool:
    """Fix pdf2docx false center/indents; preserve genuine PDF-centered title lines."""
    text = _paragraph_plain(paragraph)
    if not text.strip():
        return False

    preserve_center = _should_preserve_center_alignment(
        paragraph, src_i, title_end, pdf_hint
    )
    false_offset = _paragraph_has_false_pdf2docx_body_offset(
        paragraph, content_twips=content_twips, text=text
    )

    if preserve_center:
        if not false_offset:
            return False
        _clear_paragraph_indents_and_tabs(paragraph)
        _set_paragraph_indents(paragraph, left=0, right=0, first_line=0, hanging=0)
        if not _paragraph_is_centered(paragraph):
            _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)
        return True

    if is_heading or _is_section_heading_line(paragraph) or re.match(r"^\d+\.", text.strip()):
        ppr = paragraph._element.find(qn("w:pPr"))
        jc = ppr.find(qn("w:jc")) if ppr is not None else None
        if jc is not None:
            ppr.remove(jc)
        _clear_paragraph_indents_and_tabs(paragraph)
        _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT)
        _set_paragraph_indents(paragraph, left=0, right=0, first_line=0, hanging=0)
        return True

    needs_fix = (
        false_offset
        or _docx_center_is_likely_pdf2docx_artifact(paragraph, text)
        or (
            _paragraph_is_centered(paragraph)
            and len(text) >= 12
            and not preserve_center
        )
    )
    if not needs_fix:
        return False

    ref = src_para if src_para is not None else paragraph
    align = _resolve_newline_bilingual_alignment(
        ref,
        paragraph,
        src_i=src_i,
        pdf_hint=pdf_hint,
        abs_start=abs_start,
        abstract_end=abstract_end,
        body_start=body_start,
        title_end=title_end,
    )

    ppr = paragraph._element.find(qn("w:pPr"))
    jc = ppr.find(qn("w:jc")) if ppr is not None else None
    if jc is not None:
        ppr.remove(jc)
    _clear_paragraph_indents_and_tabs(paragraph)
    _set_paragraph_alignment(paragraph, align)
    if align != WD_PARAGRAPH_ALIGNMENT.CENTER:
        _set_paragraph_indents(paragraph, left=0, right=0, first_line=0, hanging=0)
    return True


def _paragraph_needs_newline_layout_reset(
    paragraph,
    *,
    title_end: int,
    src_i: int,
    content_twips: int,
    pdf_hint: Optional[Dict[str, Any]] = None,
) -> bool:
    text = _paragraph_plain(paragraph)
    if not text.strip():
        return False

    preserve_center = _should_preserve_center_alignment(
        paragraph, src_i, title_end, pdf_hint
    )
    false_offset = _paragraph_has_false_pdf2docx_body_offset(
        paragraph, content_twips=content_twips, text=text
    )
    if preserve_center:
        return false_offset

    if _paragraph_is_centered(paragraph) and len(text) >= 12:
        return True
    if len(text) < 12:
        return false_offset
    return false_offset or _docx_center_is_likely_pdf2docx_artifact(paragraph, text)


def apply_post_split_newline_layout(orig_para, trans_para) -> None:
    """After inline→split: keep layout like inline; alignment pass runs in recovery."""
    if orig_para is not None and trans_para is not None:
        _tighten_newline_pair_spacing(orig_para, trans_para)


def _alignment_enum_from_pdf_hint(
    hint: Optional[Dict[str, Any]],
    *,
    fallback: Optional[WD_PARAGRAPH_ALIGNMENT] = None,
) -> Optional[WD_PARAGRAPH_ALIGNMENT]:
    """Map PyMuPDF line hint to Word alignment; downgrade false center on wide lines."""
    if not hint:
        return fallback
    align_map = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }
    align_key = str(hint.get("align") or "left").lower()
    pw = float(hint.get("page_width") or 0)
    x0 = hint.get("x0")
    x1 = hint.get("x1")
    text_len = len(str(hint.get("text") or ""))
    if x0 is not None and x1 is not None and pw > 0:
        width = max(0.0, float(x1) - float(x0))
        if align_key == "center" and width >= pw * 0.38:
            align_key = "justify"
        if align_key == "center" and (width >= pw * 0.28 or text_len >= 45):
            align_key = "left"
    return align_map.get(align_key, fallback)


def _docx_center_is_likely_pdf2docx_artifact(paragraph, text: str) -> bool:
    """pdf2docx often marks full-width body paragraphs as centered."""
    if not _paragraph_is_centered(paragraph):
        return False
    plain = (text or "").strip()
    if len(plain) >= 40 or _paragraph_word_count(paragraph) >= 8:
        return True
    if _is_section_heading_line(paragraph):
        return False
    return len(plain) >= 20


def _resolve_newline_bilingual_alignment(
    src_p,
    orig_dst,
    *,
    src_i: int,
    pdf_hint: Optional[Dict[str, Any]],
    abs_start: int,
    abstract_end: int,
    body_start: int,
    title_end: int,
    prefer_justify: bool = True,
) -> WD_PARAGRAPH_ALIGNMENT:
    """Pick alignment for stacked bilingual pairs; PDF hint beats pdf2docx jc."""
    text = _paragraph_plain(orig_dst)

    if _should_preserve_center_alignment(orig_dst, src_i, title_end, pdf_hint):
        return WD_PARAGRAPH_ALIGNMENT.CENTER

    pdf_align = _alignment_enum_from_pdf_hint(pdf_hint)
    if pdf_align is not None and pdf_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
        return pdf_align

    if _is_section_heading_line(orig_dst) and _paragraph_word_count(orig_dst) <= 10:
        return WD_PARAGRAPH_ALIGNMENT.LEFT

    if _docx_center_is_likely_pdf2docx_artifact(orig_dst, text):
        return WD_PARAGRAPH_ALIGNMENT.JUSTIFY if prefer_justify else WD_PARAGRAPH_ALIGNMENT.LEFT

    jc = _paragraph_jc_val(orig_dst) or _paragraph_jc_val(src_p) or ""
    if jc == "right":
        return WD_PARAGRAPH_ALIGNMENT.RIGHT
    if jc in ("both", "justify", "distribute"):
        return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if jc == "center":
        return WD_PARAGRAPH_ALIGNMENT.JUSTIFY if prefer_justify else WD_PARAGRAPH_ALIGNMENT.LEFT

    return WD_PARAGRAPH_ALIGNMENT.JUSTIFY if prefer_justify else WD_PARAGRAPH_ALIGNMENT.LEFT


def _apply_pdf_format_hint(paragraph, hint: Dict[str, Any], *, apply_alignment: bool = False) -> None:
    if not hint:
        return
    if apply_alignment:
        align = _alignment_enum_from_pdf_hint(hint)
        if align is not None:
            _set_paragraph_alignment(paragraph, align)

    font_size = hint.get("font_size")
    if font_size:
        try:
            _apply_font_pt(paragraph, float(font_size))
        except Exception:
            pass

    bold = hint.get("bold")
    italic = hint.get("italic")
    if bold is None and italic is None:
        return
    for run in paragraph.runs:
        try:
            if bold is not None:
                run.bold = bool(bold)
            if italic is not None:
                run.italic = bool(italic)
        except Exception:
            pass


def extract_pdf_paragraph_formats(pdf_path: str) -> List[Dict[str, Any]]:
    """Extract per-line formatting hints from the PDF text layer."""
    if not pdf_path or not os.path.isfile(pdf_path):
        return []
    try:
        import fitz  # PyMuPDF
    except Exception:
        return []

    formats: List[Dict[str, Any]] = []
    doc = fitz.open(pdf_path)
    try:
        for page in doc:
            try:
                page_width = float(page.rect.width)
            except Exception:
                page_width = 0.0
            try:
                blocks = page.get_text("dict").get("blocks") or []
            except Exception:
                blocks = []

            for block in blocks:
                if int(block.get("type") or 0) != 0:
                    continue
                for line in block.get("lines") or []:
                    spans = line.get("spans") or []
                    if not spans:
                        continue
                    text = "".join(str(sp.get("text") or "") for sp in spans).strip()
                    if not text:
                        continue

                    sizes = [float(sp.get("size") or 0) for sp in spans if sp.get("text")]
                    font_size = max(sizes) if sizes else None
                    total_chars = sum(len(str(sp.get("text") or "")) for sp in spans)
                    bold_chars = sum(
                        len(str(sp.get("text") or ""))
                        for sp in spans
                        if int(sp.get("flags") or 0) & _FLAG_BOLD
                    )
                    italic_chars = sum(
                        len(str(sp.get("text") or ""))
                        for sp in spans
                        if int(sp.get("flags") or 0) & _FLAG_ITALIC
                    )

                    align = "left"
                    x0: Optional[float] = None
                    x1: Optional[float] = None
                    bbox = line.get("bbox") or block.get("bbox")
                    if bbox and page_width > 0:
                        try:
                            x0, _, x1, _ = [float(v) for v in bbox[:4]]
                        except Exception:
                            x0 = x1 = 0.0
                        width = max(0.0, x1 - x0)
                        left_m = max(0.0, x0)
                        right_m = max(0.0, page_width - x1)
                        center_x = (x0 + x1) * 0.5
                        mid = page_width * 0.5
                        sym_tol = max(8.0, page_width * 0.022)
                        cx_tol = max(6.0, page_width * 0.012)
                        if (
                            width <= page_width * 0.52
                            and abs(left_m - right_m) <= sym_tol
                            and abs(center_x - mid) <= cx_tol
                        ):
                            align = "center"
                        elif x1 >= page_width - max(24.0, page_width * 0.08) * 0.45 and center_x >= mid + page_width * 0.025:
                            align = "right"
                        elif width >= page_width * 0.58 and abs(left_m - right_m) <= max(12.0, page_width * 0.03):
                            align = "justify"

                    formats.append(
                        {
                            "text": text,
                            "text_norm": _normalize_match_text(text),
                            "font_size": font_size,
                            "bold": bool(total_chars and bold_chars * 100 >= total_chars * 60),
                            "italic": bool(total_chars and italic_chars * 100 >= total_chars * 50),
                            "align": align,
                            "font": spans[0].get("font") if spans else None,
                            "x0": x0,
                            "x1": x1,
                            "page_width": page_width,
                        }
                    )
    finally:
        doc.close()
    return formats


def _match_pdf_formats_to_paragraphs(
    para_texts: List[str],
    pdf_formats: List[Dict[str, Any]],
) -> List[Optional[Dict[str, Any]]]:
    results: List[Optional[Dict[str, Any]]] = [None] * len(para_texts)
    fmt_idx = 0
    for i, text in enumerate(para_texts):
        norm = _normalize_match_text(text)
        if not norm:
            continue
        best_j: Optional[int] = None
        best_score = 0.0
        search_end = min(len(pdf_formats), fmt_idx + 10)
        for j in range(fmt_idx, search_end):
            fmt_norm = pdf_formats[j].get("text_norm") or ""
            if not fmt_norm:
                continue
            if norm in fmt_norm or fmt_norm in norm:
                score = min(len(norm), len(fmt_norm)) / max(len(norm), len(fmt_norm), 1)
            else:
                a = set(norm.split())
                b = set(fmt_norm.split())
                if not a or not b:
                    continue
                score = len(a & b) / max(len(a), len(b))
            if score > best_score:
                best_score = score
                best_j = j
        if best_j is not None and best_score >= 0.35:
            results[i] = pdf_formats[best_j]
            fmt_idx = best_j + 1
    return results


def sync_formats_from_pdf(
    source_docx: str,
    translated_docx: str,
    pdf_path: str,
) -> Dict[str, int]:
    stats = {"pdf_formats_applied": 0}
    if not _env_bool("PDF_DOCX_PDF_FORMAT_SYNC", True):
        return stats
    if not os.path.isfile(source_docx) or not os.path.isfile(translated_docx):
        return stats

    pdf_formats = extract_pdf_paragraph_formats(pdf_path)
    if not pdf_formats:
        return stats

    src = docx.Document(source_docx)
    dst = docx.Document(translated_docx)
    src_texts = [_paragraph_plain(p) for p in src.paragraphs]
    hints = _match_pdf_formats_to_paragraphs(src_texts, pdf_formats)

    for i, hint in enumerate(hints):
        if hint is None or i >= len(dst.paragraphs):
            continue
        _apply_pdf_format_hint(
            dst.paragraphs[i],
            hint,
            apply_alignment=_env_bool("PDF_DOCX_PDF_ALIGN_SYNC", False),
        )
        stats["pdf_formats_applied"] += 1

    if stats["pdf_formats_applied"]:
        dst.save(translated_docx)
    return stats


def _sync_paragraph_style_only(src_para, dst_para) -> None:
    """Copy paragraph style only; leave indents/alignment for regional layout pass."""
    _preserve_paragraph_style(src_para, dst_para)


def _sync_paragraph_properties_only(src_para, dst_para) -> None:
    """Copy paragraph-level layout only; keep translated run text untouched."""
    _preserve_paragraph_style(src_para, dst_para)
    src_ppr = src_para._element.find(qn("w:pPr"))
    if src_ppr is None:
        return
    dst_ppr = dst_para._element.find(qn("w:pPr"))
    _copy_paragraph_properties(src_ppr, dst_ppr)


def _preserve_paragraph_layout(src_para, dst_para) -> None:
    _sync_paragraph_properties_only(src_para, dst_para)


def _append_runs_from_paragraph(src_para, dst_para, *, spacer: str = " ") -> None:
    src_runs = list(src_para.runs)
    if not src_runs:
        return
    if _paragraph_plain(dst_para):
        try:
            dst_para.add_run(spacer)
        except Exception:
            pass
    for run in src_runs:
        new_run = dst_para.add_run(run.text or "")
        try:
            if run.bold is not None:
                new_run.bold = run.bold
            if run.italic is not None:
                new_run.italic = run.italic
            if run.font.size is not None:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
        except Exception:
            pass


def _sync_region_by_offset(
    src_paras: List,
    dst_paras: List,
    src_start: int,
    dst_start: int,
) -> int:
    if src_start < 0 or dst_start < 0:
        return 0
    synced = 0
    max_k = min(len(src_paras) - src_start, len(dst_paras) - dst_start)
    for k in range(max(0, max_k)):
        _preserve_paragraph_layout(src_paras[src_start + k], dst_paras[dst_start + k])
        synced += 1
    return synced


def _should_merge_fragment(prev_para, next_para) -> bool:
    # Never merge paragraphs that contain drawings or math — merging loses the embedded objects.
    if _para_has_protected_content(prev_para) or _para_has_protected_content(next_para):
        return False
    t1 = _paragraph_plain(prev_para)
    t2 = _paragraph_plain(next_para)
    if not t1 or not t2:
        return False
    if _is_section_heading_line(prev_para) or _is_section_heading_line(next_para):
        return False
    if _is_running_header_line(prev_para) or _is_running_header_line(next_para):
        return False
    if _SECTION1_HEAD_RE.match(t2):
        return False
    if _SECTION2_HEAD_RE.match(t2) and not _FALSE_SECTION2_RE.match(t2) and _paragraph_word_count(next_para) <= 8:
        return False
    if _KEYWORDS_LINE_RE.match(t2) and not _KEYWORDS_LINE_RE.match(t1):
        return False
    if _ABSTRACT_START_RE.match(t1) and _KEYWORDS_LINE_RE.match(t2):
        return False
    if _KEYWORDS_LINE_RE.match(t1) and (_SECTION1_HEAD_RE.match(t2) or re.match(r"^1\s+\S", t2)):
        return False
    if _paragraph_is_centered(prev_para) or _paragraph_is_centered(next_para):
        if _is_short_label_line(next_para) and re.match(r"^\d+\s+\S", t2):
            return False
        return True
    if not re.search(r"[.!?:;]\s*$", t1) and len(t2.split()) >= 3:
        return True
    return False


def _collapse_paragraph_range(
    doc: docx.Document,
    paras: List,
    start: int,
    end: int,
    *,
    allow_heading_merge: bool = False,
) -> int:
    merged = 0
    i = max(0, start)
    end = min(end, len(paras))
    while i < end - 1 and i < len(paras) - 1:
        cur = paras[i]
        nxt = paras[i + 1]
        if not allow_heading_merge and (
            _is_section_heading_line(cur) or _is_section_heading_line(nxt)
        ):
            i += 1
            continue
        if not _should_merge_fragment(cur, nxt):
            i += 1
            continue
        _append_runs_from_paragraph(nxt, cur)
        try:
            nxt._element.getparent().remove(nxt._element)
        except Exception:
            break
        paras.pop(i + 1)
        end = min(end, len(paras))
        merged += 1
    return merged


def _collapse_pre_section2_fragments(doc: docx.Document, paras: List, sec2_start: int) -> int:
    if sec2_start <= 0:
        return 0

    merged = 0
    title_end = _find_title_block_end(paras)
    if title_end > 1:
        merged += _collapse_title_block_fragments(doc, paras, title_end)
        title_end = _find_title_block_end(paras)
    kw_idx = _find_keywords_paragraph_index(paras, title_end)
    sec1_start = _find_section_one_start(paras)

    abstract_end = kw_idx if kw_idx is not None else (sec1_start if sec1_start is not None else sec2_start)
    if abstract_end > title_end:
        merged += _collapse_paragraph_range(doc, paras, title_end, abstract_end)

    title_end = _find_title_block_end(paras)
    kw_idx = _find_keywords_paragraph_index(paras, title_end)
    sec1_start = _find_section_one_start(paras)
    sec2_start = _find_section_two_start(paras) or sec2_start

    if kw_idx is not None:
        kw_end = sec1_start if sec1_start is not None else sec2_start
        if kw_end > kw_idx:
            merged += _collapse_paragraph_range(doc, paras, kw_idx, kw_end)

    sec1_start = _find_section_one_start(paras)
    sec2_start = _find_section_two_start(paras) or sec2_start

    if sec1_start is not None and sec2_start > sec1_start + 1:
        merged += _collapse_paragraph_range(
            doc,
            paras,
            sec1_start + 1,
            sec2_start,
            allow_heading_merge=False,
        )

    return merged


def _normalize_title_block(
    paras: List,
    src_paras: List,
    *,
    title_end: int,
    src_title_end: int,
) -> int:
    limit = min(title_end, src_title_end, len(paras), len(src_paras))
    for i in range(limit):
        _preserve_paragraph_layout(src_paras[i], paras[i])
    return limit


def _normalize_body_block(
    paras: List,
    src_paras: List,
    *,
    start: int,
    end: int,
    body_ref,
    abstract_ref,
    region: str,
    src_heading_para=None,
    dst_heading_idx: Optional[int] = None,
    src_abs_start: Optional[int] = None,
    src_kw_idx: Optional[int] = None,
    src_sec1_start: Optional[int] = None,
) -> int:
    fixed = 0
    if src_heading_para is not None and dst_heading_idx is not None:
        if 0 <= dst_heading_idx < len(paras):
            _preserve_paragraph_layout(src_heading_para, paras[dst_heading_idx])
            fixed += 1

    first_section1_body = True
    for i in range(start, min(end, len(paras))):
        if dst_heading_idx is not None and i == dst_heading_idx:
            continue
        para = paras[i]
        if _is_running_header_line(para):
            _compact_running_header(para)
            fixed += 1
            continue
        if _is_section_heading_line(para):
            if src_heading_para is not None and i == dst_heading_idx:
                fixed += 1
                continue
            src_i = (
                src_sec1_start
                if region == "section1" and src_sec1_start is not None
                else i
            )
            if src_i < len(src_paras):
                _preserve_paragraph_layout(src_paras[src_i], para)
            fixed += 1
            continue

        if region == "abstract":
            text = _paragraph_plain(para)
            before = None
            if _KEYWORDS_LINE_RE.match(text) and src_kw_idx is not None and src_kw_idx < len(src_paras):
                before = _read_spacing_twips(src_paras[src_kw_idx]).get("before")
            elif src_abs_start is not None and src_abs_start < len(src_paras):
                before = _read_spacing_twips(src_paras[src_abs_start]).get("before")
            _apply_abstract_layout(
                para,
                body_ref=body_ref,
                abstract_ref=abstract_ref,
                spacing_before=before,
            )
        else:
            _apply_section1_body_layout(para, body_ref, first_para=first_section1_body)
            first_section1_body = False
        fixed += 1
    return fixed


def _refresh_layout_after_merge(
    paras: List,
    src_paras: List,
    *,
    title_end: int,
    src_title_end: int,
    sec1_start: Optional[int],
    src_sec1_start: Optional[int],
    sec2_start: Optional[int],
    src_sec2_start: Optional[int],
    body_ref,
    abstract_ref,
) -> int:
    fixed = 0
    sec1_end = sec2_start if sec2_start is not None else len(paras)
    abstract_end = sec1_start if sec1_start is not None else sec1_end

    fixed += _normalize_title_block(
        paras,
        src_paras,
        title_end=title_end,
        src_title_end=src_title_end,
    )
    if body_ref is not None and abstract_ref is not None:
        src_abs = src_title_end
        src_kw = _find_keywords_paragraph_index(src_paras, src_title_end)
        fixed += _normalize_body_block(
            paras,
            src_paras,
            start=title_end,
            end=abstract_end,
            body_ref=body_ref,
            abstract_ref=abstract_ref,
            region="abstract",
            src_abs_start=src_abs,
            src_kw_idx=src_kw,
        )
    if body_ref is not None and sec1_start is not None:
        src_heading = (
            src_paras[src_sec1_start]
            if src_sec1_start is not None and src_sec1_start < len(src_paras)
            else None
        )
        fixed += _normalize_body_block(
            paras,
            src_paras,
            start=sec1_start,
            end=sec1_end,
            body_ref=body_ref,
            abstract_ref=abstract_ref or body_ref,
            region="section1",
            src_heading_para=src_heading,
            dst_heading_idx=sec1_start,
            src_sec1_start=src_sec1_start,
        )
    return fixed


def normalize_document_layout(
    doc: docx.Document,
    src_doc: docx.Document,
) -> Dict[str, int]:
    stats = {
        "alignment_fixed": 0,
        "indents_fixed": 0,
        "fonts_normalized": 0,
        "title_centered": 0,
        "fragments_merged": 0,
    }
    paras = list(doc.paragraphs)
    src_paras = list(src_doc.paragraphs)

    if not _is_confident_academic_structure(src_paras):
        return stats

    src_title_end = _find_title_block_end(src_paras)
    src_sec1_start = _find_section_one_start(src_paras)
    src_sec2_start = _find_section_two_start(src_paras)

    if src_sec2_start is None:
        return stats

    body_ref_idx = _find_body_reference_paragraph(src_paras, src_sec2_start)
    if body_ref_idx is None:
        return stats
    body_ref = src_paras[body_ref_idx]

    abs_start = src_title_end
    abs_end = src_sec1_start if src_sec1_start is not None else src_sec2_start
    abs_ref_idx = _find_abstract_reference_paragraph(src_paras, abs_start, abs_end)
    abstract_ref = src_paras[abs_ref_idx if abs_ref_idx is not None else abs_start]

    dst_sec2_start = _find_section_two_start(paras)
    if _env_bool("PDF_DOCX_MERGE_FRAGMENTS", False):
        stats["fragments_merged"] = _collapse_pre_section2_fragments(
            doc,
            paras,
            dst_sec2_start if dst_sec2_start is not None else len(paras),
        )
        paras = list(doc.paragraphs)

    title_end = _find_title_block_end(paras)
    sec1_start = _find_section_one_start(paras)
    sec2_start = _find_section_two_start(paras)

    fixed = _refresh_layout_after_merge(
        paras,
        src_paras,
        title_end=title_end,
        src_title_end=src_title_end,
        sec1_start=sec1_start,
        src_sec1_start=src_sec1_start,
        sec2_start=sec2_start,
        src_sec2_start=src_sec2_start,
        body_ref=body_ref,
        abstract_ref=abstract_ref,
    )
    stats["alignment_fixed"] = fixed
    stats["indents_fixed"] = fixed
    stats["fonts_normalized"] = fixed
    stats["title_centered"] = min(title_end, len(paras))
    return stats


def _sync_table_layout_from_source(src_doc: docx.Document, dst_doc: docx.Document) -> int:
    synced = 0
    if not _env_bool("PDF_DOCX_TABLE_SYNC", True):
        return 0
    for t_idx, src_table in enumerate(src_doc.tables):
        if t_idx >= len(dst_doc.tables):
            break
        dst_table = dst_doc.tables[t_idx]
        for r_idx, src_row in enumerate(src_table.rows):
            if r_idx >= len(dst_table.rows):
                break
            dst_row = dst_table.rows[r_idx]
            for c_idx, src_cell in enumerate(src_row.cells):
                if c_idx >= len(dst_row.cells):
                    break
                dst_cell = dst_row.cells[c_idx]
                src_cell_paras = list(src_cell.paragraphs)
                dst_cell_paras = list(dst_cell.paragraphs)
                for i in range(min(len(src_cell_paras), len(dst_cell_paras))):
                    _preserve_paragraph_layout(src_cell_paras[i], dst_cell_paras[i])
                    synced += 1
    return synced


def _sync_header_footer_layout(src_doc: docx.Document, dst_doc: docx.Document) -> int:
    synced = 0
    if not _env_bool("PDF_DOCX_HEADER_FOOTER_SYNC", True):
        return 0
    try:
        for i, src_sec in enumerate(src_doc.sections):
            if i >= len(dst_doc.sections):
                break
            dst_sec = dst_doc.sections[i]
            for src_p, dst_p in zip(src_sec.header.paragraphs, dst_sec.header.paragraphs):
                _preserve_paragraph_layout(src_p, dst_p)
                synced += 1
            for src_p, dst_p in zip(src_sec.footer.paragraphs, dst_sec.footer.paragraphs):
                _preserve_paragraph_layout(src_p, dst_p)
                synced += 1
    except Exception:
        pass
    return synced


def _align_paragraph_pairs(src_paras: List, dst_paras: List) -> List[Tuple[int, int]]:
    """Pair src/dst paragraphs sequentially, skipping empty-only rows on either side."""
    pairs: List[Tuple[int, int]] = []
    si = 0
    di = 0
    while si < len(src_paras) and di < len(dst_paras):
        src_t = _paragraph_plain(src_paras[si])
        dst_t = _paragraph_plain(dst_paras[di])
        if not src_t and not dst_t:
            pairs.append((si, di))
            si += 1
            di += 1
            continue
        if not src_t:
            si += 1
            continue
        if not dst_t:
            di += 1
            continue
        pairs.append((si, di))
        si += 1
        di += 1
    return pairs


def mirror_document_layout_from_source(
    source_docx: str,
    translated_docx: str,
) -> Dict[str, int]:
    """Copy paragraph/table/header properties from source DOCX onto translated text."""
    stats = {
        "mirrored": 0,
        "mismatched_paragraphs": 0,
        "table_cells_synced": 0,
        "header_footer_synced": 0,
    }
    if not _env_bool("PDF_DOCX_MIRROR_LAYOUT", True) and not _env_bool("PDF_DOCX_LAYOUT_SYNC", True):
        return stats

    src = docx.Document(source_docx)
    dst = docx.Document(translated_docx)

    pair_count = min(len(src.paragraphs), len(dst.paragraphs))
    stats["mismatched_paragraphs"] = abs(len(src.paragraphs) - len(dst.paragraphs))
    for i in range(pair_count):
        _sync_paragraph_properties_only(src.paragraphs[i], dst.paragraphs[i])
        stats["mirrored"] += 1

    stats["table_cells_synced"] = _sync_table_layout_from_source(src, dst)
    stats["header_footer_synced"] = _sync_header_footer_layout(src, dst)
    stats["mirrored"] += stats["table_cells_synced"] + stats["header_footer_synced"]
    dst.save(translated_docx)
    return stats


def sync_docx_layout_from_source(
    source_docx: str,
    translated_docx: str,
) -> Dict[str, int]:
    """Backward-compatible alias: full mirror from source DOCX."""
    mirror_stats = mirror_document_layout_from_source(source_docx, translated_docx)
    return {
        "paragraphs_synced": int(mirror_stats.get("mirrored", 0)),
        "alignment_fixed": int(mirror_stats.get("mirrored", 0)),
        "mismatched_paragraphs": int(mirror_stats.get("mismatched_paragraphs", 0)),
        "table_cells_synced": int(mirror_stats.get("table_cells_synced", 0)),
        "header_footer_synced": int(mirror_stats.get("header_footer_synced", 0)),
    }


def _fit_paragraph_fonts(paragraph, *, min_scale: float, min_pt: float) -> int:
    changed = 0
    text = _paragraph_plain(paragraph)
    if len(text) < 20:
        return 0
    if _is_in_table_cell(paragraph):
        return 0
    base = _dominant_font_pt(paragraph)
    if base is None or base <= min_pt:
        return 0
    est_chars = max(1, int(base * 2.2))
    if len(text) <= est_chars:
        return 0
    ratio = max(min_scale, est_chars / float(len(text)))
    new_pt = max(min_pt, base * ratio)
    if new_pt < base - 0.2:
        _apply_font_pt(paragraph, new_pt)
        changed += 1
    return changed


def _fit_document_fonts(doc: docx.Document) -> int:
    if not _env_bool("PDF_DOCX_FONT_FIT", False):
        return 0
    min_scale = _env_float("PDF_DOCX_FONT_FIT_MIN_SCALE", 0.62)
    min_pt = _env_float("PDF_DOCX_FONT_FIT_MIN_PT", 7.0)
    fitted = 0
    for paragraph in iter_all_paragraphs(doc):
        fitted += _fit_paragraph_fonts(
            paragraph,
            min_scale=min_scale,
            min_pt=min_pt,
        )
    return fitted


def _recover_tables_and_images(doc: docx.Document) -> Dict[str, int]:
    out = {"table_rows_relaxed": 0, "images_resized": 0}
    try:
        from docx.enum.table import WD_ROW_HEIGHT_RULE
    except Exception:
        WD_ROW_HEIGHT_RULE = None  # type: ignore

    if WD_ROW_HEIGHT_RULE is not None:
        for table in doc.tables:
            for row in table.rows:
                row_changed = False
                try:
                    if row.height is not None:
                        row.height = None
                        row_changed = True
                except Exception:
                    pass
                try:
                    if row.height_rule != WD_ROW_HEIGHT_RULE.AT_LEAST:
                        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                        row_changed = True
                except Exception:
                    pass
                if row_changed:
                    out["table_rows_relaxed"] += 1

    max_inline_width = None
    try:
        if doc.sections:
            sec = doc.sections[0]
            max_inline_width = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
    except Exception:
        max_inline_width = None

    if max_inline_width and max_inline_width > 0:
        for shp in getattr(doc, "inline_shapes", []):
            try:
                cur_w = int(shp.width or 0)
                cur_h = int(shp.height or 0)
            except Exception:
                continue
            if cur_w <= 0 or cur_h <= 0 or cur_w <= max_inline_width:
                continue
            try:
                ratio = float(max_inline_width) / float(cur_w)
                shp.width = int(max_inline_width)
                shp.height = max(1, int(cur_h * ratio))
                out["images_resized"] += 1
            except Exception:
                pass
    return out


def _fix_indents_from_pdf_geometry(
    doc: docx.Document,
    src_doc: docx.Document,
    pdf_path: str,
    profile: Dict[str, int],
) -> int:
    """Reset DOCX indents when PDF geometry shows normal left-margin body text."""
    pdf_formats = extract_pdf_paragraph_formats(pdf_path)
    if not pdf_formats:
        return 0

    body_x0: List[float] = []
    for fmt in pdf_formats:
        pw = float(fmt.get("page_width") or 0)
        x0 = fmt.get("x0")
        x1 = fmt.get("x1")
        if pw <= 0 or x0 is None or x1 is None:
            continue
        width = float(x1) - float(x0)
        if width >= pw * 0.42 and float(x0) / pw <= 0.18:
            body_x0.append(float(x0))
    if not body_x0:
        return 0

    body_x0.sort()
    median_x0 = body_x0[len(body_x0) // 2]
    fixed = 0
    src_texts = [_paragraph_plain(p) for p in src_doc.paragraphs]
    hints = _match_pdf_formats_to_paragraphs(src_texts, pdf_formats)

    for i, hint in enumerate(hints):
        if not hint or i >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[i]
        if not _paragraph_plain(para) or _is_in_table_cell(para):
            continue

        pw = float(hint.get("page_width") or 0)
        x0 = hint.get("x0")
        x1 = hint.get("x1")
        if pw <= 0 or x0 is None or x1 is None:
            continue
        width = float(x1) - float(x0)
        near_body = abs(float(x0) - median_x0) <= pw * 0.06
        full_width = width >= pw * 0.38
        if not (near_body and full_width):
            continue

        ppr = para._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))
        right = int(ind.get("right", 0))
        content_twips = _get_content_width_twips(doc)
        wrong_width = (
            _is_indent_outlier(left, profile)
            or _is_abstract_style_indent(left, right)
            or _is_narrow_column_indent(left, right, content_twips)
        )
        if not wrong_width:
            continue

        preserved_first = int(ind.get("firstLine", 0))
        if _paragraph_has_hanging_indent(para):
            _set_paragraph_indents(
                para,
                left=int(profile.get("left", 0)),
                right=int(profile.get("right", 0)),
            )
        else:
            _set_paragraph_indents(
                para,
                left=int(profile.get("left", 0)),
                right=int(profile.get("right", 0)),
                first_line=preserved_first,
            )
        align = str(hint.get("align") or "left").lower()
        if align == "center" and width >= pw * 0.38:
            align = "justify"
        align_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        _set_paragraph_alignment(para, align_map.get(align, WD_PARAGRAPH_ALIGNMENT.JUSTIFY))
        fixed += 1
    return fixed


def _collapse_excessive_blank_paragraphs(doc: docx.Document) -> int:
    """Remove stacked empty paragraphs that create large vertical gaps.

    Paragraphs that contain drawings, OMML formulas, or embedded objects are
    never removed even when their visible text is empty.
    """
    removed = 0
    paras = list(doc.paragraphs)
    blank_streak = 0
    idx = 0
    while idx < len(paras):
        para = paras[idx]
        has_text = bool(_paragraph_plain(para))
        # Paragraphs with drawings/math are treated as non-blank regardless of text.
        if not has_text and not _para_has_protected_content(para):
            blank_streak += 1
            if blank_streak > 1:
                try:
                    para._element.getparent().remove(para._element)
                    paras.pop(idx)
                    removed += 1
                    continue
                except Exception:
                    pass
        else:
            blank_streak = 0
        idx += 1
    return removed


def fix_layout_anomalies(
    doc: docx.Document,
    src_doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
) -> Dict[str, int]:
    """Detect and fix outlier indents/spacing relative to each document's profile."""
    stats = {
        "anomalies_fixed": 0,
        "spacing_fixed": 0,
        "pdf_indent_fixed": 0,
        "empties_collapsed": 0,
    }
    if not _env_bool("PDF_DOCX_FIX_ANOMALIES", True):
        return stats

    src_body = [
        p
        for p in src_doc.paragraphs
        if _paragraph_plain(p) and not _is_in_table_cell(p) and not _paragraph_is_centered(p)
    ]
    profile = _compute_body_layout_profile(src_body)

    for i, para in enumerate(doc.paragraphs):
        if not _paragraph_plain(para) or _is_in_table_cell(para) or _paragraph_is_centered(para):
            continue

        ppr = para._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))

        if _is_indent_outlier(left, profile):
            restored = False
            if i < len(src_doc.paragraphs):
                src_para = src_doc.paragraphs[i]
                src_ind = _read_ppr_ind_twips(src_para._element.find(qn("w:pPr")))
                if not _is_indent_outlier(int(src_ind.get("left", 0)), profile):
                    _preserve_paragraph_layout(src_para, para)
                    restored = True
            if not restored:
                _set_paragraph_indents(
                    para,
                    left=int(profile.get("left", 0)),
                    right=int(profile.get("right", 0)),
                    first_line=0,
                )
            stats["anomalies_fixed"] += 1

        sp = _read_spacing_twips(para)
        med_before = int(profile.get("before", 0))
        med_after = int(profile.get("after", 0))
        cap_before = max(280, med_before * 4, 600)
        cap_after = max(280, med_after * 4, 600)
        before = sp.get("before")
        after = sp.get("after")
        new_before, new_after = before, after
        if before is not None and int(before) > cap_before:
            new_before = med_before
        if after is not None and int(after) > cap_after:
            new_after = med_after
        if new_before != before or new_after != after:
            _set_paragraph_spacing(para, before=new_before, after=new_after)
            stats["spacing_fixed"] += 1

    if pdf_path and os.path.isfile(pdf_path):
        stats["pdf_indent_fixed"] = _fix_indents_from_pdf_geometry(doc, src_doc, pdf_path, profile)

    stats["empties_collapsed"] = _collapse_excessive_blank_paragraphs(doc)
    return stats


def _normalize_paragraph_match_key(text: str) -> str:
    t = re.sub(r"\s+", " ", (text or "").replace("\u00ad", "").strip().lower())
    return t


def _paragraph_texts_match(a: str, b: str) -> bool:
    ka = _normalize_paragraph_match_key(a)
    kb = _normalize_paragraph_match_key(b)
    if not ka or not kb:
        return False
    if ka == kb:
        return True
    ka_compact = re.sub(r"\s+", "", ka)
    kb_compact = re.sub(r"\s+", "", kb)
    if ka_compact == kb_compact:
        return True
    shorter, longer = (ka, kb) if len(ka) <= len(kb) else (kb, ka)
    prefix_len = min(len(shorter), 120)
    if prefix_len >= 24 and longer.startswith(shorter[:prefix_len]):
        return True
    return False


def _is_affiliation_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if not text:
        return False
    if _AFFILIATION_LINE_RE.match(text):
        return True
    if re.match(
        r"^\d+\s*(faculty|department|university|springer|institute|fakulteit|khoa|tr[uườ]ng)\b",
        text,
        re.IGNORECASE,
    ):
        return True
    low = text.lower()
    return bool(
        re.match(r"^\d+\s*", text)
        and any(k in low for k in ("faculty", "department", "university", "institute", "khoa", "trường", "truong"))
    )


def _looks_like_subtitle_line(text: str) -> bool:
    t = (text or "").strip()
    if not t or "," in t or "@" in t:
        return False
    words = [w for w in t.split() if w]
    if not 1 <= len(words) <= 5:
        return False
    return all(w[0].isupper() for w in words if w and w[0].isalpha())


def _should_merge_title_block_fragment(prev_para, next_para) -> bool:
    if _para_has_protected_content(prev_para) or _para_has_protected_content(next_para):
        return False
    t1 = _paragraph_plain(prev_para).strip()
    t2 = _paragraph_plain(next_para).strip()
    if not t1 or not t2:
        return False
    if _ABSTRACT_START_RE.match(t2) or _ABSTRACT_START_RE.match(t1):
        return False
    if _KEYWORDS_LINE_RE.match(t2):
        return False

    if _looks_like_subtitle_line(t2) and _paragraph_word_count(prev_para) >= 5:
        return False

    if _is_affiliation_line(prev_para) or re.match(r"^\d+\s*\S", t1):
        low1 = t1.lower()
        if any(k in low1 for k in ("faculty", "department", "university", "institute", "street", "ward", "city")):
            if not re.match(
                r"^\d+\s*(faculty|department|university|institute|khoa|tr[uườ]ng)\b",
                t2,
                re.IGNORECASE,
            ):
                if "@" not in t2:
                    return True

    if "@" in t1:
        return "@" in t2 or bool(re.match(r"^[a-z0-9._-]+@", t2, re.IGNORECASE))
    if "@" in t2:
        return False

    if re.search(r",\s*[\w\-]{1,12}$", t1) and not re.match(r"^\d", t2):
        first = (t2.split() or [""])[0]
        if first and first[0].isupper():
            return True

    if _paragraph_is_centered(prev_para) and _paragraph_is_centered(next_para):
        low1 = t1.lower()
        if " and " in low1 or re.search(r"\d+\*?\s*$", t1):
            return False
        if re.match(r"^\d", t2):
            return False
        if len(t2.split()) >= 2 and not _looks_like_subtitle_line(t2):
            return True

    return False


def _collapse_title_block_fragments(doc: docx.Document, paras: List, title_end: int) -> int:
    """Merge pdf2docx line wraps inside title/author/affiliation/email block."""
    merged = 0
    i = 0
    end = min(max(0, title_end), len(paras))
    while i < end - 1 and i < len(paras) - 1:
        cur = paras[i]
        nxt = paras[i + 1]
        if not _should_merge_title_block_fragment(cur, nxt):
            i += 1
            continue
        _append_runs_from_paragraph(nxt, cur, spacer=" ")
        try:
            nxt._element.getparent().remove(nxt._element)
        except Exception:
            break
        paras.pop(i + 1)
        end = min(end, len(paras))
        merged += 1
    return merged


def _apply_body_indent_from_source(
    paragraph,
    src_para,
    body_profile: Dict[str, int],
    *,
    is_first_after_heading: bool,
) -> None:
    """Apply body left/right + first-line/hanging indent from source paragraph."""
    src_ind = _read_ppr_ind_twips(src_para._element.find(qn("w:pPr")))
    left = int(body_profile["left"])
    right = int(body_profile["right"])
    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    if is_first_after_heading:
        _set_paragraph_indents(paragraph, left=left, right=right, first_line=0, hanging=0)
        return

    first_line = int(src_ind.get("firstLine", 0) or 0)
    hanging = int(src_ind.get("hanging", 0) or 0)
    if first_line <= 0 and hanging <= 0:
        profile_fl = int(body_profile.get("firstLine", 0) or 0)
        if profile_fl > 0:
            first_line = profile_fl

    if hanging > 0:
        _set_paragraph_indents(
            paragraph,
            left=max(left, int(src_ind.get("left", left) or left)),
            right=right,
            hanging=hanging,
        )
    else:
        _set_paragraph_indents(
            paragraph,
            left=left,
            right=right,
            first_line=first_line,
        )


def _rpr_font_size(rpr) -> int:
    if rpr is None:
        return -1
    try:
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            raw = str(sz.get(qn("w:val"), "") or "").strip()
            if raw.isdigit():
                return int(raw)
    except Exception:
        pass
    return -1


def _run_is_superscript_run(run) -> bool:
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return False
        va = rpr.find(qn("w:vertAlign"))
        if va is not None:
            v = str(va.get(qn("w:val"), "") or "").strip().lower()
            return v in ("superscript", "subscript")
    except Exception:
        pass
    return False


def _dominant_body_run_rpr(paragraph):
    best_rpr = None
    best_sz = -1
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        if _run_is_superscript_run(run):
            continue
        rpr = run._element.find(qn("w:rPr"))
        sz = _rpr_font_size(rpr)
        if sz > best_sz:
            best_sz = sz
            best_rpr = rpr
    return best_rpr


def _merge_run_font_rpr(dst_run, src_rpr) -> bool:
    if src_rpr is None:
        return False
    try:
        r_el = dst_run._element
        dst_rpr = r_el.find(qn("w:rPr"))
        if dst_rpr is None:
            dst_rpr = OxmlElement("w:rPr")
            r_el.insert(0, dst_rpr)
        for tag in ("w:rFonts", "w:sz", "w:szCs"):
            src_el = src_rpr.find(qn(tag))
            if src_el is None:
                continue
            old = dst_rpr.find(qn(tag))
            if old is not None:
                dst_rpr.remove(old)
            dst_rpr.append(copy.deepcopy(src_el))
        _canonicalize_rfonts_in_rpr(dst_rpr, dst_run.text or "")
        for tag in ("w:caps", "w:smallCaps"):
            old = dst_rpr.find(qn(tag))
            if old is not None:
                dst_rpr.remove(old)
        return True
    except Exception:
        return False


def _sync_translation_layout_from_orig(orig_para, trans_para) -> None:
    """Mirror finalized orig paragraph alignment/indents onto the translation line."""
    jc = _paragraph_jc_val(orig_para) or "left"
    jc_map = {
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "both": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "distribute": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "start": WD_PARAGRAPH_ALIGNMENT.LEFT,
    }
    _set_paragraph_alignment(trans_para, jc_map.get(jc, WD_PARAGRAPH_ALIGNMENT.JUSTIFY))
    orig_ind = _read_ppr_ind_twips(orig_para._element.find(qn("w:pPr")))
    left = int(orig_ind.get("left", 0) or 0)
    right = int(orig_ind.get("right", 0) or 0)
    first_line = int(orig_ind.get("firstLine", 0) or 0)
    hanging = int(orig_ind.get("hanging", 0) or 0)
    if hanging > 0:
        _set_paragraph_indents(trans_para, left=left, right=right, hanging=hanging)
    else:
        _set_paragraph_indents(trans_para, left=left, right=right, first_line=first_line)


def _mirror_paragraph_indents_from_orig(orig_para, trans_para) -> None:
    """Copy alignment + indents from original paragraph to translation line."""
    jc = _paragraph_jc_val(orig_para) or "justify"
    if _docx_center_is_likely_pdf2docx_artifact(orig_para, _paragraph_plain(orig_para)):
        jc = "both"
    jc_map = {
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "both": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "distribute": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }
    _set_paragraph_alignment(trans_para, jc_map.get(jc, WD_PARAGRAPH_ALIGNMENT.JUSTIFY))
    orig_ind = _read_ppr_ind_twips(orig_para._element.find(qn("w:pPr")))
    left = int(orig_ind.get("left", 0) or 0)
    right = int(orig_ind.get("right", 0) or 0)
    first_line = int(orig_ind.get("firstLine", 0) or 0)
    hanging = int(orig_ind.get("hanging", 0) or 0)
    if hanging > 0:
        _set_paragraph_indents(trans_para, left=left, right=right, hanging=hanging)
    else:
        _set_paragraph_indents(trans_para, left=left, right=right, first_line=first_line)


def _sync_translation_run_font_from_orig(orig_para, trans_para) -> int:
    """Match translation run font size/name to dominant body run in original."""
    template = _dominant_body_run_rpr(orig_para)
    if template is None:
        return 0
    synced = 0
    for run in trans_para.runs:
        if not (run.text or "").strip():
            continue
        if _merge_run_font_rpr(run, template):
            synced += 1
    return synced


def _finalize_newline_translation_pair(orig_para, trans_para) -> None:
    """Ensure translation line mirrors original paragraph layout and font."""
    _mirror_paragraph_indents_from_orig(orig_para, trans_para)
    _sync_translation_run_font_from_orig(orig_para, trans_para)


# ── Newline bilingual: pair spacing ──────────────────────────────────────────

def _tighten_newline_pair_spacing(orig_para, trans_para) -> None:
    """Visually attach translation paragraph to its source: zero gap between them,
    full gap after the translation to separate from the next source paragraph.

    This makes source+translation feel like a unit while preserving document flow.
    """
    try:
        src_sp = _read_spacing_twips(orig_para)
        src_after = int(src_sp.get("after", 120))
        _set_paragraph_spacing(orig_para, after=0)
        _set_paragraph_spacing(trans_para, before=0, after=max(120, src_after))
    except Exception:
        pass


# ── Inline bilingual → 2-column table conversion ─────────────────────────────

def _find_inline_bilingual_split(p_elem: Any, delimiter: str):
    """Return (ppr_el, src_runs, trans_runs) if p_elem contains the bilingual delimiter.

    _append_inline_bilingual appends ' {delimiter} ' to the last source run and then
    adds the translation as a separate run.  We locate that suffix and split.
    Returns None if no delimiter pattern is found.
    """
    d = (delimiter or "|").strip()
    if not d:
        return None
    suffix_spaced = f" {d} "   # most common: "text | "
    suffix_bare   = f" {d}"    # edge case: "text |" at end of run

    r_elements = p_elem.findall(qn("w:r"))
    if not r_elements:
        return None

    # Fast path: find the last run ending with the delimiter suffix.
    for split_idx in range(len(r_elements) - 1, -1, -1):
        run = r_elements[split_idx]
        run_text = "".join(el.text or "" for el in run.findall(qn("w:t")))
        if run_text.endswith(suffix_spaced) or run_text.endswith(suffix_bare):
            trans_runs = r_elements[split_idx + 1:]
            if not trans_runs:
                return None
            # Strip delimiter from a deep-copied version of that run.
            new_run = copy.deepcopy(run)
            t_els = new_run.findall(qn("w:t"))
            if t_els:
                last_t = t_els[-1]
                txt = last_t.text or ""
                if txt.endswith(suffix_spaced):
                    txt = txt[: -len(suffix_spaced)]
                elif txt.endswith(suffix_bare):
                    txt = txt[: -len(suffix_bare)]
                last_t.text = txt.rstrip()
                if last_t.text:
                    last_t.set(qn("xml:space"), "preserve")
            ppr = p_elem.find(qn("w:pPr"))
            src_runs = list(r_elements[:split_idx]) + [new_run]
            return ppr, src_runs, list(trans_runs)

    # Slow path: concatenate all run text and search for delimiter.
    full_text = ""
    positions: List[Tuple[int, int, int]] = []
    for i, run in enumerate(r_elements):
        run_text = "".join(el.text or "" for el in run.findall(qn("w:t")))
        positions.append((i, len(full_text), len(full_text) + len(run_text)))
        full_text += run_text

    idx = full_text.rfind(suffix_spaced)
    if idx < 0:
        return None

    split_run_idx = None
    for r_i, c_start, c_end in positions:
        if c_start <= idx < c_end:
            split_run_idx = r_i
            break
    if split_run_idx is None:
        return None

    trans_runs = r_elements[split_run_idx + 1:]
    if not trans_runs:
        return None

    new_run = copy.deepcopy(r_elements[split_run_idx])
    t_els = new_run.findall(qn("w:t"))
    if t_els:
        r_start = positions[split_run_idx][1]
        keep_len = idx - r_start
        combined = "".join(el.text or "" for el in t_els)
        kept = combined[:keep_len].rstrip()
        for j, t_el in enumerate(t_els):
            t_el.text = kept if j == 0 else ""
            if j == 0 and kept:
                t_el.set(qn("xml:space"), "preserve")

    ppr = p_elem.find(qn("w:pPr"))
    src_runs = list(r_elements[:split_run_idx]) + [new_run]
    return ppr, src_runs, list(trans_runs)


def _build_bilingual_two_col_table(
    ppr_elem: Any,
    src_runs: List,
    trans_runs: List,
    *,
    left_w: int,
    right_w: int,
    italic_translation: bool = True,
) -> Any:
    """Build a <w:tbl> element: left cell = source runs, right cell = translation runs."""

    # ── tbl ──
    tbl = OxmlElement("w:tbl")

    tbl_pr = OxmlElement("w:tblPr")

    tbl_w_el = OxmlElement("w:tblW")
    tbl_w_el.set(qn("w:w"), str(left_w + right_w))
    tbl_w_el.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w_el)

    tbl_borders = OxmlElement("w:tblBorders")
    for bname in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{bname}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

    tbl_cell_mar = OxmlElement("w:tblCellMar")
    for side, val in [("top", "0"), ("left", "72"), ("bottom", "0"), ("right", "72")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tbl_cell_mar.append(m)
    tbl_pr.append(tbl_cell_mar)
    tbl.append(tbl_pr)

    tbl_grid = OxmlElement("w:tblGrid")
    for w in (left_w, right_w):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tbl_grid.append(gc)
    tbl.append(tbl_grid)

    def _make_cell(runs: List, cell_w: int, *, italic: bool) -> Any:
        tc = OxmlElement("w:tc")

        tc_pr = OxmlElement("w:tcPr")
        tc_w_el = OxmlElement("w:tcW")
        tc_w_el.set(qn("w:w"), str(cell_w))
        tc_w_el.set(qn("w:type"), "dxa")
        tc_pr.append(tc_w_el)

        tc_borders = OxmlElement("w:tcBorders")
        for bname in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{bname}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tc_borders.append(b)
        tc_pr.append(tc_borders)

        va = OxmlElement("w:vAlign")
        va.set(qn("w:val"), "top")
        tc_pr.append(va)
        tc.append(tc_pr)

        p = OxmlElement("w:p")
        if ppr_elem is not None:
            new_ppr = copy.deepcopy(ppr_elem)
            for tag in ("w:ind", "w:spacing"):
                el = new_ppr.find(qn(tag))
                if el is not None:
                    new_ppr.remove(el)
            jc = new_ppr.find(qn("w:jc"))
            if jc is not None:
                jc.set(qn("w:val"), "left")
            p.append(new_ppr)

        for r_el in runs:
            if r_el is None:
                continue
            new_r = copy.deepcopy(r_el)
            if italic:
                rpr = new_r.find(qn("w:rPr"))
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    new_r.insert(0, rpr)
                if rpr.find(qn("w:i")) is None:
                    rpr.append(OxmlElement("w:i"))
                if rpr.find(qn("w:iCs")) is None:
                    rpr.append(OxmlElement("w:iCs"))
            for t_el in new_r.findall(qn("w:t")):
                t_el.set(qn("xml:space"), "preserve")
            p.append(new_r)

        tc.append(p)
        return tc

    tr = OxmlElement("w:tr")
    tr.append(_make_cell(src_runs, left_w, italic=False))
    tr.append(_make_cell(trans_runs, right_w, italic=italic_translation))
    tbl.append(tr)
    return tbl


def _convert_inline_bilingual_to_tables(
    doc: docx.Document,
    delimiter: str = "|",
    *,
    col_ratio: float = 0.5,
    italic_translation: bool = True,
) -> int:
    """Replace 'Src {delimiter} Trans' body paragraphs with borderless 2-column tables.

    DISABLED by default — the inline bilingual format (source | translation in one
    paragraph) is the intended layout.  Enable via PDF_DOCX_BILINGUAL_TABLE=1 only
    for special rendering needs.

    Returns count of converted paragraphs.
    """
    if not _env_bool("PDF_DOCX_BILINGUAL_TABLE", False):
        return 0

    content_w = _get_content_width_twips(doc)
    left_w  = max(1800, int(content_w * col_ratio))
    right_w = max(1800, content_w - left_w)

    body = doc.element.body
    candidates: List[Tuple[Any, Any, List, List]] = []

    for child in list(body):
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local != "p":
            continue
        result = _find_inline_bilingual_split(child, delimiter)
        if result is None:
            continue
        ppr_el, src_runs, trans_runs = result
        if not src_runs or not trans_runs:
            continue
        candidates.append((child, ppr_el, src_runs, trans_runs))

    converted = 0
    for p_elem, ppr_el, src_runs, trans_runs in candidates:
        try:
            tbl_elem = _build_bilingual_two_col_table(
                ppr_el, src_runs, trans_runs,
                left_w=left_w,
                right_w=right_w,
                italic_translation=italic_translation,
            )
            p_elem.addprevious(tbl_elem)
            parent = p_elem.getparent()
            if parent is not None:
                parent.remove(p_elem)
            converted += 1
        except Exception:
            continue

    return converted


def _is_pre_body_line(paragraph, src_i: int, abs_start: int) -> bool:
    """Title/author/affiliation block only — strictly before abstract start index."""
    if src_i >= abs_start:
        return False
    text = _paragraph_plain(paragraph)
    if not text:
        return True
    if _is_affiliation_line(paragraph):
        return True
    if _paragraph_is_centered(paragraph) and len(text) <= 160 and _paragraph_word_count(paragraph) <= 28:
        return True
    if len(text) <= 160 and _paragraph_word_count(paragraph) <= 28:
        low = text.lower()
        if any(k in low for k in ("@", "orcid", "author", "corresponding", "email")):
            return True
    return False


def _copy_layout_from_orig_to_translation(
    orig_para,
    trans_para,
    *,
    clear_first_line: bool = False,
    strip_alignment: bool = False,
) -> None:
    """Mirror orig paragraph layout on translation line; optionally drop first-line indent only."""
    orig_ppr = orig_para._element.find(qn("w:pPr"))
    if orig_ppr is None:
        return
    cloned = copy.deepcopy(orig_ppr)
    spacing = cloned.find(qn("w:spacing"))
    if spacing is not None:
        cloned.remove(spacing)
    if strip_alignment:
        jc = cloned.find(qn("w:jc"))
        if jc is not None:
            cloned.remove(jc)
    if clear_first_line:
        ind = cloned.find(qn("w:ind"))
        if ind is not None:
            ind.attrib.pop(qn("w:firstLine"), None)
    dst_ppr = trans_para._element.find(qn("w:pPr"))
    _copy_paragraph_properties(cloned, dst_ppr)


def _iter_newline_bilingual_pairs(src_doc: docx.Document, dst_doc: docx.Document):
    """Yield (src_index, src_para, orig_dst, trans_dst|None) for stacked bilingual docs.

    Improved robustness:
    - Bounded lookahead (≤10 dst paragraphs) to find the original source paragraph.
    - If lookahead is exceeded without a match the source paragraph is skipped but
      dst_pos is NOT advanced beyond the search window, preventing runaway drift.
    """
    _MAX_LOOKAHEAD = 10

    src_items = [
        (i, p)
        for i, p in enumerate(src_doc.paragraphs)
        if _paragraph_plain(p)
    ]
    dst_items = [
        (i, p)
        for i, p in enumerate(dst_doc.paragraphs)
        if _paragraph_plain(p)
    ]

    dst_pos = 0
    for src_pos, (src_i, src_p) in enumerate(src_items):
        src_text = _paragraph_plain(src_p)

        # Bounded forward scan for the source paragraph in the translated doc.
        matched_at: Optional[int] = None
        scan_limit = min(dst_pos + _MAX_LOOKAHEAD, len(dst_items))
        for scan in range(dst_pos, scan_limit):
            if _paragraph_texts_match(src_text, _paragraph_plain(dst_items[scan][1])):
                matched_at = scan
                break

        # Unbounded fallback only if the remaining dst window is small.
        if matched_at is None and scan_limit < len(dst_items):
            remaining = len(dst_items) - scan_limit
            if remaining <= _MAX_LOOKAHEAD:
                for scan in range(scan_limit, len(dst_items)):
                    if _paragraph_texts_match(src_text, _paragraph_plain(dst_items[scan][1])):
                        matched_at = scan
                        break

        if matched_at is None:
            continue

        dst_pos = matched_at
        _, orig_dst = dst_items[dst_pos]

        next_src_text = None
        if src_pos + 1 < len(src_items):
            next_src_text = _paragraph_plain(src_items[src_pos + 1][1])

        trans_dst = None
        if dst_pos + 1 < len(dst_items):
            _, candidate = dst_items[dst_pos + 1]
            cand_text = _paragraph_plain(candidate)
            if next_src_text and _paragraph_texts_match(cand_text, next_src_text):
                dst_pos += 1
                yield src_i, src_p, orig_dst, None
                continue
            if cand_text.strip() and not _paragraph_texts_match(cand_text, src_text):
                trans_dst = candidate
                dst_pos += 2
                yield src_i, src_p, orig_dst, trans_dst
                continue

        dst_pos += 1
        yield src_i, src_p, orig_dst, trans_dst


def _sync_bilingual_newline_layout(
    src_doc: docx.Document,
    dst_doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
) -> int:
    """Copy paragraph layout from pre-translation DOCX onto stacked pairs (same as inline).

    Newline is built via inline append + split; recovery mirrors inline:
    copy source pPr → downgrade false body center → _fix_spurious_center_alignment_in_doc.
    """
    _ = pdf_path  # kept for API compatibility
    synced = 0
    for _src_i, src_p, orig_dst, trans_dst in _iter_newline_bilingual_pairs(src_doc, dst_doc):
        if _is_in_table_cell(orig_dst):
            synced += 1
            continue
        _sync_paragraph_layout_like_inline(src_p, orig_dst, include_spacing=True)
        if trans_dst is not None:
            _sync_paragraph_layout_like_inline(src_p, trans_dst, include_spacing=False)
            _tighten_newline_pair_spacing(orig_dst, trans_dst)
            _sync_translation_run_font_from_orig(orig_dst, trans_dst)
        synced += 1
    return synced


def _normalize_newline_bilingual_layout_in_doc(
    dst_doc: docx.Document,
    src_doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
    bilingual_delimiter: str = "|",
) -> Dict[str, int]:
    """Newline layout pass — reuses inline recovery (copy source layout + shared center fix)."""
    stats: Dict[str, int] = {
        "alignment_normalized": 0,
        "indents_normalized": 0,
        "title_layout_preserved": 0,
        "center_alignment_fixed": 0,
        "paragraphs_synced": 0,
    }
    if not _env_bool("PDF_DOCX_NORMALIZE_ALIGN", True):
        return stats

    stats["paragraphs_synced"] = _sync_bilingual_newline_layout(
        src_doc, dst_doc, pdf_path=pdf_path
    )
    stats["alignment_normalized"] = int(stats["paragraphs_synced"])
    stats["indents_normalized"] = int(stats["paragraphs_synced"])
    stats["center_alignment_fixed"] = int(
        _fix_spurious_center_alignment_in_doc(
            dst_doc,
            src_doc,
            pdf_path=pdf_path,
            bilingual_delimiter=bilingual_delimiter,
        )
    )
    for _, _, orig_dst, trans_dst in _iter_newline_bilingual_pairs(src_doc, dst_doc):
        if trans_dst is None or _is_in_table_cell(orig_dst):
            continue
        _sync_translation_layout_from_orig(orig_dst, trans_dst)
    return stats


def recover_docx_layout(
    source_docx: str,
    translated_docx: str,
    *,
    pdf_path: Optional[str] = None,
    analysis: Optional[Dict[str, Any]] = None,
    bilingual_mode: Optional[str] = None,
    bilingual_delimiter: Optional[str] = None,
) -> Dict[str, int]:
    """Restore paragraph properties from pre-translation DOCX (safe, index-based)."""
    bi_mode = normalize_bilingual_mode(bilingual_mode)

    stats: Dict[str, int] = {
        "changed": 0,
        "paragraphs_synced": 0,
        "table_cells_synced": 0,
        "header_footer_synced": 0,
        "mismatched_paragraphs": 0,
    }

    if not os.path.isfile(source_docx) or not os.path.isfile(translated_docx):
        return stats
    if not _env_bool("PDF_DOCX_LAYOUT_SYNC", True):
        return stats

    src = docx.Document(source_docx)
    dst = docx.Document(translated_docx)
    use_regional = _should_apply_regional_layout(doc=dst, src_doc=src, analysis=analysis)

    if bi_mode == "newline":
        stats["mismatched_paragraphs"] = abs(len(src.paragraphs) - len(dst.paragraphs))
        post_stats = _normalize_newline_bilingual_layout_in_doc(
            dst,
            src,
            pdf_path=pdf_path,
            bilingual_delimiter=bilingual_delimiter or "|",
        )
        stats["paragraphs_synced"] = int(post_stats.get("paragraphs_synced", 0))
        stats["post_alignment_normalized"] = int(post_stats.get("alignment_normalized", 0))
        stats["post_indents_normalized"] = int(post_stats.get("indents_normalized", 0))
        stats["center_alignment_fixed"] = int(post_stats.get("center_alignment_fixed", 0))
        stats["table_cells_synced"] = _sync_table_layout_from_source(src, dst)
        stats["header_footer_synced"] = _sync_header_footer_layout(src, dst)
        table_stats = _recover_tables_and_images(dst)
        stats["table_rows_relaxed"] = int(table_stats.get("table_rows_relaxed", 0))
        stats["images_resized"] = int(table_stats.get("images_resized", 0))
        stats["inline_artifacts_stripped"] = _strip_inline_pdf_artifacts(dst)
        if any(int(stats.get(k, 0) or 0) for k in (
            "paragraphs_synced", "table_cells_synced", "header_footer_synced",
            "table_rows_relaxed", "images_resized", "inline_artifacts_stripped",
            "post_alignment_normalized", "post_indents_normalized",
            "center_alignment_fixed",
        )):
            stats["changed"] = 1
            dst.save(translated_docx)
        return stats

    pair_count = min(len(src.paragraphs), len(dst.paragraphs))
    stats["mismatched_paragraphs"] = abs(len(src.paragraphs) - len(dst.paragraphs))
    sync_fn = (
        _sync_paragraph_style_only
        if use_regional
        else _sync_paragraph_style_only
    )
    for i in range(pair_count):
        sync_fn(src.paragraphs[i], dst.paragraphs[i])
        if not use_regional:
            _sync_paragraph_layout_like_inline(
                src.paragraphs[i], dst.paragraphs[i], include_spacing=True
            )
        stats["paragraphs_synced"] += 1

    stats["table_cells_synced"] = _sync_table_layout_from_source(src, dst)
    stats["header_footer_synced"] = _sync_header_footer_layout(src, dst)

    table_stats = _recover_tables_and_images(dst)
    stats["table_rows_relaxed"] = int(table_stats.get("table_rows_relaxed", 0))
    stats["images_resized"] = int(table_stats.get("images_resized", 0))

    if bi_mode in ("inline", "newline", "none"):
        center_fixed = _fix_spurious_center_alignment_in_doc(
            dst,
            src,
            pdf_path=pdf_path,
            bilingual_delimiter=bilingual_delimiter or "|",
        )
        stats["center_alignment_fixed"] = int(center_fixed)

    if bi_mode == "none":
        post_stats = normalize_converted_docx_layout_in_doc(
            dst,
            pdf_path=pdf_path,
            src_doc=src,
            analysis=analysis,
        )
        stats["post_alignment_normalized"] = int(post_stats.get("alignment_normalized", 0))
        stats["post_indents_normalized"] = int(post_stats.get("indents_normalized", 0))

        if pdf_path and os.path.isfile(pdf_path) and _env_bool("PDF_DOCX_PDF_GEOMETRY_PROFILE", True):
            body_profile = _get_fallback_body_profile(dst)
            pdf_profiles = _derive_layout_profiles_from_pdf(pdf_path, dst)
            if pdf_profiles:
                body_profile = pdf_profiles[0]
            stats["pdf_indent_fixed"] = _fix_indents_from_pdf_geometry(
                dst,
                src,
                pdf_path,
                body_profile,
            )

    stats["inline_artifacts_stripped"] = _strip_inline_pdf_artifacts(dst)

    if (
        stats["paragraphs_synced"]
        or stats["table_cells_synced"]
        or stats["header_footer_synced"]
        or stats.get("table_rows_relaxed")
        or stats.get("images_resized")
        or stats.get("post_alignment_normalized")
        or stats.get("post_indents_normalized")
        or stats.get("pdf_indent_fixed")
        or stats.get("center_alignment_fixed")
        or stats.get("inline_artifacts_stripped")
    ):
        stats["changed"] = 1
        dst.save(translated_docx)

    return stats
