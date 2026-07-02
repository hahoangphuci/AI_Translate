"""DOCX-based PDF translation pipeline.

Steps:
1) PDF Analyzer (text/scan/table/columns/images)
2) Scan OCR (create searchable PDF)
3) PDF Cleaner (autorotate/enhance/deskew/layer sanitize)
4) PDF -> DOCX Converter
5) DOCX Translation (preserve paragraph/run/style/table/heading)
6) DOCX Layout Recovery
7) DOCX -> PDF
8) Quality Checker
"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from typing import Any, Callable, Dict, List, Optional

import docx

from ...docx_service import process_docx

try:
    from .layout_recovery import (
        normalize_bilingual_mode,
        normalize_converted_docx_layout_in_doc,
        recover_docx_layout,
        sanitize_converted_docx,
        _normalize_newline_bilingual_layout_in_doc,
    )
except ImportError:
    from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
        normalize_bilingual_mode,
        normalize_converted_docx_layout_in_doc,
        recover_docx_layout,
        sanitize_converted_docx,
        _normalize_newline_bilingual_layout_in_doc,
    )


def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in ("1", "true", "yes", "on")


def _env_int(name: str, default: int) -> int:
    raw = os.getenv(name)
    if raw is None:
        return int(default)
    try:
        return int(str(raw).strip())
    except Exception:
        return int(default)


def _emit(progress_callback: Optional[Callable[[int, str], None]], pct: int, msg: str) -> None:
    if progress_callback:
        progress_callback(int(pct), msg)


_UNICODE_MATH_RE = re.compile(
    r"[\u0391-\u03C9"       # Greek Α–ω
    r"\u2200-\u22FF"         # Mathematical Operators (∀∂∃∄∅∆∇∈∉∑∏∫√∞≤≥≠≈...)
    r"\u2190-\u21FF"         # Arrows ←↑→↓↔⇒⇐⇔...
    r"\u00B1\u00D7\u00F7\u00B7"  # ± × ÷ ·
    r"\u00B2\u00B3\u00B9"   # ² ³ ¹
    r"\u2070-\u2079"         # Superscript digits ⁰–⁹
    r"\u2080-\u2089"         # Subscript digits ₀–₉
    r"\u221A-\u221F"         # √ ∛ ∜ ∝ ∞ ∟ ∠ ∡ ∢
    r"\u2248\u2260\u2264\u2265"  # ≈ ≠ ≤ ≥
    r"]"
)


def _looks_like_formula(text: str) -> bool:
    s = (text or "").strip()
    if not s:
        return False
    # Unicode math symbols (Greek, operators, arrows, superscripts, etc.)
    unicode_math = _UNICODE_MATH_RE.findall(s)
    if unicode_math:
        latin_letters = re.findall(r"[A-Za-zÀ-ỹ]", s)
        if len(unicode_math) >= len(latin_letters) or len(unicode_math) >= 2:
            return True
    # ASCII math operators.
    math_syms = re.findall(r"[=+\-*/^_{}\\<>]", s)
    if len(math_syms) >= 2:
        letters = len(re.findall(r"[A-Za-zÀ-ỹ]", s))
        if letters <= 2:
            return True
        if len(math_syms) >= 4 and letters < len(math_syms):
            return True
    if re.search(r"\b\d+\s*[+\-*/]\s*\d+\b", s) and not re.search(r"[A-Za-zÀ-ỹ]", s):
        return True
    return False


def _looks_like_table_drawings(drawings: List[dict]) -> bool:
    if not drawings:
        return False
    h_lines = 0
    v_lines = 0
    for d in drawings:
        for item in d.get("items") or []:
            if not item:
                continue
            if item[0] != "l":
                continue
            coords = [p for p in item[1:] if isinstance(p, (list, tuple)) and len(p) >= 2]
            if len(coords) < 2:
                continue
            x0, y0 = coords[0][0], coords[0][1]
            x1, y1 = coords[1][0], coords[1][1]
            if abs(float(y1) - float(y0)) <= 1.2:
                h_lines += 1
            if abs(float(x1) - float(x0)) <= 1.2:
                v_lines += 1
            if h_lines >= 4 and v_lines >= 4:
                return True
    return False


def _page_has_multiple_columns(page: Any) -> bool:
    try:
        blocks = page.get_text("blocks") or []
    except Exception:
        return False

    try:
        page_width = float(page.rect.width)
    except Exception:
        page_width = 0.0
    if page_width <= 1.0:
        return False

    text_blocks: List[tuple] = []
    for block in blocks:
        if not isinstance(block, (list, tuple)) or len(block) < 5:
            continue
        try:
            x0, y0, x1, y1 = [float(v) for v in block[:4]]
            text = str(block[4] or "").strip()
        except Exception:
            continue
        if not text:
            continue
        w = x1 - x0
        h = y1 - y0
        if w <= 8 or h <= 4:
            continue
        # Ignore near full-width blocks to reduce false positives.
        if w >= page_width * 0.9:
            continue
        text_blocks.append((x0, y0, x1, y1))

    if len(text_blocks) < 4:
        return False

    left: List[tuple] = []
    right: List[tuple] = []
    left_center = page_width * 0.47
    right_center = page_width * 0.53

    for x0, y0, x1, y1 in text_blocks:
        center_x = (x0 + x1) / 2.0
        if center_x <= left_center:
            left.append((x0, y0, x1, y1))
        elif center_x >= right_center:
            right.append((x0, y0, x1, y1))

    if len(left) < 2 or len(right) < 2:
        return False

    overlaps = 0
    for l in left:
        for r in right:
            y_overlap = min(l[3], r[3]) - max(l[1], r[1])
            if y_overlap >= 8:
                overlaps += 1
                if overlaps >= 2:
                    return True
    return False


def _resolve_tesseract_cmd() -> Optional[str]:
    env_cmd = (os.getenv("TESSERACT_CMD") or "").strip().strip('"')
    candidates = []
    if env_cmd:
        candidates.append(env_cmd)

    from_path = shutil.which("tesseract")
    if from_path:
        candidates.append(from_path)

    if os.name == "nt":
        candidates.extend(
            [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                r"C:\Tesseract-OCR\tesseract.exe",
            ]
        )

    for candidate in candidates:
        if not candidate:
            continue
        p = str(candidate).strip().strip('"')
        if not p:
            continue
        if os.path.isabs(p) or p.lower().endswith(".exe"):
            if os.path.exists(p):
                return p
        else:
            resolved = shutil.which(p)
            if resolved:
                return resolved
    return None


def _render_pdf_page_to_image(page: Any, dpi: int) -> Any:
    from PIL import Image
    import fitz  # PyMuPDF

    zoom = float(max(72, dpi)) / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    return img


def _autorotate_image_with_tesseract(img: Any, pytesseract_mod: Any) -> tuple:
    try:
        from pytesseract import Output
    except Exception:
        return img, 0

    try:
        osd = pytesseract_mod.image_to_osd(img, output_type=Output.DICT)
        rotate = int(osd.get("rotate") or 0) % 360
    except Exception:
        rotate = 0

    if rotate in (90, 180, 270):
        try:
            return img.rotate(-rotate, expand=True, fillcolor=(255, 255, 255)), rotate
        except Exception:
            return img.rotate(-rotate, expand=True), rotate
    return img, 0


def _deskew_image(img: Any) -> tuple:
    try:
        import cv2
        import numpy as np
    except Exception:
        return img, 0.0

    arr = np.array(img.convert("RGB"))
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    _, thr = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    coords = np.column_stack(np.where(thr > 0))
    if coords is None or len(coords) < 250:
        return img, 0.0

    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    if abs(float(angle)) < 0.15:
        return img, 0.0

    h, w = arr.shape[:2]
    center = (w // 2, h // 2)
    mat = cv2.getRotationMatrix2D(center, float(angle), 1.0)
    rotated = cv2.warpAffine(
        arr,
        mat,
        (w, h),
        flags=cv2.INTER_CUBIC,
        borderMode=cv2.BORDER_CONSTANT,
        borderValue=(255, 255, 255),
    )

    from PIL import Image

    return Image.fromarray(rotated).convert("RGB"), float(angle)


def _enhance_image_for_ocr(img: Any) -> Any:
    try:
        import cv2
        import numpy as np
    except Exception:
        return img

    arr = np.array(img.convert("RGB"))
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    gray = cv2.fastNlMeansDenoising(gray, None, h=9, templateWindowSize=7, searchWindowSize=21)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)
    blur = cv2.GaussianBlur(gray, (0, 0), 1.0)
    sharpen = cv2.addWeighted(gray, 1.5, blur, -0.5, 0)
    bw = cv2.adaptiveThreshold(
        sharpen,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        12,
    )

    from PIL import Image

    return Image.fromarray(bw).convert("RGB")


def _raster_clean_pdf(
    input_pdf: str,
    output_pdf: str,
    progress_callback: Optional[Callable[[int, str], None]],
    *,
    start_pct: int,
    end_pct: int,
    stage_name: str,
    ocr_langs: Optional[str] = None,
    build_searchable: bool = False,
) -> Dict[str, Any]:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise RuntimeError("PyMuPDF is required for PDF cleanup") from exc

    dpi = max(120, _env_int("PDF_CLEAN_RENDER_DPI", 250))
    enable_autorotate = _env_bool("PDF_CLEAN_AUTOROTATE", True)
    enable_deskew = _env_bool("PDF_CLEAN_DESKEW", True)
    enable_enhance = _env_bool("PDF_CLEAN_ENHANCE", True)

    pytesseract_mod = None
    if enable_autorotate or build_searchable:
        try:
            import pytesseract  # type: ignore

            cmd = _resolve_tesseract_cmd()
            if cmd:
                pytesseract.pytesseract.tesseract_cmd = cmd
            _ = pytesseract.get_tesseract_version()
            pytesseract_mod = pytesseract
        except Exception:
            pytesseract_mod = None
            if build_searchable:
                raise RuntimeError(
                    "Tesseract OCR is required to create searchable PDF from scan pages"
                )

    src = fitz.open(input_pdf)
    out = fitz.open()
    rotated_pages = 0
    deskewed_pages = 0

    try:
        total = max(1, len(src))
        for idx, page in enumerate(src):
            img = _render_pdf_page_to_image(page, dpi)

            if enable_autorotate and pytesseract_mod is not None:
                img, rot = _autorotate_image_with_tesseract(img, pytesseract_mod)
                if rot:
                    rotated_pages += 1

            if enable_deskew:
                img, angle = _deskew_image(img)
                if abs(float(angle)) >= 0.15:
                    deskewed_pages += 1

            if enable_enhance:
                img = _enhance_image_for_ocr(img)

            if build_searchable:
                langs = (ocr_langs or os.getenv("OCR_LANGS_DEFAULT") or "eng").strip() or "eng"
                pdf_bytes = pytesseract_mod.image_to_pdf_or_hocr(img, extension="pdf", lang=langs)
                one = fitz.open("pdf", pdf_bytes)
                try:
                    out.insert_pdf(one)
                finally:
                    one.close()
            else:
                width_pt = float(img.width) * 72.0 / float(dpi)
                height_pt = float(img.height) * 72.0 / float(dpi)
                out_page = out.new_page(width=width_pt, height=height_pt)
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                out_page.insert_image(fitz.Rect(0, 0, width_pt, height_pt), stream=buf.getvalue())

            pct_span = max(1, end_pct - start_pct)
            pct = start_pct + int(((idx + 1) / total) * pct_span)
            _emit(progress_callback, pct, f"{stage_name}: page {idx + 1}/{total}")
    finally:
        src.close()

    out.save(output_pdf, deflate=True, garbage=2)
    out.close()
    return {
        "path": os.path.abspath(output_pdf),
        "rotated_pages": rotated_pages,
        "deskewed_pages": deskewed_pages,
    }


def _normalize_pdf_rotation(input_pdf: str, output_pdf: str) -> bool:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return False

    src = fitz.open(input_pdf)
    changed = False
    try:
        for page in src:
            try:
                rot = int(page.rotation or 0) % 360
            except Exception:
                rot = 0
            if rot != 0:
                try:
                    page.set_rotation(0)
                    changed = True
                except Exception:
                    pass

        if changed:
            src.save(output_pdf, deflate=True, garbage=2)
    finally:
        src.close()
    return changed


def analyze_pdf(pdf_path: str) -> Dict[str, Any]:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise RuntimeError("PyMuPDF is required for PDF analysis. Install: pip install PyMuPDF") from exc

    doc = fitz.open(pdf_path)
    has_text = False
    has_images = False
    has_tables = False
    has_multiple_columns = False
    multi_column_pages = 0
    has_formulas = False
    has_references = False
    has_abstract_heading = False
    has_keywords_heading = False
    has_intro_heading = False
    numbered_section_hits = 0
    text_chars = 0

    ref_re = re.compile(r"\b(references|reference|tai lieu tham khao|tài liệu tham khảo|doi|arxiv)\b", re.IGNORECASE)
    abstract_re = re.compile(
        r"\b(abstract|t[oóô]m\s*t[aắ]t|t[oóô]mt[aắ]t|abstrak|abstrakt|r[eé]sum[eé]|zusammenfassung)\b",
        re.IGNORECASE,
    )
    keywords_re = re.compile(
        r"\b(keywords?|t[uừ]\s*kh[oó]a|t[uừ]kh[oó]a|sleutelwoorde?|fjal[eë]t\s*ky[çc]e|schl[uü]sselw[oö]rter)\b",
        re.IGNORECASE,
    )
    intro_re = re.compile(
        r"\b(1\s+introduction|introduction|1\s+gi[oớ]i\s*th[iệ]u|gi[oớ]i\s*th[iệ]u|inleiding|einleitung|hyrje)\b",
        re.IGNORECASE,
    )
    numbered_section_re = re.compile(r"(?m)^\s*\d{1,2}(?:\.\d+)?\s+[A-ZÀ-Ỹ][^\n]{2,90}$")

    page_count = 0
    try:
        page_count = int(len(doc))
        for page in doc:
            try:
                text = page.get_text("text") or ""
            except Exception:
                text = ""

            if text.strip():
                has_text = True
                text_chars += len(text)
                if not has_references and ref_re.search(text):
                    has_references = True
                if not has_formulas and _looks_like_formula(text):
                    has_formulas = True
                if not has_abstract_heading and abstract_re.search(text):
                    has_abstract_heading = True
                if not has_keywords_heading and keywords_re.search(text):
                    has_keywords_heading = True
                if not has_intro_heading and intro_re.search(text):
                    has_intro_heading = True
                if numbered_section_hits < 12:
                    numbered_section_hits += len(numbered_section_re.findall(text))

            if not has_images:
                try:
                    has_images = bool(page.get_images(full=True))
                except Exception:
                    has_images = False

            if not has_tables:
                try:
                    tables = page.find_tables()
                    if tables and getattr(tables, "tables", None):
                        has_tables = True
                except Exception:
                    try:
                        if _looks_like_table_drawings(page.get_drawings() or []):
                            has_tables = True
                    except Exception:
                        pass

            if not has_multiple_columns:
                try:
                    if _page_has_multiple_columns(page):
                        has_multiple_columns = True
                        multi_column_pages += 1
                except Exception:
                    pass
            else:
                try:
                    if _page_has_multiple_columns(page):
                        multi_column_pages += 1
                except Exception:
                    pass
    finally:
        doc.close()

    is_scan = not has_text
    academic_score = 0
    if has_abstract_heading:
        academic_score += 2
    if has_keywords_heading:
        academic_score += 1
    if has_intro_heading:
        academic_score += 1
    if has_references:
        academic_score += 2
    if has_formulas:
        academic_score += 1
    if has_multiple_columns and multi_column_pages >= max(1, page_count // 3):
        academic_score += 1
    if numbered_section_hits >= 2:
        academic_score += 1

    is_academic_like = bool(
        has_text
        and (
            (has_abstract_heading and (has_intro_heading or numbered_section_hits >= 2))
            or (academic_score >= 4 and has_references)
            or academic_score >= 5
        )
    )
    return {
        "pages": page_count,
        "has_text": has_text,
        "has_images": has_images,
        "has_tables": has_tables,
        "has_multiple_columns": has_multiple_columns,
        "multi_column_pages": multi_column_pages,
        "has_formulas": has_formulas,
        "has_references": has_references,
        "has_abstract_heading": has_abstract_heading,
        "has_keywords_heading": has_keywords_heading,
        "has_intro_heading": has_intro_heading,
        "numbered_section_hits": numbered_section_hits,
        "academic_score": academic_score,
        "is_academic_like": is_academic_like,
        "text_chars": text_chars,
        "is_scan": is_scan,
    }


def _pdf2docx_convert_settings(analysis: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Tune pdf2docx parser settings to preserve layout/format as much as possible."""
    analysis = analysis or {}

    settings: Dict[str, Any] = {
        "delete_end_line_hyphen": _env_bool("PDF2DOCX_DELETE_HYPHEN", False),
        "parse_lattice_table": _env_bool("PDF2DOCX_PARSE_LATTICE_TABLE", True),
        "parse_stream_table": _env_bool("PDF2DOCX_PARSE_STREAM_TABLE", True),
        # extract_stream_table: pdf2docx default is False (stream tables parsed but NOT saved to docx).
        # We always enable it so detected stream tables are actually written as <w:tbl>.
        "extract_stream_table": _env_bool("PDF2DOCX_EXTRACT_STREAM_TABLE", True),
        # list_not_table: keep True (default) globally — prevents bullet/list structures from being
        # misclassified as table cells, which causes content to appear twice (inside table + as paragraphs).
        # Only disable when we have confirmed bordered tables in the PDF.
        "list_not_table": _env_bool("PDF2DOCX_LIST_NOT_TABLE", True),
        "ignore_page_error": True,
        "line_separate_threshold": _env_float("PDF2DOCX_LINE_SEPARATE", 5.0),
        "new_paragraph_free_space_ratio": _env_float("PDF2DOCX_NEW_PARA_RATIO", 0.85),
        "lines_center_aligned_threshold": _env_float("PDF2DOCX_CENTER_ALIGN_TOL", 2.0),
        "clip_image_res_ratio": _env_float("PDF2DOCX_CLIP_IMAGE_RES", 4.0),
        "float_image_ignorable_gap": _env_float("PDF2DOCX_FLOAT_IMG_GAP", 5.0),
        # Lower SVG merge gap so formula vector components are merged into one image block.
        "min_svg_gap_dx": _env_float("PDF2DOCX_MIN_SVG_GAP_DX", 8.0),
        "min_svg_gap_dy": _env_float("PDF2DOCX_MIN_SVG_GAP_DY", 2.0),
    }

    if analysis.get("has_tables"):
        # Confirmed bordered tables → disable list_not_table for more aggressive cell detection.
        settings["list_not_table"] = _env_bool("PDF2DOCX_LIST_NOT_TABLE", False)
    if analysis.get("has_formulas"):
        settings["min_svg_gap_dx"] = _env_float("PDF2DOCX_MIN_SVG_GAP_DX", 5.0)
        settings["min_svg_gap_dy"] = _env_float("PDF2DOCX_MIN_SVG_GAP_DY", 1.0)
    if analysis.get("has_multiple_columns"):
        settings["line_break_width_ratio"] = _env_float("PDF2DOCX_LINE_BREAK_WIDTH", 0.45)
        settings["line_separate_threshold"] = _env_float("PDF2DOCX_LINE_SEPARATE", 8.0)
    if analysis.get("is_scan") or analysis.get("ocr_searchable"):
        settings["ocr"] = 1 if analysis.get("is_scan") else 2
    if _env_bool("PDF2DOCX_MULTI_PROCESSING", False):
        settings["multi_processing"] = True
        settings["cpu_count"] = max(0, _env_int("PDF2DOCX_CPU_COUNT", 0))
    return settings


def _env_float(name: str, default: float) -> float:
    raw = os.getenv(name)
    if raw is None:
        return float(default)
    try:
        return float(str(raw).strip())
    except Exception:
        return float(default)


_MATH_FONT_RE = re.compile(
    r"cmmi|cmsy|cmex|cmr(?!oman)|cmbx|cmtt|cmit|"
    r"msbm|msam|euex|euf|eus|eur|esint|stmary|"
    r"stix|latinmodern.math|xits|asana.math|"
    r"fira.math|mathpazo|mathtime|lucidamath|"
    r"mtmi|mtex|wasy|rsfs|bbold|doublestroke",
    re.IGNORECASE,
)


def _inject_formula_clips(pdf_path: str, docx_path: str, *, clip_res: float = 3.0) -> int:
    """Detect lines in the PDF that use mathematical fonts and inject them as
    high-resolution PNG images into corresponding paragraphs in the DOCX.

    This handles the common case where pdf2docx extracts formula characters from
    math fonts (cmmi, cmsy, stix, …) as text with wrong Unicode mapping — the
    characters appear garbled.  We clip the original PDF region and replace those
    paragraphs with an image so the formula is visible and correct.

    Returns the number of formula images injected.
    """
    if not _env_bool("PDF_DOCX_INJECT_FORMULA_CLIPS", True):
        return 0

    try:
        import fitz  # PyMuPDF
        import docx as _docx_mod
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement
        from lxml import etree
        import copy
    except Exception:
        return 0

    if not os.path.isfile(pdf_path) or not os.path.isfile(docx_path):
        return 0

    # ── Step 1: collect math-font line bboxes from the PDF ──────────────────
    try:
        pdf_doc = fitz.open(pdf_path)
    except Exception:
        return 0

    # Maps plain-text key → (page, bbox) for each formula line
    formula_lines: List[Dict[str, Any]] = []  # {"text", "page", "bbox"}

    try:
        for page_idx in range(len(pdf_doc)):
            page = pdf_doc[page_idx]
            try:
                raw = page.get_text("rawdict", flags=0)
            except Exception:
                continue
            for block in raw.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    has_math_font = False
                    line_text = ""
                    for span in line.get("spans", []):
                        font_name = span.get("font", "")
                        if _MATH_FONT_RE.search(font_name):
                            has_math_font = True
                        line_text += span.get("text", "")
                    if has_math_font and line_text.strip():
                        # Expand bbox for padding
                        bbox = list(line["bbox"])
                        bbox[1] -= 4
                        bbox[3] += 4
                        formula_lines.append({
                            "text": line_text.strip(),
                            "page": page_idx,
                            "bbox": bbox,
                        })
    finally:
        pdf_doc.close()

    if not formula_lines:
        return 0

    # ── Step 2: open DOCX, match paragraphs by text similarity ──────────────
    try:
        doc = _docx_mod.Document(docx_path)
    except Exception:
        return 0

    # Build a simple list of (para, plain_text) for all body paragraphs
    all_paras = []
    for para in doc.paragraphs:
        try:
            txt = "".join(r.text or "" for r in para.runs).strip()
        except Exception:
            txt = ""
        all_paras.append((para, txt))

    def _text_sim(a: str, b: str) -> float:
        """Naive char-set overlap similarity."""
        if not a or not b:
            return 0.0
        sa, sb = set(a.lower()), set(b.lower())
        inter = len(sa & sb)
        return inter / max(len(sa), len(sb))

    # Track which DOCX paragraphs have already been replaced to avoid duplicates
    replaced: set = set()

    # ── Step 3: for each formula line, find best-matching DOCX paragraph ────
    injected = 0
    try:
        pdf_doc2 = fitz.open(pdf_path)
    except Exception:
        return 0

    try:
        for fl in formula_lines:
            fl_text = fl["text"]
            best_idx, best_sim = -1, 0.3  # min threshold
            for idx, (para, ptxt) in enumerate(all_paras):
                if idx in replaced:
                    continue
                sim = _text_sim(fl_text, ptxt)
                if sim > best_sim:
                    best_sim = sim
                    best_idx = idx

            if best_idx < 0:
                continue

            # Clip the formula region from the PDF page
            try:
                page = pdf_doc2[fl["page"]]
                clip = fitz.Rect(fl["bbox"])
                mat = fitz.Matrix(clip_res, clip_res)
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
                img_bytes = pix.tobytes("png")
            except Exception:
                continue

            # Insert image into matched DOCX paragraph
            para = all_paras[best_idx][0]
            try:
                # Clear paragraph content while preserving pPr
                p_elem = para._element
                ppr = p_elem.find(_qn("w:pPr"))
                # Remove all w:r children
                for child in list(p_elem):
                    if child.tag != _qn("w:pPr"):
                        p_elem.remove(child)
                # Add image run via python-docx inline image mechanism
                from io import BytesIO
                from docx.shared import Inches, Pt
                img_stream = BytesIO(img_bytes)
                # Estimate display width: clip width in points / 72 * clip_res → real pt → inches
                clip_w_pt = (fl["bbox"][2] - fl["bbox"][0])
                clip_h_pt = (fl["bbox"][3] - fl["bbox"][1])
                img_w_in = min(clip_w_pt / 72.0, 6.0)  # max 6 inches
                img_h_in = clip_h_pt / 72.0 * (img_w_in / (clip_w_pt / 72.0))
                new_run = para.add_run()
                new_run.add_picture(img_stream, width=int(img_w_in * 914400))  # EMU
                replaced.add(best_idx)
                injected += 1
            except Exception:
                continue
    finally:
        pdf_doc2.close()

    if injected:
        try:
            doc.save(docx_path)
        except Exception:
            pass

    return injected


def _load_pdf2docx_converter():
    """Import pdf2docx.Converter with deploy/python_packages path setup."""
    import sys

    api_base = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "..", ".."))
    if api_base not in sys.path:
        sys.path.insert(0, api_base)

    from deps_bootstrap import (
        check_pdf2docx_converter,
        ensure_packages_on_path,
    )

    ensure_packages_on_path()
    probe = check_pdf2docx_converter()
    if not probe["import_ok"]:
        raise RuntimeError(
            "PDF -> DOCX conversion unavailable. "
            + (probe.get("error") or "pdf2docx import failed")
            + (f"; PyMuPDF: {probe['fitz_error']}" if probe.get("fitz_error") else "")
            + ". Install: pip install pdf2docx PyMuPDF python-docx"
        )

    from pdf2docx import Converter

    return Converter


def _convert_pdf_to_docx(
    pdf_path: str,
    output_docx: str,
    engine: str,
    *,
    analysis: Optional[Dict[str, Any]] = None,
) -> None:
    engine = (engine or "pdf2docx").strip().lower()
    if engine in ("pdf2docx", "library", "lib"):
        Converter = _load_pdf2docx_converter()

        cv = Converter(pdf_path)
        try:
            cv.convert(output_docx, **_pdf2docx_convert_settings(analysis))
        finally:
            cv.close()
        return

    if engine in ("word", "microsoft", "msword"):
        raise RuntimeError("PDF -> DOCX via Microsoft Word is not implemented. Use PDF_DOCX_CONVERTER=pdf2docx")

    if engine in ("adobe", "adobe_api", "adobe_pdf_services"):
        raise RuntimeError("PDF -> DOCX via Adobe PDF Services is not implemented. Use PDF_DOCX_CONVERTER=pdf2docx")

    raise RuntimeError(f"Unknown PDF_DOCX_CONVERTER: {engine}")


def _extract_runs(paragraph) -> List[Dict[str, Any]]:
    runs = []
    for r in paragraph.runs:
        try:
            size = r.font.size.pt if r.font.size else None
        except Exception:
            size = None
        try:
            color = str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None
        except Exception:
            color = None
        runs.append(
            {
                "text": r.text or "",
                "bold": bool(r.bold) if r.bold is not None else None,
                "italic": bool(r.italic) if r.italic is not None else None,
                "underline": bool(r.underline) if r.underline is not None else None,
                "superscript": bool(getattr(getattr(r, "font", None), "superscript", False)),
                "subscript": bool(getattr(getattr(r, "font", None), "subscript", False)),
                "font": r.font.name,
                "size": size,
                "color": color,
            }
        )
    return runs


def extract_docx_structure(doc_obj: docx.Document) -> Dict[str, Any]:
    blocks: List[Dict[str, Any]] = []

    for p in doc_obj.paragraphs:
        blocks.append(
            {
                "type": "paragraph",
                "text": p.text or "",
                "style": getattr(getattr(p, "style", None), "name", None),
                "runs": _extract_runs(p),
            }
        )

    for t_idx, table in enumerate(doc_obj.tables):
        rows_out = []
        for r_idx, row in enumerate(table.rows):
            row_cells = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = "\n".join(p.text or "" for p in cell.paragraphs)
                row_cells.append(
                    {
                        "row": r_idx,
                        "col": c_idx,
                        "text": cell_text,
                    }
                )
            rows_out.append(row_cells)
        blocks.append({"type": "table", "index": t_idx, "rows": rows_out})

    try:
        for s_idx, section in enumerate(doc_obj.sections):
            for p in section.header.paragraphs:
                blocks.append(
                    {
                        "type": "header",
                        "section": s_idx,
                        "text": p.text or "",
                        "style": getattr(getattr(p, "style", None), "name", None),
                        "runs": _extract_runs(p),
                    }
                )
            for p in section.footer.paragraphs:
                blocks.append(
                    {
                        "type": "footer",
                        "section": s_idx,
                        "text": p.text or "",
                        "style": getattr(getattr(p, "style", None), "name", None),
                        "runs": _extract_runs(p),
                    }
                )
    except Exception:
        pass

    images = []
    try:
        related = getattr(getattr(doc_obj, "part", None), "related_parts", None)
        if isinstance(related, dict):
            for part in related.values():
                ct = str(getattr(part, "content_type", "") or "")
                if ct.startswith("image/"):
                    images.append(
                        {
                            "content_type": ct,
                            "partname": str(getattr(part, "partname", "") or ""),
                        }
                    )
    except Exception:
        pass

    return {"blocks": blocks, "images": images}


def generate_docx_ir(docx_path: str) -> Dict[str, Any]:
    doc_obj = docx.Document(docx_path)
    structure = extract_docx_structure(doc_obj)
    return {
        "meta": {
            "source": os.path.basename(docx_path),
            "paragraphs": len(doc_obj.paragraphs),
            "tables": len(doc_obj.tables),
        },
        **structure,
    }


def _docx_stats(docx_path: str) -> Dict[str, Any]:
    doc_obj = docx.Document(docx_path)
    para_texts = [p.text or "" for p in doc_obj.paragraphs]
    cell_texts = []
    cell_count = 0
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                cell_texts.append("\n".join(p.text or "" for p in cell.paragraphs))

    all_text = "\n".join(para_texts + cell_texts)
    text_chars = len(re.sub(r"\s+", "", all_text))

    ref_re = re.compile(r"\b(references|reference|tai lieu tham khao|tài liệu tham khảo|doi|arxiv)\b", re.IGNORECASE)
    has_references = bool(ref_re.search(all_text))

    return {
        "paragraphs": len(doc_obj.paragraphs),
        "table_cells": cell_count,
        "text_chars": text_chars,
        "has_references": has_references,
    }


def _find_soffice() -> str:
    env_path = (os.getenv("LIBREOFFICE_PATH") or "").strip()
    if env_path and os.path.isfile(env_path):
        return env_path
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    for candidate in (
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/usr/lib/libreoffice/program/soffice",
    ):
        if os.path.isfile(candidate):
            return candidate
    return ""


def _export_fallback_docx_enabled() -> bool:
    raw = (os.getenv("PDF_DOCX_EXPORT_FALLBACK_DOCX") or "").strip().lower()
    if raw in ("1", "true", "yes", "on"):
        return True
    if raw in ("0", "false", "no", "off"):
        return False
    return bool(os.getenv("WEBSITE_SITE_NAME") or os.getenv("WEBSITES_PORT"))


def _convert_docx_to_pdf(docx_path: str, output_dir: str, engine: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    engine = (engine or "auto").strip().lower()
    base = os.path.splitext(os.path.basename(docx_path))[0]
    output_pdf = os.path.join(output_dir, f"{base}.pdf")

    def _try_docx2pdf() -> tuple[bool, str]:
        try:
            from docx2pdf import convert as docx2pdf_convert
        except Exception as exc:
            return False, f"docx2pdf import error: {exc}"

        try:
            if os.path.exists(output_pdf):
                os.remove(output_pdf)
        except Exception:
            pass

        try:
            docx2pdf_convert(docx_path, output_pdf)
            if os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 0:
                return True, ""
        except Exception as exc:
            last_error = str(exc)
        else:
            last_error = "docx2pdf returned without output file"

        # Retry with short ASCII temp paths. This avoids occasional COM/Word failures
        # caused by long/unicode paths in project directories.
        tmp_dir = tempfile.mkdtemp(prefix="d2p_")
        try:
            tmp_docx = os.path.join(tmp_dir, "input.docx")
            tmp_pdf = os.path.join(tmp_dir, "output.pdf")
            shutil.copy2(docx_path, tmp_docx)
            docx2pdf_convert(tmp_docx, tmp_pdf)
            if os.path.exists(tmp_pdf) and os.path.getsize(tmp_pdf) > 0:
                shutil.copy2(tmp_pdf, output_pdf)
                return True, ""
        except Exception as exc:
            last_error = str(exc)
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

        return False, last_error

    def _try_libreoffice() -> tuple[bool, str]:
        soffice = _find_soffice()
        if not soffice:
            return False, "soffice not found in PATH and LIBREOFFICE_PATH is empty"
        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_path,
        ]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
        if result.returncode != 0:
            stderr = (result.stderr or b"").decode(errors="ignore").strip()
            return False, stderr or f"soffice exited with code {result.returncode}"
        if os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 0:
            return True, ""
        return False, "LibreOffice returned success but output PDF was not created"

    if engine == "auto":
        d2p_error = ""
        if os.name == "nt":
            ok, d2p_error = _try_docx2pdf()
            if ok:
                return output_pdf
        ok, lo_error = _try_libreoffice()
        if ok:
            return output_pdf
        raise RuntimeError(
            "DOCX -> PDF failed. "
            f"docx2pdf_error={d2p_error or 'n/a'}; "
            f"libreoffice_error={lo_error}. "
            "Install Microsoft Word or LibreOffice, or set PDF_DOCX_EXPORT_ENGINE accordingly."
        )

    if engine in ("docx2pdf", "word"):
        ok, d2p_error = _try_docx2pdf()
        if ok:
            return output_pdf
        # Be tolerant: if LibreOffice exists, use it as a fallback even when docx2pdf is requested.
        ok, lo_error = _try_libreoffice()
        if ok:
            return output_pdf
        raise RuntimeError(
            "DOCX -> PDF via docx2pdf failed. "
            f"docx2pdf_error={d2p_error}; libreoffice_error={lo_error}."
        )

    if engine in ("libreoffice", "soffice"):
        ok, lo_error = _try_libreoffice()
        if ok:
            return output_pdf
        raise RuntimeError(
            "DOCX -> PDF via LibreOffice failed. "
            f"Details: {lo_error}. Ensure soffice is in PATH or set LIBREOFFICE_PATH."
        )

    raise RuntimeError(f"Unknown PDF_DOCX_EXPORT_ENGINE: {engine}")


def _deliver_translated_docx_fallback(translated_docx: str, output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    base = os.path.splitext(os.path.basename(translated_docx))[0]
    dest = os.path.join(output_dir, f"{base}_translated.docx")
    shutil.copy2(translated_docx, dest)
    return dest


def quality_check(
    source_docx: str,
    translated_docx: str,
    output_pdf: str,
    analysis: Dict[str, Any],
) -> List[str]:
    warnings: List[str] = []
    try:
        src_stats = _docx_stats(source_docx)
        dst_stats = _docx_stats(translated_docx)
    except Exception:
        return ["quality_check_failed"]

    if dst_stats.get("text_chars", 0) <= 0:
        warnings.append("missing_text")

    if analysis.get("has_tables") and dst_stats.get("table_cells", 0) <= 0:
        warnings.append("tables_missing")

    if analysis.get("has_references") and not dst_stats.get("has_references"):
        warnings.append("references_missing")

    if not os.path.exists(output_pdf):
        warnings.append("output_pdf_missing")
    elif not str(output_pdf).lower().endswith(".pdf"):
        warnings.append("output_docx_fallback")

    return warnings


def run_pdf_docx_pipeline(
    service: Any,
    file_path: str,
    target_lang: str,
    progress_callback: Optional[Callable[[int, str], None]] = None,
    *,
    ocr_images: Optional[bool] = False,
    ocr_langs: Optional[str] = None,
    ocr_mode: Optional[str] = None,
    bilingual_mode: Optional[str] = None,
    bilingual_delimiter: Optional[str] = None,
) -> str:
    _emit(progress_callback, 2, "PDF Analyzer: checking text/scan/table/columns/images...")
    analysis = analyze_pdf(file_path)
    source_is_scan = bool(analysis.get("is_scan"))

    _emit(
        progress_callback,
        4,
        (
            "PDF Analyzer: "
            f"scan={source_is_scan}, "
            f"tables={bool(analysis.get('has_tables'))}, "
            f"columns={bool(analysis.get('has_multiple_columns'))}, "
            f"images={bool(analysis.get('has_images'))}"
        ),
    )

    if _env_bool("PDF_ANALYZE_DEBUG", False):
        dbg = os.path.join(service.download_folder, f"pdf_analyze_{uuid.uuid4().hex[:8]}.json")
        with open(dbg, "w", encoding="utf-8") as f:
            json.dump(analysis, f, ensure_ascii=False, indent=2)
        _emit(progress_callback, 4, f"PDF Analyzer: debug -> {os.path.basename(dbg)}")

    keep_intermediate = _env_bool("PDF_DOCX_KEEP_INTERMEDIATE", False)
    intermediate_paths: List[str] = []
    working_pdf = os.path.abspath(file_path)
    scan_ocr_applied = False

    if source_is_scan:
        _emit(progress_callback, 6, "Scan OCR: detected scan PDF, creating searchable PDF...")
        searchable_pdf = os.path.join(service.upload_folder, f"pdf_scan_ocr_{uuid.uuid4().hex[:8]}.pdf")
        try:
            stats = _raster_clean_pdf(
                working_pdf,
                searchable_pdf,
                progress_callback,
                start_pct=7,
                end_pct=24,
                stage_name="Scan OCR",
                ocr_langs=ocr_langs,
                build_searchable=True,
            )
            working_pdf = os.path.abspath(searchable_pdf)
            scan_ocr_applied = True
            analysis["ocr_searchable"] = True
            analysis["is_scan"] = False
            analysis["has_text"] = True
            intermediate_paths.append(working_pdf)
            _emit(
                progress_callback,
                24,
                (
                    "Scan OCR: searchable PDF ready "
                    f"(rotated={int(stats.get('rotated_pages', 0))}, "
                    f"deskewed={int(stats.get('deskewed_pages', 0))})"
                ),
            )
        except Exception as exc:
            msg = str(exc).strip().splitlines()[0][:200]
            _emit(progress_callback, 24, f"Scan OCR: skipped ({msg})")
            if _env_bool("PDF_SCAN_OCR_STRICT", False):
                raise RuntimeError(f"Scan OCR failed: {msg}") from exc

    _emit(progress_callback, 26, "PDF Cleaner: autorotate/enhance/deskew/layer sanitize...")
    if scan_ocr_applied:
        _emit(progress_callback, 30, "PDF Cleaner: preprocessing already applied during Scan OCR")
    else:
        cleaned_pdf = working_pdf
        try:
            if source_is_scan:
                cleaned_pdf = os.path.join(service.upload_folder, f"pdf_clean_{uuid.uuid4().hex[:8]}.pdf")
                _raster_clean_pdf(
                    working_pdf,
                    cleaned_pdf,
                    progress_callback,
                    start_pct=27,
                    end_pct=30,
                    stage_name="PDF Cleaner",
                    ocr_langs=ocr_langs,
                    build_searchable=False,
                )
            elif _env_bool("PDF_CLEAN_NORMALIZE_ROTATION", True):
                maybe_clean = os.path.join(service.upload_folder, f"pdf_rot_{uuid.uuid4().hex[:8]}.pdf")
                if _normalize_pdf_rotation(working_pdf, maybe_clean):
                    cleaned_pdf = maybe_clean
                    _emit(progress_callback, 30, "PDF Cleaner: page rotation normalized")
                else:
                    _emit(progress_callback, 30, "PDF Cleaner: no major cleanup needed")
            else:
                _emit(progress_callback, 30, "PDF Cleaner: skipped")
        except Exception as exc:
            msg = str(exc).strip().splitlines()[0][:200]
            _emit(progress_callback, 30, f"PDF Cleaner: fallback to original ({msg})")

        if os.path.abspath(cleaned_pdf) != os.path.abspath(working_pdf):
            working_pdf = os.path.abspath(cleaned_pdf)
            intermediate_paths.append(working_pdf)

    _emit(progress_callback, 32, "PDF -> DOCX: converting...")
    tmp_docx = os.path.join(
        service.upload_folder,
        f"pdf_{uuid.uuid4().hex[:8]}.docx",
    )
    _convert_pdf_to_docx(
        working_pdf,
        tmp_docx,
        os.getenv("PDF_DOCX_CONVERTER", "pdf2docx"),
        analysis=analysis,
    )

    noise_stats = sanitize_converted_docx(
        tmp_docx,
        pdf_path=working_pdf,
        analysis=analysis,
    )
    if any(int(v or 0) for v in noise_stats.values()):
        _emit(
            progress_callback,
            34,
            (
                "PDF -> DOCX: cleaned "
                f"noise={int(noise_stats.get('noise_removed', 0))}, "
                f"merge={int(noise_stats.get('fragments_merged', 0))}, "
                f"align={int(noise_stats.get('alignment_normalized', 0))}, "
                f"indent={int(noise_stats.get('indents_normalized', 0))}"
            ),
        )

    # Determine bilingual mode early so we can skip formula clipping for bilingual modes.
    # (Formula clipper replaces paragraphs with images before translation; that breaks bilingual
    # text insertion which appends translated text to the same paragraph elements.)
    _bi_mode_early = normalize_bilingual_mode(bilingual_mode)

    # Inject formula clips: detect math-font lines in the PDF and replace the
    # corresponding garbled-text paragraphs in the DOCX with clipped PNG images.
    # Only run for normal (non-bilingual) mode — bilingual modes need the original paragraph
    # structure intact so translated text can be appended correctly.
    if _bi_mode_early == "none" and (analysis.get("has_formulas") or analysis.get("is_academic_like")):
        try:
            n_clips = _inject_formula_clips(working_pdf, tmp_docx, clip_res=3.0)
            if n_clips:
                _emit(progress_callback, 35, f"Formula Clipper: injected {n_clips} formula image(s)")
        except Exception as _fclip_exc:
            _emit(progress_callback, 35, f"Formula Clipper: skipped ({str(_fclip_exc)[:80]})")

    _emit(progress_callback, 36, "DOCX Structure Extractor: reading...")
    docx_ir = generate_docx_ir(tmp_docx)

    if _env_bool("DOCX_IR_DEBUG", False) or _env_bool("PDF_DOCX_IR_DEBUG", False):
        dbg = os.path.join(service.download_folder, f"docx_ir_{uuid.uuid4().hex[:8]}.json")
        with open(dbg, "w", encoding="utf-8") as f:
            json.dump(docx_ir, f, ensure_ascii=False, indent=2)
        _emit(progress_callback, 38, f"DOCX IR: debug -> {os.path.basename(dbg)}")

    _emit(progress_callback, 40, "DOCX Translation: preserving paragraph/run/style/table/heading...")
    _emit(progress_callback, 41, "Translation Guard: skipping URL/DOI/reference/formula/code")

    def _docx_progress(pct: int, msg: str) -> None:
        mapped = 42 + int((max(0, min(100, pct)) / 100.0) * 40)
        _emit(progress_callback, mapped, msg)

    use_ocr_images = bool(ocr_images) or (source_is_scan and not scan_ocr_applied)
    requested_ocr_mode = (str(ocr_mode).strip().lower() if ocr_mode else None)
    if requested_ocr_mode not in (None, "image", "text", "both", "auto"):
        requested_ocr_mode = None
    use_ocr_mode = (requested_ocr_mode or ("auto" if source_is_scan else "image")) if use_ocr_images else None

    translated_docx = process_docx(
        service,
        tmp_docx,
        target_lang,
        progress_callback=_docx_progress,
        ocr_images=use_ocr_images,
        ocr_langs=ocr_langs,
        ocr_mode=use_ocr_mode,
        bilingual_mode=bilingual_mode,
        bilingual_delimiter=bilingual_delimiter,
        from_pdf=True,
    )

    if not str(translated_docx).lower().endswith(".docx"):
        raise RuntimeError("DOCX rebuild failed; cannot export to PDF. Check DOCX output in downloads.")

    bi_mode = normalize_bilingual_mode(bilingual_mode)
    if bi_mode != "none":
        _emit(
            progress_callback,
            83,
            (
                "Bilingual mode: "
                + ("inline (adjacent)" if bi_mode == "inline" else "newline (stacked)")
            ),
        )

    _emit(
        progress_callback,
        84,
        "DOCX Layout Recovery: copy paragraph layout from source DOCX...",
    )
    recovery_stats = recover_docx_layout(
        tmp_docx,
        translated_docx,
        pdf_path=working_pdf,
        analysis=analysis,
        bilingual_mode=bi_mode,
        bilingual_delimiter=bilingual_delimiter,
    )
    if recovery_stats.get("changed"):
        _bi_tables = int(recovery_stats.get("bilingual_tables_converted", 0))
        _bi_table_msg = f", bilingual_tables={_bi_tables}" if _bi_tables else ""
        _emit(
            progress_callback,
            88,
            (
                "DOCX Layout Recovery: "
                f"paras={int(recovery_stats.get('paragraphs_synced', 0))}, "
                f"mismatch={int(recovery_stats.get('mismatched_paragraphs', 0))}, "
                f"table_cells={int(recovery_stats.get('table_cells_synced', 0))}, "
                f"images={int(recovery_stats.get('images_resized', 0))}"
                + _bi_table_msg
            ),
        )
    else:
        _emit(progress_callback, 88, "DOCX Layout Recovery: no structural fixes needed")

    if bi_mode == "none":
        _emit(progress_callback, 89, "DOCX Layout: final regional profile pass before PDF export...")
        try:
            import docx as _docx_mod

            final_doc = _docx_mod.Document(translated_docx)
            src_doc = _docx_mod.Document(tmp_docx)
            final_stats = normalize_converted_docx_layout_in_doc(
                final_doc,
                pdf_path=working_pdf,
                src_doc=src_doc,
                analysis=analysis,
            )
            if any(int(v or 0) for v in final_stats.values()):
                final_doc.save(translated_docx)
        except Exception as exc:
            msg = str(exc).strip().splitlines()[0][:160]
            _emit(progress_callback, 89, f"DOCX Layout final pass skipped ({msg})")
    elif bi_mode == "newline":
        _emit(progress_callback, 89, "Bilingual newline: inline-style layout sync before PDF export...")
        try:
            final_doc = docx.Document(translated_docx)
            src_doc = docx.Document(tmp_docx)
            final_stats = _normalize_newline_bilingual_layout_in_doc(
                final_doc,
                src_doc,
                pdf_path=working_pdf,
                bilingual_delimiter=bilingual_delimiter or "|",
            )
            if any(int(v or 0) for v in final_stats.values()):
                final_doc.save(translated_docx)
                _emit(
                    progress_callback,
                    89,
                    (
                        "Bilingual newline layout (inline path): "
                        f"synced={int(final_stats.get('paragraphs_synced', 0))}, "
                        f"center_fixed={int(final_stats.get('center_alignment_fixed', 0))}"
                    ),
                )
        except Exception as exc:
            msg = str(exc).strip().splitlines()[0][:160]
            _emit(progress_callback, 89, f"Bilingual newline final pass skipped ({msg})")
    else:
        _emit(progress_callback, 89, "Bilingual inline: skip layout re-normalize (preserve bilingual text)")

    try:
        from .layout_recovery import sanitize_document_text_fonts

        final_doc = docx.Document(translated_docx)
        font_fixed = sanitize_document_text_fonts(final_doc)
        if font_fixed:
            final_doc.save(translated_docx)
            _emit(progress_callback, 89, f"DOCX font sanitize: fixed {font_fixed} run(s) before PDF export")
    except Exception as exc:
        msg = str(exc).strip().splitlines()[0][:160]
        _emit(progress_callback, 89, f"DOCX font sanitize skipped ({msg})")

    _emit(progress_callback, 90, "DOCX -> PDF: exporting...")
    output_path = ""
    try:
        output_path = _convert_docx_to_pdf(
            translated_docx,
            service.download_folder,
            os.getenv("PDF_DOCX_EXPORT_ENGINE", "auto"),
        )
    except RuntimeError as exc:
        if not _export_fallback_docx_enabled():
            raise
        output_path = _deliver_translated_docx_fallback(translated_docx, service.download_folder)
        _emit(
            progress_callback,
            93,
            f"PDF export unavailable ({exc}); delivered translated DOCX instead",
        )

    _emit(progress_callback, 95, "Quality Checker: validating output...")
    warnings = quality_check(tmp_docx, translated_docx, output_path, analysis)
    if warnings:
        _emit(progress_callback, 98, f"Quality Checker warnings: {', '.join(warnings)}")

    _emit(progress_callback, 100, "PDF translation completed")

    if not keep_intermediate:
        for path in [tmp_docx, translated_docx, *intermediate_paths]:
            try:
                if os.path.abspath(path) != os.path.abspath(output_path):
                    os.remove(path)
            except Exception:
                pass

    return os.path.abspath(output_path)
