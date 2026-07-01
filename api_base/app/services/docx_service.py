"""DOCX translation service — extracted from FileService._process_docx.

This module contains the full DOCX translation pipeline including:
- Bilingual modes (inline, newline)
- OCR image handling (overlay, text insertion, auto mode)
- Format-group translation preserving per-run styling
- Table, header/footer, TOC processing
- Post-processing fixes (hyperlinks, leaders, phrase fixes)
"""

import os
import re
import unicodedata
import uuid
import io
import zipfile
import docx
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, RGBColor, Pt
from .file_service import ProviderRateLimitError


# Matches Vietnamese-specific diacritic letters (both lower- and upper-case).
# NOTE: do NOT compile this with re.IGNORECASE — case-folding the à-ỹ range under
# IGNORECASE makes it match plain ASCII letters (e.g. "Heading", "Format"), which
# previously caused bilingual translations to be detected as "still untranslated"
# and appended a second time.
_VN_DIACRITIC_RE = re.compile(r"[à-ỹÀ-ỸđĐ]")


def _has_vietnamese_diacritics(text) -> bool:
    """Return True if the text contains any Vietnamese-specific diacritic letter."""
    return bool(_VN_DIACRITIC_RE.search(text or ""))


def _word_count(text) -> int:
    return len(re.findall(r"[\w\u00C0-\u1EF9]+", text or "", flags=re.UNICODE))


def _is_fragmented_prose_segments(segments) -> bool:
    """True when many tiny runs/segments should be translated as one prose block."""
    parts = [(s or "").strip() for s in (segments or []) if (s or "").strip()]
    if len(parts) < 2:
        return False
    total_words = sum(_word_count(p) for p in parts)
    if total_words < 3:
        return False
    short_parts = sum(1 for p in parts if len(p) <= 14 and _word_count(p) <= 2)
    if short_parts >= max(2, len(parts) * 2 // 3):
        return True
    avg_len = sum(len(p) for p in parts) / len(parts)
    return avg_len <= 10 and total_words >= 4


def _run_has_drawing_xml(run) -> bool:
    """Return True if the run element contains any drawing, image, VML, or OLE child.

    Setting run.text on such a run calls clear_content() internally, which would
    destroy embedded drawings/formulas.  Always check this before mutating run.text.
    """
    try:
        return bool(
            run._element.xpath(
                './/*[local-name()="drawing" or local-name()="pict"'
                ' or local-name()="object" or local-name()="OLEObject"]'
            )
        )
    except Exception:
        return False


def _para_has_protected_xml(paragraph) -> bool:
    """Return True if the paragraph contains drawings, OMML math, or OLE objects.

    Used as a fast early-exit guard before any run-level text mutation on the paragraph.
    Mirrors layout_recovery._para_has_protected_content for use inside docx_service.
    """
    try:
        return bool(
            paragraph._element.xpath(
                './/*[local-name()="drawing" or local-name()="pict"'
                ' or local-name()="object" or local-name()="OLEObject"'
                ' or local-name()="oMath" or local-name()="oMathPara"]'
            )
        )
    except Exception:
        return False


def _docx_run_boundary_needs_space(prev: str, next_t: str) -> bool:
    """Heuristic: pdf2docx often emits word-sized <w:r> without spaces between them."""
    ps = (prev or "").rstrip()
    ns = (next_t or "").lstrip()
    if not ps or not ns:
        return False
    if ps.endswith((" ", "\n", "\t")) or ns.startswith((" ", "\n", "\t")):
        return False
    pc = ps[-1]
    nc = ns[0]
    if ps.endswith(("-", "'", "\u2019")):
        return False
    # Comma / clause punctuation → next token (2025,bao | giá.Kết split across runs)
    if pc in ",.;:!?":
        if nc.isalnum():
            if pc in ",." and len(ps) >= 2 and ps[-2].isdigit() and nc.isdigit():
                return False
            return True
        return False
    if pc in ")]}" and nc.isalnum():
        return True
    if pc.isdigit() and nc.isdigit():
        return False
    if pc.isalnum() and nc.isalnum():
        return True
    return False


def _intrarun_insert_missing_spaces(text: str) -> str:
    """Inside a single run: add space after punctuation if pdf2docx glued (e.g. 'giá.Kết', '2025,bao')."""
    if not text:
        return text
    out = []
    n = len(text)
    for i, c in enumerate(text):
        out.append(c)
        if i + 1 >= n:
            continue
        nxt = text[i + 1]
        if nxt in " \n\t\r":
            continue
        prev_ch = text[i - 1] if i > 0 else ""
        if c in ",;:":
            if c == "," and prev_ch.isdigit() and nxt.isdigit():
                continue
            if nxt.isalnum():
                out.append(" ")
        elif c == ".":
            if prev_ch.isdigit() and nxt.isdigit():
                continue
            # Add space after "." for any letter (not just uppercase), consistent with
            # inter-run logic in _docx_run_boundary_needs_space.
            # This fixes "word.lowercase" glue artifacts from pdf2docx.
            if nxt.isalpha():
                out.append(" ")
        elif c in "!?":
            if nxt.isalnum():
                out.append(" ")
    return "".join(out)


def apply_docx_paragraph_spacing(paragraph) -> None:
    """Mutate runs only (preserve bold per run); fixes pdf2docx glue + punctuation.

    Skips any run that contains a drawing, image, or VML element to avoid calling
    run.text (which internally calls clear_content() and would destroy the drawing).
    """
    try:
        runs = list(paragraph.runs)
    except Exception:
        return
    if not runs:
        return
    # Fast-exit: skip the whole paragraph if it has any protected content.
    if _para_has_protected_xml(paragraph):
        return
    for i in range(len(runs) - 1):
        try:
            if _run_has_drawing_xml(runs[i]) or _run_has_drawing_xml(runs[i + 1]):
                continue
            a = runs[i].text or ""
            b = runs[i + 1].text or ""
            if _docx_run_boundary_needs_space(a, b):
                runs[i].text = a + " "
        except Exception:
            continue
    for r in runs:
        try:
            if _run_has_drawing_xml(r):
                continue
            t = r.text or ""
            nt = _intrarun_insert_missing_spaces(t)
            if nt != t:
                r.text = nt
        except Exception:
            continue


def merged_paragraph_plain(paragraph) -> str:
    """Plain text as seen after spacing heuristics (for translation + cache keys)."""
    try:
        texts = [r.text or "" for r in paragraph.runs]
    except Exception:
        return ""
    return _intrarun_insert_missing_spaces(join_docx_run_texts(texts))


def join_docx_run_texts(run_texts):
    """Join DOCX run strings; infer spaces between runs when pdf2docx glued words."""
    seq = [t or "" for t in (run_texts or [])]
    if not seq:
        return ""
    parts = [seq[0]]
    for t in seq[1:]:
        prev = parts[-1]
        if _docx_run_boundary_needs_space(prev, t):
            parts.append(" ")
        parts.append(t)
    return "".join(parts)


_RUN_BOUNDARY_MARKER_RE = re.compile(r"\uE000(\d+)\uE001")
_RUN_BOUNDARY_MARKER_SPLIT_RE = re.compile(r"(\uE000\d+\uE001)")


def _run_boundary_marker(index: int) -> str:
    return f"\uE000{index}\uE001"


def _run_format_signature(run, *, ignore_color: bool = False) -> bytes:
    """Stable key for run formatting: XML rPr plus resolved font properties."""
    from docx.oxml.ns import qn as _qn

    try:
        from lxml import etree

        chunks = []
        rpr = run._element.find(_qn("w:rPr"))
        if rpr is not None:
            if ignore_color:
                rpr_copy = etree.fromstring(etree.tostring(rpr))
                for tag in ("w:color", "w:highlight", "w:shd"):
                    for el in rpr_copy.findall(_qn(tag)):
                        rpr_copy.remove(el)
                chunks.append(etree.tostring(rpr_copy))
            else:
                chunks.append(etree.tostring(rpr))

        def _tri(name, val):
            if val is True:
                chunks.append(f"{name}=1")
            elif val is False:
                chunks.append(f"{name}=0")

        _tri("bold", run.bold)
        _tri("italic", run.italic)
        _tri("underline", run.underline)
        font = getattr(run, "font", None)
        if font is not None:
            for name, attr in (
                ("strike", "strike"),
                ("dblStrike", "double_strike"),
                ("shadow", "shadow"),
                ("outline", "outline"),
                ("emboss", "emboss"),
                ("imprint", "imprint"),
                ("caps", "all_caps"),
                ("smallCaps", "small_caps"),
                ("sup", "superscript"),
                ("sub", "subscript"),
            ):
                try:
                    _tri(name, getattr(font, attr, None))
                except Exception:
                    pass
            if not ignore_color:
                try:
                    rgb = font.color.rgb if font.color else None
                    if rgb is not None:
                        chunks.append(f"rgb={rgb}")
                except Exception:
                    pass
                try:
                    tc = font.color.theme_color if font.color else None
                    if tc is not None:
                        chunks.append(f"theme={tc}")
                except Exception:
                    pass
                try:
                    if font.highlight_color:
                        chunks.append(f"hl={font.highlight_color}")
                except Exception:
                    pass
            try:
                if font.name:
                    chunks.append(f"font={font.name}")
            except Exception:
                pass
            try:
                if font.size:
                    chunks.append(f"sz={font.size.pt}")
            except Exception:
                pass

        encoded = []
        for chunk in chunks:
            if isinstance(chunk, bytes):
                encoded.append(chunk.decode("utf-8", errors="ignore"))
            else:
                encoded.append(str(chunk))
        return "|".join(encoded).encode("utf-8")
    except Exception:
        return b""


def _run_format_signature_semantic(run) -> bytes:
    """Ignore color-only differences when deciding translation grouping."""
    return _run_format_signature(run, ignore_color=True)


def copy_run_rpr_preserve_all(src_run, dst_run) -> None:
    """Deep-copy w:rPr from src to dst without stripping decorative properties."""
    try:
        from docx.oxml.ns import qn as _qn
        import copy as _copy

        src_rpr = src_run._element.find(_qn("w:rPr"))
        if src_rpr is None:
            return
        new_rpr = _copy.deepcopy(src_rpr)
        old_rpr = dst_run._element.find(_qn("w:rPr"))
        if old_rpr is not None:
            dst_run._element.remove(old_rpr)
        dst_run._element.insert(0, new_rpr)
    except Exception:
        pass


def distribute_text_by_source_weights(translated_text, source_chunks):
    chunks = list(source_chunks or [])
    n = len(chunks)
    txt = translated_text or ""
    if n <= 0:
        return []
    if n == 1:
        return [txt]
    if not txt:
        return [""] * n

    weights = [max(1, len(c or "")) for c in chunks]
    total_weight = max(1, sum(weights))
    txt_len = len(txt)
    boundaries = []
    acc = 0
    for i in range(1, n):
        acc += weights[i - 1]
        target = int(round((acc / total_weight) * txt_len))
        target = max(1, min(txt_len - 1, target))

        left = max(1, target - 24)
        right = min(txt_len - 1, target + 24)
        cut = target

        found = None
        p = target
        while p >= left:
            if txt[p - 1].isspace():
                found = p
                break
            p -= 1
        if found is None:
            p = target + 1
            while p <= right:
                if txt[p - 1].isspace():
                    found = p
                    break
                p += 1
        if found is not None:
            cut = found

        if boundaries and cut <= boundaries[-1]:
            cut = min(txt_len - 1, boundaries[-1] + 1)
        boundaries.append(cut)

    out = []
    prev = 0
    for b in boundaries:
        out.append(txt[prev:b])
        prev = b
    out.append(txt[prev:])

    if len(out) < n:
        out.extend([""] * (n - len(out)))
    elif len(out) > n:
        out = out[: n - 1] + ["".join(out[n - 1 :])]

    # Avoid orphan 1-letter chunks mid-word when source segments were whole words.
    for i in range(len(out)):
        chunk = out[i] or ""
        src = chunks[i] if i < len(chunks) else ""
        stripped = chunk.strip()
        if len(stripped) != 1 or not stripped.isalpha():
            continue
        if len((src or "").strip()) <= 2:
            continue
        if i + 1 < len(out):
            out[i + 1] = chunk + (out[i + 1] or "")
            out[i] = ""
        elif i > 0:
            out[i - 1] = (out[i - 1] or "") + chunk
            out[i] = ""

    return out


def build_run_marker_translation_text(segments):
    parts = []
    for i, seg in enumerate(segments or []):
        parts.append(_run_boundary_marker(i))
        parts.append(seg or "")
    return "".join(parts)


def strip_run_boundary_markers(text: str) -> str:
    """Remove run-boundary placeholder tokens leaked by the translator."""
    if not text:
        return text
    out = text or ""
    # Corrupted wrappers: keep inner text, drop marker envelope.
    out = re.sub(r"\\uE000([^\\]+)\\uE001", r"\1", out)
    out = re.sub(r"\uE000([^\uE001]+)\uE001", r"\1", out)
    out = _RUN_BOUNDARY_MARKER_SPLIT_RE.sub("", out)
    out = re.sub(r"\\uE000\d+\\uE001", "", out)
    return out


def parse_run_marker_translation(translated: str, segment_count: int):
    text = translated or ""
    if segment_count <= 0:
        return []
    if segment_count == 1:
        m = _RUN_BOUNDARY_MARKER_RE.search(text)
        if m and m.start() == 0:
            return [text[m.end() :]]
        return [text]

    mapped = {}
    pieces = _RUN_BOUNDARY_MARKER_SPLIT_RE.split(text)
    i = 0
    while i < len(pieces):
        part = pieces[i]
        if not part:
            i += 1
            continue
        mm = _RUN_BOUNDARY_MARKER_RE.fullmatch(part)
        if mm:
            idx = int(mm.group(1))
            seg_text = pieces[i + 1] if i + 1 < len(pieces) else ""
            mapped[idx] = seg_text
            i += 2
        else:
            i += 1

    if len(mapped) >= segment_count and all(j in mapped for j in range(segment_count)):
        return [mapped.get(j, "") for j in range(segment_count)]
    return None


def paragraph_has_distinct_run_formats(runs, original_texts=None) -> bool:
    signatures = set()
    for i, run in enumerate(runs or []):
        if _run_has_drawing_xml(run):
            continue
        if original_texts is not None:
            txt = original_texts[i] if i < len(original_texts) else (run.text or "")
        else:
            txt = run.text or ""
        if not (txt or "").strip():
            continue
        signatures.add(_run_format_signature_semantic(run))
        if len(signatures) > 1:
            return True
    return False


def write_translated_to_run_indices(runs, indices, original_texts, new_text, *, ensure_font_compat_fn=None):
    if not indices:
        return

    safe_indices = [i for i in indices if not _run_has_drawing_xml(runs[i])]
    if not safe_indices:
        return

    candidate_indices = [i for i in safe_indices if (original_texts[i] or "").strip()]
    if not candidate_indices:
        candidate_indices = list(safe_indices)

    if len(candidate_indices) == 1:
        idx = candidate_indices[0]
        text_value = new_text or ""
        orig = original_texts[idx] if idx < len(original_texts) else ""
        if orig and orig != orig.strip():
            lead = orig[: len(orig) - len(orig.lstrip())]
            tail = orig[len(orig.rstrip()) :]
            core = text_value.strip() if (text_value or "").strip() else text_value
            text_value = f"{lead}{core}{tail}"
        if ensure_font_compat_fn:
            ensure_font_compat_fn(runs[idx], text_value)
        runs[idx].text = text_value
        for i in indices:
            if i != idx:
                runs[i].text = ""
        return

    chunks = distribute_text_by_source_weights(
        new_text or "",
        [original_texts[i] for i in candidate_indices],
    )
    for idx, chunk in zip(candidate_indices, chunks):
        if ensure_font_compat_fn:
            ensure_font_compat_fn(runs[idx], chunk)
        runs[idx].text = chunk
    for i in indices:
        if i not in candidate_indices and not _run_has_drawing_xml(runs[i]):
            runs[i].text = ""


def append_formatted_translation_runs(paragraph, translated_text, template_runs, *, ensure_font_compat_fn=None):
    """Append translated text as multiple runs mirroring source formatting segments."""
    templates = [
        r for r in (template_runs or [])
        if (r.text or "").strip() and not _run_has_drawing_xml(r)
    ]
    text_value = translated_text or ""
    if not text_value:
        return False
    if not templates:
        paragraph.add_run(text_value)
        return True

    if len(templates) == 1:
        rr = paragraph.add_run(text_value)
        copy_run_rpr_preserve_all(templates[0], rr)
        if ensure_font_compat_fn:
            ensure_font_compat_fn(rr, text_value)
        return True

    source_segments = [r.text or "" for r in templates]
    chunks = distribute_text_by_source_weights(text_value, source_segments)
    for tmpl, chunk in zip(templates, chunks):
        rr = paragraph.add_run(chunk)
        copy_run_rpr_preserve_all(tmpl, rr)
        if ensure_font_compat_fn:
            ensure_font_compat_fn(rr, chunk)
        try:
            from docx.oxml.ns import qn as _qn
            for t_el in rr._element.findall(_qn("w:t")):
                t_el.set(_qn("xml:space"), "preserve")
        except Exception:
            pass
    return True


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_VML_NS = "urn:schemas-microsoft-com:vml"
_DRAWINGML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_EMOJI_ONLY_RUN_RE = re.compile(
    r"^[\U0001F300-\U0001FAFF\U00002600-\U000027BF"
    r"\U0001F1E6-\U0001F1FF\s\u200d\ufe0f]+$",
    flags=re.UNICODE,
)


def _is_translatable_text_core(text: str) -> bool:
    core = (text or "").strip()
    if not core:
        return False
    return bool(re.search(r"[\w\u00C0-\u1EF9]", core, flags=re.UNICODE))


def _collect_docx_special_text_roots(document):
    """Body, headers/footers, footnotes/endnotes — anywhere text boxes may hide."""
    roots = [document.element.body]
    try:
        for section in document.sections:
            roots.append(section.header._element)
            roots.append(section.footer._element)
    except Exception:
        pass
    try:
        for rel in document.part.rels.values():
            reltype = str(getattr(rel, "reltype", "") or "")
            if "footnotes" in reltype or "endnotes" in reltype:
                roots.append(rel.target_part.element)
    except Exception:
        pass
    return roots


def _iter_docx_special_text_nodes(root_element):
    """Yield text nodes in text boxes / WordArt (not exposed via python-docx Run.text)."""
    if root_element is None:
        return
    try:
        for t_el in root_element.xpath(
            './/*[local-name()="t" and ancestor::*[local-name()="txbxContent"]]'
        ):
            yield t_el, t_el.text or "", "txbx_wt"
    except Exception:
        pass
    try:
        for tp in root_element.xpath('.//*[local-name()="textpath"]'):
            val = tp.get(f"{{{_VML_NS}}}string") or tp.get("string")
            if val is not None:
                yield tp, val, "vml_textpath"
    except Exception:
        pass
    # DrawingML shape text outside txbxContent (legacy WordArt / floating shapes).
    try:
        for t_el in root_element.xpath(
            './/*[namespace-uri()="http://schemas.openxmlformats.org/drawingml/2006/main"'
            ' and local-name()="t"'
            ' and not(ancestor::*[local-name()="txbxContent"])]'
        ):
            raw = t_el.text or ""
            if (raw or "").strip():
                yield t_el, raw, "drawingml_t"
    except Exception:
        pass


def _iter_docx_extra_paragraphs(document):
    """Paragraphs in footnotes/endnotes (not exposed on Document.paragraphs)."""
    try:
        from docx.text.paragraph import Paragraph
    except Exception:
        return
    try:
        for rel in document.part.rels.values():
            reltype = str(getattr(rel, "reltype", "") or "")
            if "footnotes" not in reltype and "endnotes" not in reltype:
                continue
            part = rel.target_part
            for p_el in part.element.xpath('.//*[local-name()="p"]'):
                if p_el.xpath('.//*[local-name()="txbxContent"]'):
                    continue
                try:
                    yield Paragraph(p_el, part)
                except Exception:
                    continue
    except Exception:
        pass


def collect_docx_special_text_segments(document):
    """Collect translatable strings embedded in text boxes / WordArt."""
    segments = []
    seen = set()
    for root in _collect_docx_special_text_roots(document):
        for el, raw, kind in _iter_docx_special_text_nodes(root):
            core = (raw or "").strip()
            if not core or core in seen or not _is_translatable_text_core(core):
                continue
            seen.add(core)
            segments.append({"element": el, "raw": raw or "", "core": core, "kind": kind})
    return segments


def apply_docx_special_text_translations(
    document,
    segments,
    translated_value_fn,
    *,
    bilingual_mode=None,
    bilingual_delimiter=None,
    join_inline_fn=None,
):
    """Write translated text into textbox / WordArt XML nodes (all duplicate copies)."""
    from app.services.document_v2.pdf_docx_pipeline.layout_recovery import normalize_bilingual_mode

    bi_mode = normalize_bilingual_mode(bilingual_mode)
    delimiter = (bilingual_delimiter or "|").strip() or "|"
    translations = {}
    for seg in segments or []:
        core = seg.get("core") or ""
        if not core or core in translations:
            continue
        try:
            translated = translated_value_fn(core)
        except Exception:
            continue
        if translated is None:
            continue
        translated = str(translated)
        if not translated.strip() or translated.strip() == core:
            continue
        if bi_mode == "inline" and join_inline_fn:
            translations[core] = join_inline_fn(core, translated.strip(), delimiter)
        elif bi_mode == "newline":
            translations[core] = f"{core}\n{translated.strip()}"
        else:
            translations[core] = translated.strip()

    if not translations:
        return 0

    touched = 0
    for root in _collect_docx_special_text_roots(document):
        for t_el, raw, kind in _iter_docx_special_text_nodes(root):
            core = (raw or "").strip()
            translated = translations.get(core)
            if not translated:
                continue
            if raw != core:
                lead = raw[: len(raw) - len(raw.lstrip())]
                tail = raw[len(raw.rstrip()) :]
                out = f"{lead}{translated}{tail}"
            else:
                out = translated
            try:
                if kind == "vml_textpath":
                    t_el.set(f"{{{_VML_NS}}}string", out)
                elif kind == "drawingml_t":
                    t_el.text = out
                else:
                    t_el.text = out
                    try:
                        from docx.oxml.ns import qn as _qn
                        t_el.set(_qn("xml:space"), "preserve")
                    except Exception:
                        pass
                touched += 1
            except Exception:
                continue
    return touched


def dedupe_consecutive_identical_paragraphs(document) -> int:
    """Remove consecutive body paragraphs with identical text (common pdf2docx artifact)."""
    removed = 0
    prev_key = None
    for para in list(document.paragraphs):
        try:
            if _para_has_protected_xml(para):
                prev_key = None
                continue
        except Exception:
            pass
        try:
            texts = [r.text or "" for r in para.runs]
        except Exception:
            texts = []
        plain = join_docx_run_texts(texts).strip()
        if not plain:
            prev_key = None
            continue
        key = re.sub(r"\s+", " ", plain).strip().lower()
        if prev_key and key == prev_key:
            try:
                para._element.getparent().remove(para._element)
                removed += 1
                continue
            except Exception:
                pass
        prev_key = key
    return removed


def split_inline_bilingual_to_newline(paragraph, delimiter):
    """Move inline bilingual translation (source | trans) onto the next paragraph."""
    try:
        from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
            _find_inline_bilingual_split,
        )
    except Exception:
        return None

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    import copy as _copy

    d = (delimiter or "|").strip() or "|"
    p_elem = paragraph._element
    split = _find_inline_bilingual_split(p_elem, d)
    if split is None:
        return None
    ppr, src_run_els, trans_run_els = split

    new_p = OxmlElement("w:p")
    if ppr is not None:
        new_ppr = _copy.deepcopy(ppr)
        num_pr = new_ppr.find(_qn("w:numPr"))
        if num_pr is not None:
            new_ppr.remove(num_pr)
        new_p.insert(0, new_ppr)

    for r_el in trans_run_els:
        new_p.append(_copy.deepcopy(r_el))

    for child in list(p_elem):
        if child.tag != _qn("w:pPr"):
            p_elem.remove(child)
    for r_el in src_run_els:
        p_elem.append(_copy.deepcopy(r_el))

    p_elem.addnext(new_p)
    try:
        from docx.text.paragraph import Paragraph

        trans_para = Paragraph(new_p, paragraph._parent)
        try:
            from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                apply_post_split_newline_layout,
            )

            apply_post_split_newline_layout(paragraph, trans_para)
        except Exception:
            pass
        return trans_para
    except Exception:
        return None


def process_docx(service, file_path, target_lang, progress_callback=None, *, ocr_images=False, ocr_langs=None, ocr_mode=None, bilingual_mode=None, bilingual_delimiter=None, from_pdf=False):
    """Translate DOCX while preserving original formatting, layout, images.
    
    Bilingual modes:
      - none: normal translation (replace original with translation)
      - inline: song ngữ liền kề (Original | Translated in same paragraph)
      - newline: song ngữ xuống dòng (inline path first, then split translation to line below)
      - preserve_layout: alias for 'inline' mode (dịch song ngữ liền kề, giữ layout)
      - line_by_line: alias for 'newline' mode (dịch song ngữ xuống dòng)
    """
    # python-docx rejects macro-enabled documents (.docm content type).
    # Strip the macros by re-packaging as a standard .docx before opening.
    try:
        doc = docx.Document(file_path)
    except ValueError:
        import tempfile, zipfile as _zf
        _tmp = tempfile.mktemp(suffix='.docx')
        with _zf.ZipFile(file_path, 'r') as zin, _zf.ZipFile(_tmp, 'w') as zout:
            for item in zin.infolist():
                # Skip VBA / macro parts entirely
                if 'vbaProject' in item.filename or 'vbaData' in item.filename:
                    continue
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    data = data.replace(
                        b'application/vnd.ms-word.document.macroEnabled.main+xml',
                        b'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml',
                    )
                # Remove relationship entries that reference VBA parts
                if item.filename.endswith('.rels'):
                    data = re.sub(rb'<Relationship[^>]*Target="[^"]*vba[^"]*"[^/]*/>', b'', data, flags=re.IGNORECASE)
                zout.writestr(item, data)
        doc = docx.Document(_tmp)
        os.remove(_tmp)

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy
        pass
    except Exception:
        pass

    api_only = str(os.getenv('AI_DISABLE_FALLBACK', '0')).strip().lower() in ('1', 'true', 'yes', 'on')

    try:
        from app.services.document_v2.pdf_docx_pipeline.layout_recovery import normalize_bilingual_mode
    except Exception:
        def normalize_bilingual_mode(mode):
            bi = (str(mode or "").strip().lower() or "none")
            if bi in ("preserve_layout", "inline"):
                return "inline"
            if bi in ("line_by_line", "newline"):
                return "newline"
            return "none"

    bi_mode = normalize_bilingual_mode(bilingual_mode)

    mode = (str(ocr_mode).strip().lower() if ocr_mode else 'image')
    if mode not in ('image', 'text', 'both', 'auto'):
        mode = 'auto'

    def _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode=None):
        try:
            raw = (ocr_text or '').strip()
            if not raw:
                return 'text'

            char_count = len(raw)
            words = re.findall(r'\w+', raw, flags=re.UNICODE)
            word_count = len(words)

            if char_count >= 120 or word_count >= 25:
                print(f"  [MODE] Prose detected (chars={char_count}, words={word_count}), AI={ai_recommended_mode} -> text")
                return 'text'

            lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
            low_raw = raw.lower()
            low_trans = (translated_text or '').lower()
            promo_keywords = (
                'sale', 'discount', 'offer', 'book now', 'vacation',
                'summer', 'up to', '% off', 'promo', 'hotline',
                'free', 'limited', 'special', 'deal', 'subscribe',
            )
            has_promo = any(k in low_raw or k in low_trans for k in promo_keywords)

            alpha_chars = [ch for ch in raw if ch.isalpha()]
            upper_ratio = (
                sum(1 for ch in alpha_chars if ch == ch.upper()) / max(1, len(alpha_chars))
                if alpha_chars else 0.0
            )
            line_word_counts = [len(re.findall(r'\w+', ln, flags=re.UNICODE)) for ln in lines] if lines else [0]
            avg_wpl = (sum(line_word_counts) / len(line_word_counts)) if line_word_counts else 0.0
            short_lines = sum(1 for c in line_word_counts if c <= 3)

            looks_banner = (
                has_promo or
                (upper_ratio >= 0.50 and avg_wpl <= 4) or
                (short_lines >= 3 and avg_wpl <= 3)
            )

            ai_mode = (ai_recommended_mode or '').lower()

            if looks_banner or ai_mode == 'image':
                final = 'image'
            else:
                final = 'text'

            print(
                f"  [MODE] AI={ai_mode}, banner={looks_banner}, "
                f"chars={char_count}, words={word_count}, upper={upper_ratio:.2f} -> {final}"
            )
            return final
        except Exception:
            return 'text'

    def iter_all_paragraphs(document):
        paras = []
        try:
            paras.extend(list(document.paragraphs))
        except Exception:
            pass
        try:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paras.extend(list(cell.paragraphs))
        except Exception:
            pass
        try:
            for section in document.sections:
                paras.extend(list(section.header.paragraphs))
                paras.extend(list(section.footer.paragraphs))
        except Exception:
            pass
        return paras

    def paragraph_image_rids(paragraph):
        rids = []
        try:
            runs = list(paragraph.runs)
        except Exception:
            runs = []
        if not runs:
            return rids

        rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        for run in runs:
            try:
                blips = run._element.xpath('.//*[local-name()="blip"]')
            except Exception:
                blips = []
            for blip in blips:
                try:
                    rid = blip.get(rel_attr)
                except Exception:
                    rid = None
                if rid:
                    rids.append(rid)
        seen = set()
        out = []
        for rid in rids:
            if rid in seen:
                continue
            seen.add(rid)
            out.append(rid)
        return out

    def rid_to_image_part(paragraph, rid):
        try:
            part = paragraph.part
            related = getattr(part, 'related_parts', None)
            if isinstance(related, dict) and rid in related:
                return related[rid]
        except Exception:
            pass
        try:
            rels = getattr(paragraph.part, 'rels', None)
            if rels and rid in rels:
                return rels[rid].target_part
        except Exception:
            pass
        return None

    def _collect_header_footer_image_partnames(document):
        protected = set()
        try:
            for section in document.sections:
                for hf in (section.header, section.footer):
                    part = getattr(hf, 'part', None)
                    related = getattr(part, 'related_parts', None)
                    if not isinstance(related, dict):
                        continue
                    for _rid, target in related.items():
                        try:
                            ct = str(getattr(target, 'content_type', '') or '').lower()
                            if not ct.startswith('image/'):
                                continue
                            pn = str(getattr(target, 'partname', '') or '').lstrip('/')
                            if pn:
                                protected.add(pn)
                        except Exception:
                            continue
        except Exception:
            pass
        return protected

    def replace_image_with_text(paragraph, rid, translated_text):
        txt = (translated_text or '').strip()
        if not txt:
            return False

        rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        replaced = False
        try:
            runs = list(paragraph.runs)
        except Exception:
            runs = []

        for run in runs:
            try:
                blips = run._element.xpath('.//*[local-name()="blip"]')
            except Exception:
                blips = []
            if not blips:
                continue

            has_target = False
            for blip in blips:
                try:
                    if blip.get(rel_attr) == rid:
                        has_target = True
                        break
                except Exception:
                    continue
            if not has_target:
                continue

            try:
                drawings = run._element.xpath('./*[local-name()="drawing"]')
                for dr in drawings:
                    parent = dr.getparent()
                    if parent is not None:
                        parent.remove(dr)
            except Exception:
                pass
            run.text = ""
            replaced = True
            break

        try:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        except Exception:
            pass

        try:
            new_run = paragraph.add_run(txt)
            _ = new_run
            replaced = True
        except Exception:
            pass

        if not replaced:
            try:
                paragraph.add_run(txt)
                replaced = True
            except Exception:
                replaced = False
        return replaced

    def _normalize_ocr_text_for_docx(text):
        raw = (text or '').replace('\r\n', '\n').replace('\r', '\n')
        if not raw.strip():
            return ''

        cleaned_lines = []
        for ln in raw.split('\n'):
            ln2 = re.sub(r'\s+', ' ', (ln or '').strip())
            if not ln2:
                continue
            if len(ln2) <= 1 and not re.search(r'[0-9]', ln2):
                continue
            cleaned_lines.append(ln2)

        if not cleaned_lines:
            return ''

        out_parts = []
        cur = ''
        for ln in cleaned_lines:
            if not cur:
                cur = ln
                continue

            end_punct = cur.endswith(('.', '!', '?', ':', ';'))
            starts_bullet = bool(re.match(r'^(\-|\*|\d+[\.)])\s+', ln))
            if end_punct or starts_bullet:
                out_parts.append(cur)
                cur = ln
            else:
                cur = f"{cur} {ln}".strip()

        if cur:
            out_parts.append(cur)

        normalized = '\n'.join(out_parts)
        return normalized.strip()

    _pdf_subset_font_re = re.compile(r"^[A-Z]{6}\+(.+)$")
    _symbol_font_hint_re = re.compile(r"(symbol|wingdings|webdings|dingbats?|zapf|mt\s*extra)", re.IGNORECASE)

    def _strip_pdf_subset_font_prefix(name: str) -> str:
        out = str(name or '').strip()
        if not out:
            return ''
        while True:
            m = _pdf_subset_font_re.match(out)
            if not m:
                break
            nxt = str(m.group(1) or '').strip()
            if not nxt or nxt == out:
                break
            out = nxt
        return out

    def _run_font_names(run) -> list[str]:
        names = []
        seen = set()

        def _push(value):
            v = _strip_pdf_subset_font_prefix(value)
            if not v:
                return
            key = v.lower()
            if key in seen:
                return
            seen.add(key)
            names.append(v)

        try:
            _push(getattr(getattr(run, 'font', None), 'name', None))
        except Exception:
            pass

        try:
            rpr = run._element.find(_qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(_qn('w:rFonts'))
                if rfonts is not None:
                    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                        _push(rfonts.get(_qn(attr)))
        except Exception:
            pass

        return names

    def _normalize_run_subset_fonts(run) -> int:
        changed = 0
        try:
            current = getattr(getattr(run, 'font', None), 'name', None)
            cleaned = _strip_pdf_subset_font_prefix(current)
            if cleaned and cleaned != (current or ''):
                run.font.name = cleaned
                changed += 1
        except Exception:
            pass

        try:
            rpr = run._element.find(_qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(_qn('w:rFonts'))
                if rfonts is not None:
                    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                        cur = rfonts.get(_qn(attr))
                        if not cur:
                            continue
                        cleaned = _strip_pdf_subset_font_prefix(cur)
                        if cleaned and cleaned != cur:
                            rfonts.set(_qn(attr), cleaned)
                            changed += 1
        except Exception:
            pass

        return changed

    def _run_uses_symbol_font(run) -> bool:
        for name in _run_font_names(run):
            if _symbol_font_hint_re.search(name or ''):
                return True
        return False

    def _pick_text_target_run(runs):
        if not runs:
            return None
        best = None
        best_score = -10**9

        for run in runs:
            txt = run.text or ''
            token_chars = len(re.findall(r"[A-Za-zÀ-ỹ0-9]", txt, flags=re.UNICODE))
            ws = len((txt or '').strip())
            score = token_chars * 10 + ws
            if not re.search(r"\w", txt or '', flags=re.UNICODE):
                score -= 8
            if _run_uses_symbol_font(run):
                score -= 10000
            if best is None or score > best_score:
                best = run
                best_score = score

        if best is not None and not _run_uses_symbol_font(best):
            return best

        for run in runs:
            if not _run_uses_symbol_font(run):
                return run
        return runs[0]

    def _ensure_text_font_compat(run, text: str) -> None:
        if run is None:
            return
        _normalize_run_subset_fonts(run)

        raw = text or ''
        has_text_tokens = bool(re.search(r"[A-Za-zÀ-ỹ0-9]", raw, flags=re.UNICODE))
        has_list_marker = bool(re.search(r"[+\-•*]", raw))
        if not (has_text_tokens or has_list_marker):
            return

        if bool(from_pdf):
            try:
                from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                    _sanitize_run_text_font,
                )

                _sanitize_run_text_font(run, raw)
                return
            except Exception:
                pass

        if str(os.getenv('PDF_DOCX_FIX_SYMBOL_FONT_RUNS', '1')).strip().lower() not in ('1', 'true', 'yes', 'on'):
            return
        if not _run_uses_symbol_font(run):
            return

        fallback = str(os.getenv('PDF_DOCX_TEXT_FONT_FALLBACK', 'Times New Roman') or '').strip()
        if not fallback:
            return

        try:
            run.font.name = fallback
        except Exception:
            pass

        try:
            rpr = run._element.find(_qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run._element.insert(0, rpr)
            rfonts = rpr.find(_qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                rfonts.set(_qn(attr), fallback)
        except Exception:
            pass

    def _normalize_subset_fonts_in_document(document):
        fixed = 0
        for para in iter_all_paragraphs(document):
            try:
                runs = list(para.runs)
            except Exception:
                runs = []
            for run in runs:
                try:
                    fixed += _normalize_run_subset_fonts(run)
                except Exception:
                    continue
        return fixed

    DB_IDENTIFIER_MAP = {
        "ma_khach_hang": "customer_id",
        "ngay_ban": "sale_date",
    }

    def _env_on(name: str, default: str = "0") -> bool:
        return str(os.getenv(name, default)).strip().lower() in ("1", "true", "yes", "on")

    def _apply_db_identifier_map(text: str) -> str:
        out = text or ""
        for src, dst in DB_IDENTIFIER_MAP.items():
            out = re.sub(rf"\b{re.escape(src)}\b", dst, out, flags=re.IGNORECASE)
        return out

    def _fix_duplicate_label_punctuation(text: str) -> str:
        """Fix 'Abstract. .', 'Abstrak..', 'Keywords: :', and 'Fig. 1. . body' artifacts."""
        out = text or ""
        out = re.sub(r"([:：])\s+[:：]", r"\1", out)
        # Double period with optional whitespace: ". ." / ".." after labels
        out = re.sub(r"(\.)\s*\.+(?=\s|$|\w)", r".", out)
        # Abbrev caption: Fig. . 1 -> Fig. 1
        out = re.sub(
            r"(\b(?:Fig\.|Tab\.|Abb\.|Figuur|Tabel|Figuut|Hình|Bảng|Illustration|Afbeelding|Illustrasie))\s+\.\s*(\d)",
            r"\1 \2",
            out,
            flags=re.IGNORECASE,
        )
        # Known label words: collapse trailing duplicate periods (Abstrak.. / Fig. 1..)
        out = re.sub(
            r"(\b(?:Abstract|Abstrak|Opsomming|T[oóô]m\s*t[aắ]t|Keywords|Sleutelwoorde|T[uừ]\s*kh[oó]a"
            r"|Fig\.|Tab\.|Abb\.|Figuur|Hình|Bảng|Illustration|Afbeelding|Illustrasie|Tabel|Figuut)"
            r"(?:\s+\d+(?:[\.:\-–—]\d+)*)?)"
            r"\.{2,}",
            r"\1.",
            out,
            flags=re.IGNORECASE,
        )
        return out

    def _fix_method_call_tail_artifacts(text: str) -> str:
        """Fix orphaned '()' runs merged as '.().' after translation (e.g. viewLichSuMuon().())."""
        out = text or ""
        if not out.strip():
            return out
        out = re.sub(r"(\w\))\.\(\)\.?$", r"\1()", out)
        out = re.sub(r"(\w\))\.\(\)\(\)\.?$", r"\1()", out)
        out = re.sub(r"(\w\))\.+\(\)\.?$", r"\1()", out)
        return out

    def _fix_known_proper_noun_mistranslations(text: str) -> str:
        """Correct common literal translations of Vietnamese proper nouns/form labels."""
        out = text or ""
        replacements = (
            (r"\bSOUTH\s+CAN\s+THO\s+UNIVERSITY\b", "NAM CAN THO UNIVERSITY"),
            (r"\bSouth\s+Can\s+Tho\s+University\b", "Nam Can Tho University"),
            (r"\bSOUTH\s+CAN\s+THO\b", "NAM CAN THO"),
            (r"\bSouth\s+Can\s+Tho\b", "Nam Can Tho"),
            (r"\bMODEL\s+GT\s*([0-9]+)\b", r"FORM GT \1"),
            (r"\bModel\s+GT\s*([0-9]+)\b", r"Form GT \1"),
        )
        for pattern, replacement in replacements:
            out = re.sub(pattern, replacement, out)
        return out

    def _strip_duplicate_leading_punct(label: str, body: str) -> str:
        if not body:
            return body
        label_end = (label or "").rstrip()
        rest = (body or "").lstrip()
        if label_end.endswith(":") or label_end.endswith("："):
            rest = re.sub(r"^[:：\s]+", "", rest)
        elif label_end.endswith("."):
            rest = re.sub(r"^[\.\s]+", "", rest)
        return rest

    def _normalize_label_text(label: str) -> str:
        """Ensure label ends with at most one trailing . or : (no 'Abstrak. .' / 'Fig. 1. .')."""
        out = (label or "").rstrip()
        if not out:
            return out
        out = _fix_duplicate_label_punctuation(out)
        if out.endswith("."):
            out = re.sub(r"\.+$", ".", out)
        elif out.endswith(":") or out.endswith("："):
            out = re.sub(r"[:：]+$", ":", out) if out.endswith(":") else out
        return out

    def _finalize_label_body_pair(label: str, body: str) -> tuple[str, str]:
        """Normalize label/body boundary so runs don't render as 'Fig. 1. . text'."""
        lbl = _normalize_label_text(label or "")
        bdy = _strip_duplicate_leading_punct(lbl, body or "")
        bdy = re.sub(r"^[\.\s,_]+", "", (bdy or "").lstrip())
        # Caption: label ends with "1." — body must not restart with "." or "._"
        if _CAPTION_LABEL_HEAD_RE.match(lbl.strip()) or re.search(r"\d\.\s*$", lbl):
            bdy = re.sub(r"^[\.\s]+", "", bdy)
        if bdy and lbl:
            if not lbl.endswith((" ", "\t")):
                lbl = lbl.rstrip() + " "
        return lbl, bdy

    def _cleanup_translated_text(text: str) -> str:
        out = "" if text is None else str(text)
        out = out.replace("\u200b", "").replace("\ufeff", "")
        out = strip_run_boundary_markers(out)
        out = _fix_duplicate_label_punctuation(out)
        out = _fix_method_call_tail_artifacts(out)
        out = _fix_known_proper_noun_mistranslations(out)
        try:
            out = service._sanitize_text(out)
        except Exception:
            pass
        if not _env_on("DOCX_DOMAIN_FIXUPS", "0"):
            return out
        out = re.sub(r"\bNAMEBUILDING\b", "NAME BUILDING", out, flags=re.IGNORECASE)
        out = _apply_db_identifier_map(out)
        out = re.sub(r"\bNot\s+nul\b", "Not null", out, flags=re.IGNORECASE)
        out = re.sub(r"\bInfo\s+tin\s+basic\b", "Basic Information", out, flags=re.IGNORECASE)
        out = re.sub(r"\bWhere\s+the\s+topic\s+is\s+applied\b", "Application of the project", out, flags=re.IGNORECASE)
        out = re.sub(r"\bDevelopment\s+direction\s*:\s*Is\s+there\b", "Development direction: Yes", out, flags=re.IGNORECASE)
        out = re.sub(r"\bSTUDENT\s+ID\b", "Student ID", out, flags=re.IGNORECASE)
        out = re.sub(r"\bData\s+types\b", "Data Types", out, flags=re.IGNORECASE)
        out = re.sub(r"\bKHOA\s*CÔNG\s*NGHỆ\s*THÔNG\s*TIN\b", "FACULTY OF INFORMATION TECHNOLOGY", out, flags=re.IGNORECASE)
        out = re.sub(r"Gửi\s+lại\s+phiếu\s+đăng\s+ký\s+qua\s+Email\s*:", "Resubmit the registration form via Email:", out, flags=re.IGNORECASE)
        try:
            out = service._sanitize_text(out)
        except Exception:
            pass
        return out

    _guard_enabled = str(os.getenv("DOCX_TRANSLATION_GUARD", "1")).strip().lower() in ("1", "true", "yes", "on")
    _guard_url_re = re.compile(r"(https?://[^\s<>()]+|www\.[^\s<>()]+)", re.IGNORECASE)
    _guard_email_re = re.compile(r"\b[\w.+\-]+@[\w.\-]+\.[A-Za-z]{2,}\b")
    _guard_doi_re = re.compile(r"\b10\.\d{4,9}/[^\s\"'<>]+", re.IGNORECASE)
    _guard_doi_label_re = re.compile(r"\bdoi\s*:\s*10\.\d{4,9}/[^\s\"'<>]+", re.IGNORECASE)
    _guard_arxiv_re = re.compile(r"\barxiv\s*:\s*\S+", re.IGNORECASE)
    _guard_ref_line_re = re.compile(r"^\s*(\[\d+\]|\d{1,3}[.)])\s+\S+")
    _guard_code_block_re = re.compile(r"```.+?```", re.DOTALL)
    _guard_inline_code_re = re.compile(r"`[^`]+`")

    # Unicode math: Greek letters, Math Operators block (∀∂∃∈∉∑∏∫√∞≤≥≠≈...), Arrows,
    # ± × ÷ · superscript/subscript digits, etc.
    _UNICODE_MATH_RE = re.compile(
        r"[\u0391-\u03C9"       # Greek Α–ω
        r"\u2200-\u22FF"         # Mathematical Operators (∀∂∃∄∅∆∇∈∉∊∋∌∍∎∏∐∑−∓∔∕∖∗∘∙√∝∞∟∠∡∢∣∤∥∦∧∨∩∪∫∬∭∮∯∰∱∲∳∴∵∶∷∸∹∺∻∼∽∾∿≀≁≂≃≄≅≆≇≈≉≊≋≌≍≎≏≐≑≒≓≔≕≖≗≘≙≚≛≜≝≞≟≠≡≢≣≤≥≦≧≨≩≪≫≬≭≮≯≰≱≲≳≴≵≶≷≸≹≺≻≼≽≾≿⊀⊁⊂⊃⊄⊅⊆⊇⊈⊉⊊⊋⊌⊍⊎⊏⊐⊑⊒⊓⊔⊕⊖⊗⊘⊙⊚⊛⊜⊝⊞⊟⊠⊡⊢⊣⊤⊥⊦⊧⊨⊩⊪⊫⊬⊭⊮⊯⊰⊱⊲⊳⊴⊵⊶⊷⊸⊹⊺⊻⊼⊽⊾⊿⋀⋁⋂⋃⋄⋅⋆⋇⋈⋉⋊⋋⋌⋍⋎⋏⋐⋑⋒⋓⋔⋕⋖⋗⋘⋙⋚⋛⋜⋝⋞⋟⋠⋡⋢⋣⋤⋥⋦⋧⋨⋩⋪⋫⋬⋭⋮⋯⋰⋱)"
        r"\u2190-\u21FF"         # Arrows ←↑→↓↔⇒⇐⇔...
        r"\u00B1\u00D7\u00F7\u00B7"  # ± × ÷ ·
        r"\u00B2\u00B3\u00B9"   # ² ³ ¹
        r"\u2070-\u2079"         # Superscript digits ⁰–⁹
        r"\u2080-\u2089"         # Subscript digits ₀–₉
        r"\u221A-\u221F"         # √ ∛ ∜ ∝ ∞ ∟ ∠ ∡ ∢
        r"\u2248\u2260\u2264\u2265"  # ≈ ≠ ≤ ≥
        r"]"
    )

    def _looks_like_formula(text: str) -> bool:
        s = (text or "").strip()
        if not s:
            return False
        # Unicode math symbols (Greek, operators, arrows, etc.) → treat as formula.
        unicode_math = _UNICODE_MATH_RE.findall(s)
        if unicode_math:
            # Any paragraph that has Unicode math and is not clearly a Vietnamese/English sentence.
            latin_letters = re.findall(r"[A-Za-zÀ-ỹ]", s)
            # If Unicode math chars outnumber or equal Latin letters → formula
            if len(unicode_math) >= len(latin_letters) or len(unicode_math) >= 2:
                return True
        # ASCII math operators.
        math_syms = re.findall(r"[=+\-*/^_{}\\<>]", s)
        if len(math_syms) >= 2:
            letters = len(re.findall(r"[A-Za-zÀ-ỹ]", s))
            if letters <= 2:
                return True
            if len(math_syms) >= 4 and letters < len(math_syms):
                return True
        if re.search(r"\b\d+\s*[+\-*/]\s*\d+\b", s) and not re.search(r"[A-Za-zÀ-ỹ]", s):
            return True
        return False

    def _looks_like_code(text: str) -> bool:
        s = (text or "").strip()
        if not s:
            return False
        if _guard_code_block_re.search(s) or _guard_inline_code_re.search(s):
            return True
        if re.search(r"\b(def|class|function|var|let|const|public|private)\b", s):
            return True
        if re.search(r"[{}<>;]|->|=>|::", s) and len(s) <= 200:
            letters = len(re.findall(r"[A-Za-zÀ-ỹ]", s))
            if letters <= 4:
                return True
        return False

    _ref_heading_re = re.compile(
        r"^\s*(references|bibliography|works\s+cited|tài\s+liệu\s+tham\s+khảo|tai\s+lieu\s+tham\s+khao)\s*:?\s*$",
        re.IGNORECASE | re.UNICODE,
    )
    _ref_entry_line_re = re.compile(r"^\s*(?:\[\d+\]|\d{1,3}[.)])\s+\S")

    def _looks_like_english_reference_block(text: str) -> bool:
        """Whole paragraph looks like an English bibliography — pointless to translate to EN."""
        s = (text or "").strip()
        if len(s) < 80:
            return False
        try:
            if service._looks_vietnamese_like_text(s):
                return False
        except Exception:
            if _has_vietnamese_diacritics(s):
                return False
        low = s.lower()
        if "doi:" in low or "doi.org/" in low or "arxiv:" in low:
            return True
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        if len(lines) < 2:
            return False
        numbered = sum(1 for ln in lines if _ref_entry_line_re.match(ln))
        if numbered >= 2:
            return True
        if len(re.findall(r"\(\s*\d{4}\s*[a-z]?\)", s)) >= 5:
            return True
        return False

    def _references_skip_translation_whole(text: str) -> bool:
        tl = str(target_lang).strip().lower()
        if not tl.startswith("en"):
            return False
        return _looks_like_english_reference_block(text)

    def _split_reference_entries(text: str):
        lines = (text or "").splitlines()
        chunks = []
        cur = []
        for ln in lines:
            if _ref_entry_line_re.match(ln):
                if cur:
                    chunks.append("\n".join(cur).strip())
                cur = [ln]
            else:
                cur.append(ln)
        if cur:
            chunks.append("\n".join(cur).strip())
        out = [c for c in chunks if c.strip()]
        return out if len(out) >= 2 else []

    def _paragraph_is_multi_reference_list(text: str) -> bool:
        return len(_split_reference_entries(text)) >= 2

    def _skip_reference_entry_translation(entry: str) -> bool:
        """Skip API translation for headings / citation lines already in English (or non‑VI when target is VI)."""
        e = (entry or "").strip()
        if not e:
            return True
        lines = e.splitlines()
        first = lines[0].strip() if lines else ""
        if _ref_heading_re.match(first) and not _ref_entry_line_re.search(e):
            return True
        citation_like = bool(
            _ref_entry_line_re.search(e)
            or "doi:" in e.lower()
            or "doi.org" in e.lower()
            or re.search(r"\b(?:19|20)\d{2}\b", e)
        )
        if not citation_like or len(e) < 22:
            return False
        try:
            looks_vi = service._looks_vietnamese_like_text(e)
        except Exception:
            looks_vi = _has_vietnamese_diacritics(e)
        return not looks_vi

    def _should_skip_translation(text: str) -> bool:
        s = (text or "").strip()
        if _references_skip_translation_whole(s):
            return True
        if not _guard_enabled:
            return False
        if not s:
            return False
        if _guard_url_re.fullmatch(s) or _guard_email_re.fullmatch(s):
            return True
        if _guard_doi_label_re.fullmatch(s) or _guard_doi_re.fullmatch(s):
            return True
        if _guard_arxiv_re.fullmatch(s):
            return True
        if _guard_ref_line_re.match(s):
            low = s.lower()
            if "doi" in low or "arxiv" in low:
                return True
            # Only skip numbered items that look like actual bibliography references:
            # require BOTH a year-like number AND a comma (e.g. "Smith, J. (2020). Title.")
            # Plain numbered list items like "1. First item, second part" must still be translated.
            if re.search(r"\b(19|20)\d{2}\b", s) and "," in s:
                return True
        if _looks_like_formula(s):
            return True
        if _looks_like_code(s):
            return True
        return False

    def _paragraph_has_superscript_runs(paragraph) -> bool:
        if paragraph is None:
            return False
        try:
            from docx.oxml.ns import qn as _qn
            for run in paragraph.runs:
                try:
                    if bool(getattr(getattr(run, "font", None), "superscript", False)):
                        return True
                except Exception:
                    pass
                try:
                    rpr = run._element.find(_qn("w:rPr"))
                    if rpr is not None:
                        va = rpr.find(_qn("w:vertAlign"))
                        if va is not None:
                            v = str(va.get(_qn("w:val"), "") or "").strip().lower()
                            if v in ("superscript", "subscript"):
                                return True
                except Exception:
                    pass
        except Exception:
            pass
        return False

    def _should_preserve_author_metadata(paragraph, text: str) -> bool:
        """Keep author name lines (with affiliation superscripts) unchanged."""
        s = (text or "").strip()
        if not s or len(s) > 320:
            return False
        if re.search(r"\borcid\b", s, re.IGNORECASE):
            return True
        if _paragraph_has_superscript_runs(paragraph):
            if s.count(",") >= 1 or re.search(r"\b(?:and|en|và|va)\b", s, re.IGNORECASE):
                if len(s) <= 240:
                    return True
        return False

    _guard_token_re = re.compile(r"\uE000(\d+)\uE001")
    _corrupted_guard_re = re.compile(r"__?\s*[\wÀ-Ỹ]{2,12}\s*_\s*(\d+)\s*__?", re.IGNORECASE)

    def _guard_placeholder_key(index: int) -> str:
        return f"\uE000{index}\uE001"

    def _guard_mask_tokens(text: str) -> tuple[str, dict]:
        if not _guard_enabled:
            return text, {}
        placeholders = {}

        def _sub(pattern, src: str) -> str:
            def repl(match):
                key = _guard_placeholder_key(len(placeholders))
                placeholders[key] = match.group(0)
                return key

            return pattern.sub(repl, src)

        out = text
        out = _sub(_guard_doi_label_re, out)
        out = _sub(_guard_doi_re, out)
        out = _sub(_guard_arxiv_re, out)
        out = _sub(_guard_url_re, out)
        out = _sub(_guard_email_re, out)
        return out, placeholders

    def _guard_restore_tokens(text: str, placeholders: dict) -> str:
        out = text or ""
        indexed_values = {}
        for key, value in (placeholders or {}).items():
            out = out.replace(key, value)
            m = _guard_token_re.search(key)
            if m:
                indexed_values[int(m.group(1))] = value

        if indexed_values:
            def _restore_corrupted(match):
                idx = int(match.group(1))
                return indexed_values.get(idx, match.group(0))

            out = _corrupted_guard_re.sub(_restore_corrupted, out)
        return out

    def _translate_preserve_exact_lines(text):
        raw = text or ""
        if not raw:
            return raw
        parts = re.split(r"(\r\n|\r|\n)", raw)
        out = []
        for part in parts:
            if part in ("\r\n", "\r", "\n"):
                out.append(part)
                continue
            if not part.strip():
                out.append(part)
                continue
            m = re.match(r"^(\s*)(.*?)(\s*)$", part, flags=re.DOTALL)
            if m:
                lead, core, tail = m.group(1), m.group(2), m.group(3)
            else:
                lead, core, tail = "", part, ""
            if not core.strip():
                out.append(part)
                continue
            translated_core = _translate_preserve_form_leaders(core)
            out.append(f"{lead}{translated_core}{tail}")
        return "".join(out)

    def _handle_multi_reference_bilingual(paragraph, paragraph_text: str) -> None:
        """Split numbered bibliography into entries; skip EN citations; newline → one § per entry below."""
        if bi_mode not in ("inline", "newline"):
            return
        entries = _split_reference_entries(paragraph_text)
        if len(entries) < 2:
            return
        if _references_skip_translation_whole(paragraph_text):
            return

        if bi_mode == "newline":
            anchor = paragraph
            for entry in entries:
                es = entry.strip()
                if _skip_reference_entry_translation(es):
                    continue
                try:
                    te = _translate_preserve_exact_lines(es).strip()
                except ProviderRateLimitError:
                    raise
                except Exception:
                    continue
                if not te or te == es:
                    continue
                np = _insert_paragraph_after(anchor, te, italic=False)
                if np is not None:
                    anchor = np
                    try:
                        _seen_para_elems.add(id(np))
                    except Exception:
                        pass
                else:
                    _append_translation_linebreak(anchor, te, italic=False)
            return

        # inline: one line per reference — "entry | trans"
        parts_out = []
        for entry in entries:
            es = entry.strip()
            if _skip_reference_entry_translation(es):
                parts_out.append(es)
                continue
            try:
                te = _translate_preserve_exact_lines(es).strip()
            except ProviderRateLimitError:
                raise
            except Exception:
                parts_out.append(es)
                continue
            if not te or te == es:
                parts_out.append(es)
            else:
                parts_out.append(service._join_inline_bilingual(es, te, bilingual_delimiter))
        combined = "\n".join(parts_out)
        _set_paragraph_text_preserve_runs(paragraph, combined)

    def image_part_ext(image_part):
        try:
            partname = str(getattr(image_part, 'partname', '') or '')
            base = os.path.basename(partname)
            ext = os.path.splitext(base)[1].lower()
            if ext:
                return ext
        except Exception:
            pass
        try:
            ct = str(getattr(image_part, 'content_type', '') or '').lower()
            mapping = {
                'image/png': '.png',
                'image/jpeg': '.jpg',
                'image/jpg': '.jpg',
                'image/gif': '.gif',
                'image/bmp': '.bmp',
                'image/tiff': '.tif',
                'image/webp': '.webp',
            }
            return mapping.get(ct, '.png')
        except Exception:
            return '.png'

    def _overlay_bytes_to_original_format(png_bytes: bytes, desired_ext: str) -> bytes:
        desired_ext = (desired_ext or '.png').lower()
        try:
            from PIL import Image
        except Exception:
            return png_bytes

        fmt_map = {
            '.png': 'PNG',
            '.jpg': 'JPEG',
            '.jpeg': 'JPEG',
            '.bmp': 'BMP',
            '.tif': 'TIFF',
            '.tiff': 'TIFF',
            '.webp': 'WEBP',
            '.gif': 'PNG',
        }
        out_fmt = fmt_map.get(desired_ext, 'PNG')
        try:
            img = Image.open(io.BytesIO(png_bytes))
            if out_fmt in ('JPEG', 'BMP', 'TIFF'):
                if img.mode not in ('RGB', 'L'):
                    img = img.convert('RGB')
            buf = io.BytesIO()
            img.save(buf, format=out_fmt)
            return buf.getvalue()
        except Exception:
            return png_bytes

    def _copy_run_rpr_to_run(src_run, dst_run):
        copy_run_rpr_preserve_all(src_run, dst_run)

    def _write_text_with_line_breaks(paragraph, style_run, text):
        normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        if "\n" not in normalized:
            _ensure_text_font_compat(style_run, normalized)
            style_run.text = normalized
            return

        chunks = normalized.split("\n")
        _ensure_text_font_compat(style_run, chunks[0] or "")
        style_run.text = chunks[0] or ""
        for chunk in chunks[1:]:
            br = paragraph.add_run("")
            _copy_run_rpr_to_run(style_run, br)
            br.add_break()
            if chunk:
                rr = paragraph.add_run(chunk)
                _copy_run_rpr_to_run(style_run, rr)
                _ensure_text_font_compat(rr, chunk)

    _ACADEMIC_LABEL_HEAD_RE = re.compile(
        r"^\s*(abstract|abstrakt|abstrak|opsomming|t[oóô]m\s*t[aắ]t|t[oóô]mt[aắ]t|"
        r"keywords?|sleutelwoorde?|t[uừ]\s*kh[oó]a|schl[uü]sselw[oö]rter)\s*[:：\.]",
        re.IGNORECASE,
    )
    _CAPTION_LABEL_HEAD_RE = re.compile(
        r"^\s*("
        r"figure|figures|fig\.?|figuur|figuut|"
        r"hình|hinh|"
        r"table|tables|tab\.?|tabel|"
        r"bảng|bang|"
        r"chart|charts|diagram|diagrams|"
        r"image|images|ảnh|anh|"
        r"graph|graphs|"
        r"abb\.?|abbildung|bild|"
        r"afbeelding|illustration|illustrasie"
        r")\s*"
        r"[\dIVXLCivxlc]+(?:[\.:\-–—][\dIVXLCivxlc]+)*"
        r"\s*[:：\.]",
        re.IGNORECASE,
    )
    _CAPTION_LABEL_SPLIT_RE = re.compile(
        r"^(\s*(?:"
        r"figure|figures|fig\.?|figuur|figuut|"
        r"hình|hinh|"
        r"table|tables|tab\.?|tabel|"
        r"bảng|bang|"
        r"chart|charts|diagram|diagrams|"
        r"image|images|ảnh|anh|"
        r"graph|graphs|"
        r"abb\.?|abbildung|bild|"
        r"afbeelding|illustration|illustrasie"
        r")\s*"
        r"[\dIVXLCivxlc]+(?:[\.:\-–—][\dIVXLCivxlc]+)*"
        r"\s*[:：\.]\s*)(.*)$",
        re.IGNORECASE | re.DOTALL,
    )
    _TRANSLATED_LABEL_SPLIT_RE = re.compile(
        r"^(\s*[^\n]{1,64}?(?:[:：\.])\s*)(.*)$",
        re.DOTALL,
    )

    def _run_is_bold_for_label(run) -> bool:
        try:
            if run.bold is True:
                return True
        except Exception:
            pass
        try:
            if getattr(getattr(run, "font", None), "bold", None) is True:
                return True
        except Exception:
            pass
        try:
            from docx.oxml.ns import qn as _qn
            rpr = run._element.find(_qn("w:rPr"))
            if rpr is not None:
                b = rpr.find(_qn("w:b"))
                if b is not None:
                    bv = str(b.get(_qn("w:val"), "") or "").strip().lower()
                    if bv not in ("0", "false", "off"):
                        return True
        except Exception:
            pass
        return False

    def _force_run_bold(run) -> None:
        """Ensure explicit bold on a run (style inheritance alone is not enough after run split)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn

        if run is None:
            return
        try:
            rpr = run._element.find(_qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run._element.insert(0, rpr)
            for tag in ("w:b", "w:bCs"):
                el = rpr.find(_qn(tag))
                if el is not None:
                    rpr.remove(el)
                rpr.append(OxmlElement(tag))
            try:
                run.bold = True
            except Exception:
                pass
        except Exception:
            pass

    def _run_is_superscript_for_label(run) -> bool:
        try:
            if bool(getattr(getattr(run, "font", None), "superscript", False)):
                return True
        except Exception:
            pass
        try:
            from docx.oxml.ns import qn as _qn
            rpr = run._element.find(_qn("w:rPr"))
            if rpr is not None:
                va = rpr.find(_qn("w:vertAlign"))
                if va is not None:
                    v = str(va.get(_qn("w:val"), "") or "").strip().lower()
                    if v in ("superscript", "subscript"):
                        return True
        except Exception:
            pass
        return False

    def _run_is_italic_for_label(run) -> bool:
        try:
            if run.italic is True:
                return True
        except Exception:
            pass
        try:
            if getattr(getattr(run, "font", None), "italic", None) is True:
                return True
        except Exception:
            pass
        try:
            from docx.oxml.ns import qn as _qn
            rpr = run._element.find(_qn("w:rPr"))
            if rpr is not None:
                i_el = rpr.find(_qn("w:i"))
                if i_el is not None:
                    iv = str(i_el.get(_qn("w:val"), "") or "").strip().lower()
                    if iv not in ("0", "false", "off"):
                        return True
        except Exception:
            pass
        return False

    def _pick_body_run_index(runs, content_indices, bold_idx):
        best_idx = None
        best_score = -1
        for i in content_indices:
            if i == bold_idx:
                continue
            if _run_is_bold_for_label(runs[i]):
                continue
            t = (runs[i].text or "").strip()
            if not t or re.fullmatch(r"[\.\s,_]+", t):
                continue
            score = len(t) + (10000 if _run_is_italic_for_label(runs[i]) else 0)
            if score > best_score:
                best_score = score
                best_idx = i
        if best_idx is not None:
            return best_idx
        for i in content_indices:
            if i != bold_idx and not _run_is_bold_for_label(runs[i]):
                return i
        return None

    def _paragraph_has_mixed_bold_label_prefix(paragraph) -> bool:
        prefix = ""
        saw_after = False
        for run in paragraph.runs:
            rt = run.text or ""
            if not rt.strip() or _run_is_superscript_for_label(run):
                continue
            if _run_is_bold_for_label(run) and not saw_after:
                prefix += rt
                continue
            if prefix.strip():
                saw_after = True
            break
        return bool(
            prefix.strip()
            and saw_after
            and len(prefix.strip()) <= 64
            and re.search(r"[:：\.]\s*$", prefix.strip())
        )

    _FORM_PROFILE_LABEL_RE = re.compile(
        r"^\s*(?:full\s*name|họ\s*và\s*tên|ho\s*(?:va|ten)|mssv|student\s*id|"
        r"email|e-mail|class|lớp|lop)\s*[:：]",
        re.IGNORECASE,
    )

    def _paragraph_has_list_formatting(paragraph) -> bool:
        try:
            from docx.oxml.ns import qn as _qn
            p_pr = paragraph._element.find(_qn("w:pPr"))
            if p_pr is not None and p_pr.find(_qn("w:numPr")) is not None:
                return True
        except Exception:
            pass
        return False

    def _extract_mixed_bold_prefix(paragraph) -> str:
        prefix = ""
        saw_after = False
        for run in paragraph.runs:
            rt = run.text or ""
            if not rt.strip() or _run_is_superscript_for_label(run):
                continue
            if _run_is_bold_for_label(run) and not saw_after:
                prefix += rt
                continue
            if prefix.strip():
                saw_after = True
            break
        return prefix.strip()

    def _fix_duplicate_list_markers_in_text(text: str) -> str:
        out = text or ""
        if not out.strip():
            return out
        out = re.sub(r"^(\s*)([+\-•*])(\s+\2)+", r"\1\2 ", out)
        return out

    def _is_form_profile_or_list_paragraph(paragraph, text: str) -> bool:
        s = (text or "").strip()
        if _FORM_PROFILE_LABEL_RE.match(s):
            return True
        if re.search(r"(\.{3,}|_{3,}|-{3,}|…+)", s):
            return True
        if re.match(r"^[\s\u00A0]*[+\-•*]\s", s):
            return True
        if _paragraph_has_list_formatting(paragraph):
            return True
        return False

    def _strip_duplicate_leading_list_marker(paragraph, text: str) -> str:
        out = _fix_duplicate_list_markers_in_text(text or "")
        if not out.strip():
            return out
        runs = list(paragraph.runs)
        content_indices = [
            i for i, r in enumerate(runs) if not _is_structural_text(r.text or "")
        ]
        if content_indices:
            first_content_idx = min(content_indices)
            if first_content_idx > 0:
                prefix_text = "".join((runs[i].text or "") for i in range(first_content_idx))
                if prefix_text and re.search(r"[+\-•*\uF000-\uF8FF]", prefix_text):
                    # Only remove a duplicate marker already present in structural prefix runs.
                    out = re.sub(
                        r"^[\s\u00A0]*[+\-•*\uF000-\uF8FF][\s\t\u00A0]*",
                        "",
                        out,
                        count=1,
                    )
        return out

    def _strip_duplicate_leading_emoji_runs(paragraph, text: str) -> str:
        """If paragraph starts with emoji/icon runs, drop the same prefix from translated text."""
        out = text or ""
        if not out.strip():
            return out
        runs = list(paragraph.runs)
        prefix = ""
        for run in runs:
            rt = run.text or ""
            if _run_has_drawing_xml(run):
                break
            if _is_decorative_leading_run_text(rt):
                prefix += rt
                continue
            break
        pt = prefix.strip()
        if not pt:
            return out
        stripped = out.lstrip()
        if stripped.startswith(pt):
            return stripped[len(pt) :].lstrip() if len(runs) > 1 else out
        pt_compact = pt.strip()
        if pt_compact and stripped.startswith(pt_compact):
            return stripped[len(pt_compact) :].lstrip() if len(runs) > 1 else out
        return out

    def _paragraph_translation_source_text(paragraph, paragraph_text=None):
        """Split leading icon/bullet runs from translatable body (e.g. title after 🌐)."""
        merged = (
            paragraph_text
            if paragraph_text is not None
            else merged_paragraph_plain(paragraph)
        )
        merged = (merged or "").strip()
        if not merged:
            return merged, ""
        try:
            runs = list(paragraph.runs)
        except Exception:
            return merged, ""
        prefix = ""
        for run in runs:
            rt = run.text or ""
            if _run_has_drawing_xml(run):
                break
            if _is_decorative_leading_run_text(rt):
                prefix += rt
                continue
            break
        if not prefix:
            return merged, ""
        if merged.startswith(prefix):
            core = merged[len(prefix) :].strip()
            if core:
                return core, prefix
        pt = prefix.strip()
        if pt and merged.lstrip().startswith(pt):
            core = merged.lstrip()[len(pt) :].strip()
            if core:
                return core, prefix
        return merged, ""

    def _is_caption_style_paragraph(paragraph) -> bool:
        try:
            sn = str(getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
            if not sn:
                return False
            if "caption" in sn:
                return True
            if "chú thích" in sn or "chu thich" in sn or "chuthich" in sn:
                return True
        except Exception:
            pass
        return False

    def _split_labeled_body(text: str):
        """Split section/caption label from body. Caption patterns (Fig. 1.) take priority."""
        raw = text or ""
        stripped = raw.strip()
        if _CAPTION_LABEL_HEAD_RE.match(stripped):
            m_cap = _CAPTION_LABEL_SPLIT_RE.match(raw)
            if m_cap:
                return _finalize_label_body_pair(m_cap.group(1), m_cap.group(2) or "")
        if _ACADEMIC_LABEL_HEAD_RE.match(stripped):
            m_acad = _TRANSLATED_LABEL_SPLIT_RE.match(raw)
            if m_acad:
                return _finalize_label_body_pair(m_acad.group(1), m_acad.group(2) or "")
        m_cap = _CAPTION_LABEL_SPLIT_RE.match(raw)
        if m_cap:
            return _finalize_label_body_pair(m_cap.group(1), m_cap.group(2) or "")
        m_gen = _TRANSLATED_LABEL_SPLIT_RE.match(raw)
        if m_gen:
            return _finalize_label_body_pair(m_gen.group(1), m_gen.group(2) or "")
        return None, None

    def _split_academic_label_body(text: str):
        return _split_labeled_body(text)

    def _normalize_translated_label(src_label: str, translated_label: str) -> str:
        out = _normalize_label_text((translated_label or "").strip())
        if not out:
            return _normalize_label_text((src_label or "").strip())
        src = (src_label or "").rstrip()
        if src.endswith("."):
            if not out.endswith("."):
                out = out.rstrip() + "."
        elif src.endswith(":") or src.endswith("："):
            if not out.endswith(":") and not out.endswith("："):
                out = out.rstrip() + ":"
        out = _normalize_label_text(out)
        if (src_label or "").endswith(" ") and not out.endswith(" "):
            out += " "
        return out

    def _merge_label_body_text(label: str, body: str) -> str:
        lbl, bdy = _finalize_label_body_pair(label, body)
        if not lbl.strip():
            return bdy
        if not bdy:
            return lbl.rstrip()
        return f"{lbl.rstrip()} {bdy}"

    def _translate_academic_label_paragraph(paragraph, paragraph_text, translate_fn) -> None:
        """Translate labeled paragraphs (Abstract/Keywords/Figure/Table captions) without losing bold labels."""
        src_label, src_body = _split_labeled_body(paragraph_text)
        if not src_label:
            translated = translate_fn(paragraph_text)
            if _paragraph_should_use_bold_label_layout(paragraph, paragraph_text):
                if not _apply_translation_with_bold_label(paragraph, translated):
                    _apply_translation_to_runs(paragraph, translated)
            else:
                _apply_translation_to_runs(paragraph, translated)
            return

        try:
            translated_label = _normalize_translated_label(
                src_label,
                _cleanup_translated_text(translate_fn(src_label.strip())),
            )
        except ProviderRateLimitError:
            raise
        except Exception:
            translated_label = src_label.strip()

        if (src_body or "").strip():
            try:
                translated_body = translate_fn(src_body)
            except ProviderRateLimitError:
                raise
            except Exception:
                translated_body = src_body
            _, translated_body = _finalize_label_body_pair(translated_label, translated_body or "")
            combined = _merge_label_body_text(translated_label, translated_body)
        else:
            combined = translated_label

        combined = _fix_duplicate_label_punctuation(combined)
        if not _apply_translation_with_bold_label(paragraph, combined):
            _apply_translation_to_runs(paragraph, combined)

    def _paragraph_should_use_bold_label_layout(paragraph, text: str) -> bool:
        s = (text or "").strip()
        if _is_form_profile_or_list_paragraph(paragraph, s):
            return False
        if _ACADEMIC_LABEL_HEAD_RE.match(s):
            return True
        if _CAPTION_LABEL_HEAD_RE.match(s):
            return True
        if _is_caption_style_paragraph(paragraph) and _CAPTION_LABEL_HEAD_RE.match(s):
            return True
        if _paragraph_has_mixed_bold_label_prefix(paragraph):
            prefix = _extract_mixed_bold_prefix(paragraph)
            if prefix and (
                _ACADEMIC_LABEL_HEAD_RE.match(prefix)
                or _CAPTION_LABEL_HEAD_RE.match(prefix)
            ):
                return True
        return False

    def _collect_label_body_rpr(paragraph):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy

        bold_rpr = None
        body_rpr = None
        body_should_be_non_bold = False
        src_runs = [r for r in paragraph.runs if (r.text or "").strip()]
        non_sup = [r for r in src_runs if not _run_is_superscript_for_label(r)]
        content = [r for r in non_sup if not re.match(r"^\d+(?:[\.)]\d+)*[\.)]?$", (r.text or "").strip())]

        for r in src_runs:
            if _run_is_bold_for_label(r) and not _run_is_superscript_for_label(r):
                found = r._element.find(_qn("w:rPr"))
                if found is not None:
                    bold_rpr = found
                    break

        chosen = None
        for r in content:
            if not _run_is_bold_for_label(r) and _run_is_italic_for_label(r):
                chosen = r
                break
        if chosen is None:
            for r in content:
                if not _run_is_bold_for_label(r):
                    chosen = r
                    break
        if chosen is None and content:
            chosen = content[0]
        if chosen is not None:
            body_rpr = chosen._element.find(_qn("w:rPr"))

        body_should_be_non_bold = any(not _run_is_bold_for_label(r) for r in content)

        if bold_rpr is None and body_rpr is not None:
            bold_rpr = _copy.deepcopy(body_rpr)
            try:
                for tag in ("w:b", "w:bCs"):
                    el = bold_rpr.find(_qn(tag))
                    if el is not None:
                        bold_rpr.remove(el)
                    bold_rpr.append(OxmlElement(tag))
            except Exception:
                pass
        elif bold_rpr is None:
            bold_rpr = OxmlElement("w:rPr")
            bold_rpr.append(OxmlElement("w:b"))
            bold_rpr.append(OxmlElement("w:bCs"))

        return bold_rpr, body_rpr, body_should_be_non_bold

    def _apply_rpr_to_run(run, src_rpr, *, force_non_bold: bool = False) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy

        if src_rpr is None:
            return
        try:
            new_rpr = _copy.deepcopy(src_rpr)
            sp = new_rpr.find(_qn("w:spacing"))
            if sp is not None:
                new_rpr.remove(sp)
            if force_non_bold:
                for tag in ("w:b", "w:bCs"):
                    b_el = new_rpr.find(_qn(tag))
                    if b_el is not None:
                        new_rpr.remove(b_el)
                    off_el = OxmlElement(tag)
                    off_el.set(_qn("w:val"), "0")
                    new_rpr.append(off_el)
            old_rpr = run._element.find(_qn("w:rPr"))
            if old_rpr is not None:
                run._element.remove(old_rpr)
            run._element.insert(0, new_rpr)
        except Exception:
            pass

    def _apply_translation_with_bold_label(paragraph, translated_text) -> bool:
        """Keep Abstract/Keywords label bold while updating translated body text."""
        txt = _fix_duplicate_label_punctuation(translated_text or "")
        if not txt.strip():
            return False
        if not _paragraph_should_use_bold_label_layout(paragraph, txt):
            return False

        m_label = None
        if _CAPTION_LABEL_HEAD_RE.match(txt.strip()) or _is_caption_style_paragraph(paragraph):
            m_label = _CAPTION_LABEL_SPLIT_RE.match(txt)
        if not m_label and _ACADEMIC_LABEL_HEAD_RE.match(txt.strip()):
            m_label = _TRANSLATED_LABEL_SPLIT_RE.match(txt)
        if not m_label:
            m_label = _CAPTION_LABEL_SPLIT_RE.match(txt)
        if not m_label:
            m_label = _TRANSLATED_LABEL_SPLIT_RE.match(txt)
        if not m_label:
            return False

        label_text = m_label.group(1)
        body_text = m_label.group(2) or ""
        label_text, body_text = _finalize_label_body_pair(label_text, body_text)
        label_only = not (body_text or "").strip()

        # Trim label for bold run (no trailing space); body carries leading space if needed.
        bold_text = label_text.rstrip()
        if body_text and not body_text.startswith((" ", "\t")) and bold_text.endswith((".", ":", "：", "!", "?")):
            body_text = " " + body_text.lstrip()

        runs = list(paragraph.runs)
        content_indices = [
            i for i, run in enumerate(runs) if not _is_structural_text(run.text or "")
        ]
        if not content_indices:
            return False

        bold_rpr, body_rpr, force_non_bold = _collect_label_body_rpr(paragraph)

        bold_idx = None
        for i in content_indices:
            if _run_is_bold_for_label(runs[i]):
                bold_idx = i
                break
        if bold_idx is None:
            bold_idx = content_indices[0]

        body_idx = _pick_body_run_index(runs, content_indices, bold_idx)

        bold_run = runs[bold_idx]
        bold_text = _normalize_label_text(bold_text)

        _apply_rpr_to_run(bold_run, bold_rpr, force_non_bold=False)
        _force_run_bold(bold_run)
        _ensure_text_font_compat(bold_run, bold_text)
        bold_run.text = bold_text

        if label_only:
            for i in content_indices:
                if runs[i] is bold_run:
                    continue
                runs[i].text = ""
            return True

        if body_idx is None:
            body_run = paragraph.add_run("")
        else:
            body_run = runs[body_idx]

        _apply_rpr_to_run(body_run, body_rpr, force_non_bold=force_non_bold)
        _ensure_text_font_compat(body_run, body_text)
        body_run.text = body_text

        for i in content_indices:
            run = runs[i]
            if run is bold_run or run is body_run:
                continue
            run.text = ""
        return True

    def _apply_translation_to_runs(paragraph, translated_text):
        translated_value = _fix_duplicate_label_punctuation(translated_text or "")
        if _apply_translation_with_bold_label(paragraph, translated_value):
            return

        runs = list(paragraph.runs)
        original_texts = [(r.text or "") for r in runs]

        if not runs:
            if "\n" in translated_value.replace("\r\n", "\n").replace("\r", "\n"):
                starter = paragraph.add_run("")
                _write_text_with_line_breaks(paragraph, starter, translated_value)
            else:
                paragraph.add_run(translated_value)
            return

        content_indices = []
        for i, run in enumerate(runs):
            if _run_has_drawing_xml(run):
                continue
            rt = run.text or ""
            if _is_structural_text(rt) or _is_decorative_leading_run_text(rt):
                continue
            content_indices.append(i)

        if not content_indices:
            return

        # When a paragraph starts with structural marker runs (e.g. Symbol bullet + tab),
        # translated text may still come back with the same marker at its start.
        # If we keep both, output becomes duplicated markers ("+ + Text" or tofu + square).
        translated_value = _strip_duplicate_leading_list_marker(paragraph, translated_value)
        translated_value = _strip_duplicate_leading_emoji_runs(paragraph, translated_value)

        structural_indices = [i for i in range(len(runs)) if i not in content_indices]
        marker_structural_found = False
        if bool(from_pdf) and structural_indices:
            for i in structural_indices:
                st = runs[i].text or ""
                if re.match(r"^\s*[+\uF000-\uF8FF](?:\s*[+\uF000-\uF8FF])*\s*$", st):
                    marker_structural_found = True
                    break

        candidate_runs = [runs[i] for i in content_indices]
        target_run = _pick_text_target_run(candidate_runs)
        if target_run is None:
            return

        if "\n" in translated_value.replace("\r\n", "\n").replace("\r", "\n"):
            for idx, run in enumerate(runs):
                if run is target_run:
                    continue
                if run in candidate_runs:
                    run.text = ""
                    continue
                # For multiline from-PDF paragraphs with symbol-marker scaffolding,
                # remove old structural runs (marker/tab/newline artifacts) and let
                # translated_value drive the final line structure.
                if marker_structural_found and idx in structural_indices:
                    run.text = ""
            _write_text_with_line_breaks(paragraph, target_run, translated_value)
            return

        if len(content_indices) > 1:
            write_translated_to_run_indices(
                runs,
                content_indices,
                original_texts,
                translated_value,
                ensure_font_compat_fn=_ensure_text_font_compat,
            )
            for i in structural_indices:
                if marker_structural_found:
                    runs[i].text = ""
            return

        _ensure_text_font_compat(target_run, translated_value)
        try:
            tidx = runs.index(target_run)
            orig = original_texts[tidx] if tidx < len(original_texts) else (target_run.text or "")
        except Exception:
            orig = target_run.text or ""
        if orig and orig != orig.strip():
            lead = orig[: len(orig) - len(orig.lstrip())]
            tail = orig[len(orig.rstrip()) :]
            core = translated_value.strip() if (translated_value or "").strip() else translated_value
            translated_value = f"{lead}{core}{tail}"
        target_run.text = translated_value

        for run in candidate_runs:
            if run is not target_run:
                run.text = ""

    def _get_run_format_key(run):
        return _run_format_signature_semantic(run)

    def _paragraph_should_translate_as_whole(text):
        raw = (text or "").strip()
        if not raw or _is_structural_text(raw):
            return False
        if leader_re.search(raw):
            return False
        words = _word_count(raw)
        if words < 2 or len(raw) < 6:
            return False
        # Vietnamese prose must stay in sentence context (đất nước, thắng cảnh, ...).
        if _has_vietnamese_diacritics(raw) and words >= 2 and len(raw) >= 8:
            return True
        if re.search(r"[.!?;:]", raw):
            return True
        if words >= 3 and len(raw) >= 12:
            return True
        return words >= 5 or len(raw) >= 30

    def _group_runs_by_format(runs):
        groups = []
        for i, run in enumerate(runs):
            if _run_has_drawing_xml(run):
                continue
            text = run.text or ""
            fmt = _get_run_format_key(run)

            if text.strip() and _is_structural_text(text):
                groups.append((b'__structural__' + fmt, [i]))
                continue

            if not text.strip():
                continue

            if groups and groups[-1][0] == fmt:
                groups[-1][1].append(i)
            else:
                groups.append((fmt, [i]))
        return groups

    def _write_group_text(runs, indices, original_texts, new_text):
        write_translated_to_run_indices(
            runs,
            indices,
            original_texts,
            new_text,
            ensure_font_compat_fn=_ensure_text_font_compat,
        )

    def _split_translated_by_source_weights(translated_text, source_chunks):
        return distribute_text_by_source_weights(translated_text, source_chunks)

    def _try_marker_translate_content_runs(runs, content_indices, original_texts, translate_fn):
        if len(content_indices) < 2:
            return False
        segments = [original_texts[i] for i in content_indices]
        payload = build_run_marker_translation_text(segments)
        try:
            translated_payload = translate_fn(payload)
        except ProviderRateLimitError:
            raise
        except Exception:
            return False
        if re.search(r"\\uE000|\\uE001", translated_payload or ""):
            return False
        parsed = parse_run_marker_translation(translated_payload, len(segments))
        if parsed is None:
            return False
        if any(re.search(r"\\uE000\d+\\uE001", p or "") for p in parsed):
            return False
        cleaned = [strip_run_boundary_markers(p or "") for p in parsed]
        if not any((c or "").strip() for c in cleaned):
            return False
        if any(_RUN_BOUNDARY_MARKER_RE.search(p or "") for p in parsed):
            return False
        for run_idx, chunk in zip(content_indices, cleaned):
            _ensure_text_font_compat(runs[run_idx], chunk)
            runs[run_idx].text = chunk
        return True

    def _translate_runs_individually(runs, content_indices, original_texts, translate_fn):
        """Never translate tiny styled fragments in isolation — keep sentence context."""
        if not content_indices:
            return False
        combined = "".join(original_texts[i] or "" for i in content_indices)
        if not combined.strip():
            return False
        try:
            translated = translate_fn(combined)
        except ProviderRateLimitError:
            raise
        except Exception:
            return False
        if not (translated or "").strip():
            return False
        write_translated_to_run_indices(
            runs,
            content_indices,
            original_texts,
            translated,
            ensure_font_compat_fn=_ensure_text_font_compat,
        )
        return True

    def _translate_format_groups(paragraph, translate_fn):
        from docx.oxml.ns import qn as _qn
        runs = list(paragraph.runs)
        if not runs:
            return

        original_texts = [(r.text or "") for r in runs]
        paragraph_text = "".join(original_texts)
        if not paragraph_text.strip():
            return

        if _should_preserve_author_metadata(paragraph, paragraph_text):
            return

        if _paragraph_should_translate_as_whole(paragraph_text):
            source_text, prefix = _paragraph_translation_source_text(paragraph, paragraph_text)
            translated = translate_fn(source_text if prefix else paragraph_text)
            _apply_translation_to_runs(paragraph, translated)
            return

        has_structural_runs = any(
            (original_texts[i] or "").strip() and _is_structural_text(original_texts[i])
            for i in range(len(runs))
        )

        groups = _group_runs_by_format(runs)

        if _paragraph_should_use_bold_label_layout(paragraph, paragraph_text):
            _translate_academic_label_paragraph(paragraph, paragraph_text, translate_fn)
            return

        content_indices = [
            i for i, t in enumerate(original_texts)
            if (t or "").strip()
            and not _is_structural_text(t or "")
            and not _is_decorative_leading_run_text(t or "")
            and not _run_has_drawing_xml(runs[i])
        ]
        if len(content_indices) >= 2 and paragraph_has_distinct_run_formats(runs, original_texts):
            seg_texts = [original_texts[i] for i in content_indices]
            if not _is_fragmented_prose_segments(seg_texts):
                try:
                    if _try_marker_translate_content_runs(
                        runs, content_indices, original_texts, translate_fn
                    ):
                        return
                except ProviderRateLimitError:
                    raise

        if len(content_indices) >= 2 and not paragraph_has_distinct_run_formats(runs, original_texts):
            source_text, _prefix = _paragraph_translation_source_text(paragraph, paragraph_text)
            translated = translate_fn(source_text if _prefix else paragraph_text)
            _apply_translation_to_runs(paragraph, translated)
            return

        if len(groups) <= 1 and not has_structural_runs:
            if len(content_indices) >= 2 and paragraph_has_distinct_run_formats(runs, original_texts):
                try:
                    if _translate_runs_individually(
                        runs, content_indices, original_texts, translate_fn
                    ):
                        return
                except ProviderRateLimitError:
                    raise
            source_text, _prefix = _paragraph_translation_source_text(paragraph, paragraph_text)
            translated = translate_fn(source_text if _prefix else paragraph_text)
            _apply_translation_to_runs(paragraph, translated)
            return

        if len(groups) <= 1 and has_structural_runs:
            if len(content_indices) >= 2 and paragraph_has_distinct_run_formats(runs, original_texts):
                try:
                    if _translate_runs_individually(
                        runs, content_indices, original_texts, translate_fn
                    ):
                        return
                except ProviderRateLimitError:
                    raise
            source_text, _prefix = _paragraph_translation_source_text(paragraph, paragraph_text)
            translated = translate_fn(source_text if _prefix else paragraph_text)
            _apply_translation_to_runs(paragraph, translated)
            return

        # pdf2docx often creates noisy multi-run paragraphs; skip redistribution only
        # when runs share the same effective formatting.
        if bool(from_pdf) and not paragraph_has_distinct_run_formats(runs, original_texts):
            translated = translate_fn(paragraph_text)
            _apply_translation_to_runs(paragraph, translated)
            return

        # Semantic-quality path: translate whole paragraph once, redistribute by style.
        # Works even when runs differ by bold/color — avoids fragment-level drift.
        has_form_leaders = bool(re.search(r"(\.{3,}|_{3,}|-{3,}|…+|\t+)", paragraph_text or ""))
        if not has_form_leaders:
            non_struct_groups = []
            for _, indices in groups:
                group_text = "".join(original_texts[i] for i in indices)
                if not group_text.strip() or _is_structural_text(group_text):
                    continue
                non_struct_groups.append((indices, group_text))

            if non_struct_groups:
                try:
                    source_text, _prefix = _paragraph_translation_source_text(
                        paragraph, paragraph_text
                    )
                    translated_para = translate_fn(
                        source_text if _prefix else paragraph_text
                    )
                    if _prefix:
                        translated_para = _strip_duplicate_leading_emoji_runs(
                            paragraph, translated_para
                        )
                    translated_chunks = _split_translated_by_source_weights(
                        translated_para,
                        [g[1] for g in non_struct_groups],
                    )
                    prev_written = ""
                    first_chunk = True
                    for (indices, _src), t_chunk in zip(non_struct_groups, translated_chunks):
                        chunk = t_chunk or ""
                        if first_chunk:
                            chunk = _strip_duplicate_leading_list_marker(paragraph, chunk)
                            first_chunk = False
                        if prev_written:
                            chunk = _strip_duplicate_leading_punct(prev_written, chunk)
                        _write_group_text(runs, indices, original_texts, chunk)
                        prev_written = (prev_written or "") + chunk
                    return
                except ProviderRateLimitError:
                    raise
                except Exception as e:
                    print(f"Paragraph-level redistribution failed, fallback group mode: {e}")

        # Safety net: never translate multi-word prose one format-group at a time.
        if _word_count(paragraph_text) >= 3 and not has_form_leaders:
            try:
                source_text, _prefix = _paragraph_translation_source_text(
                    paragraph, paragraph_text
                )
                translated_para = translate_fn(source_text if _prefix else paragraph_text)
                _apply_translation_to_runs(paragraph, translated_para)
                return
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Whole-paragraph safety fallback failed: {e}")

        prev_written = ""
        first_chunk = True
        for fmt_key, indices in groups:
            group_text = "".join(original_texts[i] for i in indices)
            if not group_text.strip():
                continue

            if _is_structural_text(group_text):
                continue
            # Leave punctuation/spacing-only groups unchanged.
            if not re.search(r"[\w\u00C0-\u1EF9]", group_text, flags=re.UNICODE):
                continue

            try:
                translated_group = translate_fn(group_text)
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Format-group translation failed: {e}")
                if api_only:
                    raise
                translated_group = group_text
            chunk = translated_group or ""
            chunk = strip_run_boundary_markers(chunk)
            if first_chunk:
                chunk = _strip_duplicate_leading_list_marker(paragraph, chunk)
                first_chunk = False
            if prev_written:
                chunk = _strip_duplicate_leading_punct(prev_written, chunk)
            _write_group_text(runs, indices, original_texts, chunk)
            if chunk.strip():
                prev_written = (prev_written or "") + chunk

    def _insert_paragraph_after(ref_para, text, italic=True, *, clear_first_line=False):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy

        def _to_twips(v):
            if v is None:
                return None
            try:
                return int(v.twips)
            except Exception:
                try:
                    return int(v)
                except Exception:
                    return None

        new_p = OxmlElement('w:p')

        try:
            pPr_src = ref_para._element.find(_qn('w:pPr'))
            if pPr_src is not None:
                new_pPr = _copy.deepcopy(pPr_src)

                try:
                    if new_pPr.find(_qn('w:pStyle')) is None:
                        style_id = str(getattr(getattr(ref_para, 'style', None), 'style_id', '') or '').strip()
                        if style_id:
                            p_style = OxmlElement('w:pStyle')
                            p_style.set(_qn('w:val'), style_id)
                            new_pPr.insert(0, p_style)
                except Exception:
                    pass

                numPr = new_pPr.find(_qn('w:numPr'))
                if numPr is not None:
                    src_left = None
                    src_hanging = None
                    src_firstLine = None
                    try:
                        fmt = ref_para.paragraph_format
                        src_left = _to_twips(fmt.left_indent)
                        src_firstLine = _to_twips(fmt.first_line_indent)
                    except Exception:
                        pass
                    if src_left is None:
                        try:
                            orig_ind = pPr_src.find(_qn('w:ind'))
                            if orig_ind is not None:
                                l_val = orig_ind.get(_qn('w:left')) or orig_ind.get(_qn('w:start'))
                                if l_val:
                                    src_left = int(l_val)
                                h_val = orig_ind.get(_qn('w:hanging'))
                                if h_val:
                                    src_hanging = int(h_val)
                                fl_val = orig_ind.get(_qn('w:firstLine'))
                                if fl_val:
                                    src_firstLine = int(fl_val)
                        except Exception:
                            pass

                    new_pPr.remove(numPr)

                    ind = new_pPr.find(_qn('w:ind'))
                    if ind is None and src_left is not None:
                        ind = OxmlElement('w:ind')
                        ind.set(_qn('w:left'), str(src_left))
                        if src_hanging is not None:
                            ind.set(_qn('w:hanging'), str(src_hanging))
                        elif src_firstLine is not None:
                            if src_firstLine < 0:
                                ind.set(_qn('w:hanging'), str(abs(src_firstLine)))
                            elif src_firstLine > 0:
                                ind.set(_qn('w:firstLine'), str(src_firstLine))
                        new_pPr.append(ind)

                jc = new_pPr.find(_qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(_qn('w:val'), '')
                    if jc_val in ('distribute', 'thai-distribute'):
                        new_pPr.remove(jc)

                new_p.insert(0, new_pPr)
            else:
                try:
                    style_id = str(getattr(getattr(ref_para, 'style', None), 'style_id', '') or '').strip()
                    if style_id:
                        new_pPr = OxmlElement('w:pPr')
                        p_style = OxmlElement('w:pStyle')
                        p_style.set(_qn('w:val'), style_id)
                        new_pPr.append(p_style)
                        new_p.insert(0, new_pPr)
                except Exception:
                    pass
        except Exception:
            pass

        def _run_is_superscript(run):
            try:
                if bool(getattr(getattr(run, 'font', None), 'superscript', False)):
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    va = rpr.find(_qn('w:vertAlign'))
                    if va is not None:
                        v = str(va.get(_qn('w:val'), '') or '').strip().lower()
                        if v in ('superscript', 'subscript'):
                            return True
            except Exception:
                pass
            return False

        def _run_is_bold(run):
            try:
                if run.bold is True:
                    return True
            except Exception:
                pass
            try:
                if getattr(getattr(run, 'font', None), 'bold', None) is True:
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(_qn('w:b'))
                    if b is not None:
                        bv = str(b.get(_qn('w:val'), '') or '').strip().lower()
                        if bv not in ('0', 'false', 'off'):
                            return True
            except Exception:
                pass
            return False

        def _clone_rpr(src_rpr, *, force_non_bold=False):
            if src_rpr is not None:
                rpr = _copy.deepcopy(src_rpr)
            else:
                rpr = OxmlElement('w:rPr')
            try:
                sp = rpr.find(_qn('w:spacing'))
                if sp is not None:
                    rpr.remove(sp)
            except Exception:
                pass
            if force_non_bold:
                try:
                    for tag in ('w:b', 'w:bCs'):
                        b_el = rpr.find(_qn(tag))
                        if b_el is not None:
                            rpr.remove(b_el)
                        off_el = OxmlElement(tag)
                        off_el.set(_qn('w:val'), '0')
                        rpr.append(off_el)
                except Exception:
                    pass
            return rpr

        def _append_run_with_rpr(p_el, txt, src_rpr, *, force_non_bold=False):
            if txt is None or txt == '':
                return
            chunks = str(txt).replace('\r\n', '\n').replace('\r', '\n').split('\n')
            for idx, chunk in enumerate(chunks):
                if idx > 0:
                    br_el = OxmlElement('w:r')
                    br_rpr = _clone_rpr(src_rpr, force_non_bold=force_non_bold)
                    if italic:
                        try:
                            if br_rpr.find(_qn('w:i')) is None:
                                br_rpr.append(OxmlElement('w:i'))
                        except Exception:
                            pass
                    br_el.insert(0, br_rpr)
                    br_el.append(OxmlElement('w:br'))
                    p_el.append(br_el)
                if chunk == '':
                    continue
                r_el = OxmlElement('w:r')
                rpr = _clone_rpr(src_rpr, force_non_bold=force_non_bold)
                if italic:
                    try:
                        if rpr.find(_qn('w:i')) is None:
                            rpr.append(OxmlElement('w:i'))
                    except Exception:
                        pass
                r_el.insert(0, rpr)
                t_el = OxmlElement('w:t')
                t_el.set(_qn('xml:space'), 'preserve')
                t_el.text = chunk
                r_el.append(t_el)
                p_el.append(r_el)

        def _append_with_super_markers(p_el, txt, base_rpr, sup_rpr, markers, *, force_non_bold_base=False):
            if txt is None or txt == '':
                return
            marker_values = [m for m in (markers or []) if m]
            if sup_rpr is None or not marker_values:
                _append_run_with_rpr(p_el, txt, base_rpr, force_non_bold=force_non_bold_base)
                return
            marker_values = sorted(set(marker_values), key=len, reverse=True)
            pattern = r'(?<!\d)(' + '|'.join(re.escape(m) for m in marker_values) + r')(?!\d)'
            try:
                parts = re.split(pattern, txt)
            except Exception:
                _append_run_with_rpr(p_el, txt, base_rpr, force_non_bold=force_non_bold_base)
                return
            if not parts:
                _append_run_with_rpr(p_el, txt, base_rpr, force_non_bold=force_non_bold_base)
                return
            for idx, part in enumerate(parts):
                if part is None or part == '':
                    continue
                if idx % 2 == 1:
                    _append_run_with_rpr(p_el, part, sup_rpr)
                else:
                    _append_run_with_rpr(p_el, part, base_rpr, force_non_bold=force_non_bold_base)

        body_rpr = None
        bold_rpr = None
        super_rpr = None
        super_digit_markers = []
        has_mixed_bold_prefix = False
        has_leading_sup_marker = False
        body_should_be_non_bold = False

        def _is_heading_like(p):
            try:
                sn = str(getattr(getattr(p, 'style', None), 'name', '') or '').strip().lower()
                if sn.startswith('heading') or sn in ('title', 'subtitle'):
                    return True
                if 'title' in sn or 'tiêu đề' in sn or 'tieu de' in sn:
                    return True
                return False
            except Exception:
                return False

        def _is_centered_like(p):
            try:
                if getattr(p, 'alignment', None) == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    return True
            except Exception:
                pass
            try:
                ppr = p._element.find(_qn('w:pPr'))
                if ppr is not None:
                    jc = ppr.find(_qn('w:jc'))
                    if jc is not None:
                        return str(jc.get(_qn('w:val'), '') or '').strip().lower() == 'center'
            except Exception:
                pass
            return False

        def _is_marker_run_text(t):
            s = (t or '').strip()
            if not s:
                return False
            return bool(re.match(r'^\d+(?:[\.)]\d+)*[\.)]?$', s))

        def _extract_sup_markers_from_text(t):
            s = (t or '').strip()
            if not s:
                return []
            # Capture short affiliation/reference markers (1, 2, 10, ...)
            # and ignore long numbers like years/ORCID chunks.
            return re.findall(r'(?<!\d)(\d{1,2})(?=(?:[\)\]\.,;\s]|$|\[))', s)

        def _rpr_size(rpr):
            if rpr is None:
                return -1
            try:
                sz = rpr.find(_qn('w:sz'))
                if sz is not None:
                    raw = str(sz.get(_qn('w:val'), '') or '').strip()
                    if raw.isdigit():
                        return int(raw)
            except Exception:
                pass
            return -1

        def _best_rpr_from_runs(runs):
            best = None
            best_sz = -1
            for rr in runs or []:
                try:
                    if _run_is_superscript(rr):
                        continue
                    rpr = rr._element.find(_qn('w:rPr'))
                    if rpr is None:
                        continue
                    sz = _rpr_size(rpr)
                    if best is None or sz > best_sz:
                        best = rpr
                        best_sz = sz
                except Exception:
                    continue
            return best

        try:
            src_runs = [r for r in list(ref_para.runs) if (r.text or '').strip()]
            if src_runs:
                non_sup_runs = [r for r in src_runs if not _run_is_superscript(r)]
                content_runs = [r for r in non_sup_runs if not _is_marker_run_text(r.text or '')]
                para_text = join_docx_run_texts([(r.text or '') for r in src_runs])

                total_chars = sum(len((r.text or '').strip()) for r in content_runs)
                bold_chars = sum(len((r.text or '').strip()) for r in content_runs if _run_is_bold(r))
                bold_ratio = (float(bold_chars) / float(max(1, total_chars))) if total_chars > 0 else 0.0
                line_count = len([ln for ln in (para_text or '').splitlines() if ln.strip()])
                centered_title_like = _is_centered_like(ref_para) and total_chars <= 220 and line_count <= 3
                prefer_bold_body = (
                    _is_heading_like(ref_para)
                    or (centered_title_like and bold_ratio >= 0.35)
                    or (bold_chars >= 4 and bold_ratio >= 0.60)
                )

                body_should_be_non_bold = (not prefer_bold_body) and any(
                    (not _run_is_bold(r)) for r in content_runs
                )
                # Body style: prefer non-bold + non-superscript to avoid whole-line bold/superscript drift.
                chosen_run = None
                if prefer_bold_body:
                    for r in content_runs:
                        if _run_is_bold(r):
                            chosen_run = r
                            break
                    if chosen_run is None:
                        for r in non_sup_runs:
                            if not _is_marker_run_text(r.text or ''):
                                chosen_run = r
                                break
                else:
                    for r in content_runs:
                        if not _run_is_bold(r):
                            chosen_run = r
                            break
                    if chosen_run is None:
                        for r in non_sup_runs:
                            if _run_is_bold(r):
                                continue
                            chosen_run = r
                            break
                if chosen_run is None:
                    chosen_run = (content_runs[0] if content_runs else (non_sup_runs[0] if non_sup_runs else None))
                if chosen_run is None:
                    chosen_run = src_runs[0]

                body_rpr = chosen_run._element.find(_qn('w:rPr'))
                if body_rpr is None:
                    body_rpr = _best_rpr_from_runs(content_runs if content_runs else non_sup_runs)

                for r in src_runs:
                    if _run_is_bold(r) and not _run_is_superscript(r):
                        bold_rpr = r._element.find(_qn('w:rPr'))
                        if bold_rpr is not None:
                            break
                for r in src_runs:
                    if _run_is_superscript(r):
                        super_rpr = r._element.find(_qn('w:rPr'))
                        if super_rpr is not None:
                            break
                for r in src_runs:
                    if not _run_is_superscript(r):
                        continue
                    token = (r.text or '').strip()
                    if token and len(token) <= 64 and (not re.search(r'\s', token)) and any(ch.isdigit() for ch in token):
                        if token not in super_digit_markers:
                            super_digit_markers.append(token)
                    for marker in _extract_sup_markers_from_text(token):
                        if marker not in super_digit_markers:
                            super_digit_markers.append(marker)

                first_run = src_runs[0]
                first_text = (first_run.text or '').strip()
                has_leading_sup_marker = bool(
                    _run_is_superscript(first_run)
                    and re.match(r'^\d{1,2}(?:[\)\.]?)?(?:$|[\s\[,;])', first_text)
                )

                prefix = ''
                saw_after = False
                for r in src_runs:
                    rt = r.text or ''
                    if not rt.strip() or _run_is_superscript(r):
                        continue
                    if _run_is_bold(r) and not saw_after:
                        prefix += rt
                        continue
                    if prefix.strip():
                        saw_after = True
                    break
                if prefix.strip() and saw_after and len(prefix.strip()) <= 64:
                    has_mixed_bold_prefix = bool(re.search(r'[:：\.]\s*$', prefix.strip()))

            if body_rpr is None:
                body_rpr = _best_rpr_from_runs(src_runs if 'src_runs' in locals() else [])
            if not italic:
                try:
                    dom_rpr = _best_rpr_from_runs(
                        content_runs if content_runs else (non_sup_runs if 'non_sup_runs' in locals() else [])
                    )
                    if dom_rpr is not None:
                        body_rpr = dom_rpr
                except Exception:
                    pass
        except Exception:
            try:
                body_rpr = _best_rpr_from_runs(list(ref_para.runs))
            except Exception:
                body_rpr = None

        remaining = '' if text is None else str(text)

        if has_leading_sup_marker and super_rpr is not None:
            m_ref = re.match(r'^(\s*\d{1,2}(?:[\)\.]?)(?:\s+|(?=[\[,;])))(.*)$', remaining, flags=re.DOTALL)
            if m_ref:
                _append_run_with_rpr(new_p, m_ref.group(1), super_rpr)
                remaining = m_ref.group(2)

        wrote_split = False
        if has_mixed_bold_prefix and bold_rpr is not None:
            m_label = re.match(r'^(\s*[^\n]{1,64}?(?:[:：\.])\s*)(.*)$', remaining, flags=re.DOTALL)
            if m_label and (m_label.group(2) or '').strip():
                label_text, body_text = _finalize_label_body_pair(m_label.group(1), m_label.group(2))
                _append_run_with_rpr(new_p, label_text, bold_rpr)
                _append_with_super_markers(
                    new_p,
                    body_text,
                    body_rpr,
                    super_rpr,
                    super_digit_markers,
                    force_non_bold_base=body_should_be_non_bold,
                )
                wrote_split = True

        if not wrote_split:
            src_for_mirror = src_runs if 'src_runs' in locals() and src_runs else list(ref_para.runs)
            mirror_templates = [
                r for r in (src_for_mirror or [])
                if (r.text or '').strip() and not _run_is_superscript(r) and not _is_marker_run_text(r.text or '')
            ]
            mirrored = False
            if len(mirror_templates) >= 2 and paragraph_has_distinct_run_formats(mirror_templates):
                mirror_segments = [r.text or "" for r in mirror_templates]
                mirror_chunks = distribute_text_by_source_weights(remaining, mirror_segments)
                for tmpl, chunk in zip(mirror_templates, mirror_chunks):
                    tmpl_rpr = tmpl._element.find(_qn('w:rPr'))
                    _append_run_with_rpr(
                        new_p,
                        chunk,
                        tmpl_rpr,
                        force_non_bold=body_should_be_non_bold,
                    )
                mirrored = True
            if not mirrored:
                _append_with_super_markers(
                    new_p,
                    remaining,
                    body_rpr,
                    super_rpr,
                    super_digit_markers,
                    force_non_bold_base=body_should_be_non_bold,
                )
        try:
            ref_para._element.addnext(new_p)
            if clear_first_line:
                ppr = new_p.find(_qn('w:pPr'))
                if ppr is not None:
                    ind = ppr.find(_qn('w:ind'))
                    if ind is not None:
                        for key in (_qn('w:firstLine'), _qn('w:hanging')):
                            ind.attrib.pop(key, None)
            return new_p
        except Exception:
            return None

    def _append_translation_linebreak(paragraph, text, italic=True):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy

        txt = (text or '').strip()
        if not txt:
            return False

        def _run_is_superscript(run):
            try:
                if bool(getattr(getattr(run, 'font', None), 'superscript', False)):
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    va = rpr.find(_qn('w:vertAlign'))
                    if va is not None:
                        v = str(va.get(_qn('w:val'), '') or '').strip().lower()
                        if v in ('superscript', 'subscript'):
                            return True
            except Exception:
                pass
            return False

        def _run_is_bold(run):
            try:
                if run.bold is True:
                    return True
            except Exception:
                pass
            try:
                if getattr(getattr(run, 'font', None), 'bold', None) is True:
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(_qn('w:b'))
                    if b is not None:
                        bv = str(b.get(_qn('w:val'), '') or '').strip().lower()
                        if bv not in ('0', 'false', 'off'):
                            return True
            except Exception:
                pass
            return False

        def _apply_rpr_to_run(dst_run, src_rpr, *, force_non_bold=False):
            if src_rpr is None:
                return
            try:
                new_rpr = _copy.deepcopy(src_rpr)
                sp = new_rpr.find(_qn('w:spacing'))
                if sp is not None:
                    new_rpr.remove(sp)
                if force_non_bold:
                    for tag in ('w:b', 'w:bCs'):
                        b_el = new_rpr.find(_qn(tag))
                        if b_el is not None:
                            new_rpr.remove(b_el)
                        off_el = OxmlElement(tag)
                        off_el.set(_qn('w:val'), '0')
                        new_rpr.append(off_el)
                if italic and new_rpr.find(_qn('w:i')) is None:
                    new_rpr.append(OxmlElement('w:i'))
                old_rpr = dst_run._element.find(_qn('w:rPr'))
                if old_rpr is not None:
                    dst_run._element.remove(old_rpr)
                dst_run._element.insert(0, new_rpr)
            except Exception:
                pass

        def _append_styled_text_with_breaks(paragraph_obj, raw_text, src_rpr, *, force_non_bold=False):
            chunks = str(raw_text or '').replace('\r\n', '\n').replace('\r', '\n').split('\n')
            wrote_any = False
            for idx, chunk in enumerate(chunks):
                if idx > 0:
                    br = paragraph_obj.add_run('')
                    _apply_rpr_to_run(br, src_rpr, force_non_bold=force_non_bold)
                    br.add_break()
                    wrote_any = True
                if chunk == '':
                    continue
                rr = paragraph_obj.add_run(chunk)
                _apply_rpr_to_run(rr, src_rpr, force_non_bold=force_non_bold)
                wrote_any = True
            return wrote_any

        try:
            paragraph.add_run('').add_break()

            src_runs = [r for r in list(paragraph.runs) if (r.text or '').strip()]
            body_rpr = None
            bold_rpr = None
            has_mixed_bold_prefix = False
            body_should_be_non_bold = False

            def _is_heading_like(p):
                try:
                    sn = str(getattr(getattr(p, 'style', None), 'name', '') or '').strip().lower()
                    if sn.startswith('heading') or sn in ('title', 'subtitle'):
                        return True
                    if 'title' in sn or 'tiêu đề' in sn or 'tieu de' in sn:
                        return True
                    return False
                except Exception:
                    return False

            def _is_centered_like(p):
                try:
                    if getattr(p, 'alignment', None) == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        return True
                except Exception:
                    pass
                try:
                    ppr = p._element.find(_qn('w:pPr'))
                    if ppr is not None:
                        jc = ppr.find(_qn('w:jc'))
                        if jc is not None:
                            return str(jc.get(_qn('w:val'), '') or '').strip().lower() == 'center'
                except Exception:
                    pass
                return False

            def _is_marker_run_text(t):
                s = (t or '').strip()
                if not s:
                    return False
                return bool(re.match(r'^\d+(?:[\.)]\d+)*[\.)]?$', s))

            def _rpr_size(rpr):
                if rpr is None:
                    return -1
                try:
                    sz = rpr.find(_qn('w:sz'))
                    if sz is not None:
                        raw = str(sz.get(_qn('w:val'), '') or '').strip()
                        if raw.isdigit():
                            return int(raw)
                except Exception:
                    pass
                return -1

            def _best_rpr_from_runs(runs):
                best = None
                best_sz = -1
                for rr in runs or []:
                    try:
                        if _run_is_superscript(rr):
                            continue
                        rpr = rr._element.find(_qn('w:rPr'))
                        if rpr is None:
                            continue
                        sz = _rpr_size(rpr)
                        if best is None or sz > best_sz:
                            best = rpr
                            best_sz = sz
                    except Exception:
                        continue
                return best

            try:
                non_sup_runs = [r for r in src_runs if not _run_is_superscript(r)]
                content_runs = [r for r in non_sup_runs if not _is_marker_run_text(r.text or '')]
                para_text = join_docx_run_texts([(r.text or '') for r in src_runs])

                total_chars = sum(len((r.text or '').strip()) for r in content_runs)
                bold_chars = sum(len((r.text or '').strip()) for r in content_runs if _run_is_bold(r))
                bold_ratio = (float(bold_chars) / float(max(1, total_chars))) if total_chars > 0 else 0.0
                line_count = len([ln for ln in (para_text or '').splitlines() if ln.strip()])
                centered_title_like = _is_centered_like(paragraph) and total_chars <= 220 and line_count <= 3
                prefer_bold_body = (
                    _is_heading_like(paragraph)
                    or (centered_title_like and bold_ratio >= 0.35)
                    or (bold_chars >= 4 and bold_ratio >= 0.60)
                )

                body_should_be_non_bold = (not prefer_bold_body) and any(
                    (not _run_is_bold(r)) for r in content_runs
                )

                chosen = None
                if prefer_bold_body:
                    for r in content_runs:
                        if _run_is_bold(r):
                            chosen = r
                            break
                    if chosen is None:
                        for r in non_sup_runs:
                            if not _is_marker_run_text(r.text or ''):
                                chosen = r
                                break
                else:
                    for r in content_runs:
                        if not _run_is_bold(r):
                            chosen = r
                            break
                    if chosen is None:
                        for r in non_sup_runs:
                            if _run_is_bold(r):
                                continue
                            chosen = r
                            break
                if chosen is None:
                    chosen = (content_runs[0] if content_runs else (non_sup_runs[0] if non_sup_runs else None))
                if chosen is None and src_runs:
                    chosen = src_runs[0]
                if chosen is not None:
                    body_rpr = chosen._element.find(_qn('w:rPr'))
                if body_rpr is None:
                    body_rpr = _best_rpr_from_runs(content_runs if content_runs else non_sup_runs)
                if body_rpr is None:
                    body_rpr = _best_rpr_from_runs(src_runs)

                for r in src_runs:
                    if _run_is_bold(r) and not _run_is_superscript(r):
                        bold_rpr = r._element.find(_qn('w:rPr'))
                        if bold_rpr is not None:
                            break

                prefix = ''
                saw_after = False
                for r in src_runs:
                    rt = r.text or ''
                    if not rt.strip() or _run_is_superscript(r):
                        continue
                    if _run_is_bold(r) and not saw_after:
                        prefix += rt
                        continue
                    if prefix.strip():
                        saw_after = True
                    break
                if prefix.strip() and saw_after and len(prefix.strip()) <= 64:
                    has_mixed_bold_prefix = bool(re.search(r'[:：\.]\s*$', prefix.strip()))
            except Exception:
                pass

            m_label = None
            if has_mixed_bold_prefix and bold_rpr is not None:
                m_label = re.match(r'^(\s*[^\n]{1,64}?(?:[:：\.])\s*)(.*)$', txt, flags=re.DOTALL)

            if m_label and (m_label.group(2) or '').strip():
                label_text, body_text = _finalize_label_body_pair(m_label.group(1), m_label.group(2))
                _append_styled_text_with_breaks(paragraph, label_text, bold_rpr)
                _append_styled_text_with_breaks(
                    paragraph,
                    body_text,
                    body_rpr,
                    force_non_bold=body_should_be_non_bold,
                )
            else:
                mirror_templates = [
                    r for r in (src_runs or [])
                    if (r.text or '').strip() and not _run_is_superscript(r)
                ]
                if len(mirror_templates) >= 2 and paragraph_has_distinct_run_formats(mirror_templates):
                    mirror_chunks = distribute_text_by_source_weights(
                        txt,
                        [r.text or "" for r in mirror_templates],
                    )
                    for tmpl, chunk in zip(mirror_templates, mirror_chunks):
                        tmpl_rpr = tmpl._element.find(_qn('w:rPr'))
                        _append_styled_text_with_breaks(
                            paragraph,
                            chunk,
                            tmpl_rpr,
                            force_non_bold=body_should_be_non_bold,
                        )
                else:
                    _append_styled_text_with_breaks(
                        paragraph,
                        txt,
                        body_rpr,
                        force_non_bold=body_should_be_non_bold,
                    )
            return True
        except Exception:
            return False

    leader_re = re.compile(r"(\.{3,}|_{3,}|-{3,}|…+|\t+)")
    _pdf_list_line_re = re.compile(r"^\s*(?:[-*]|\(?\d{1,3}[.)]|[A-Za-z][.)])\s+\S")
    _pdf_form_hint_re = re.compile(
        r"\b(?:tel|phone|fax|email|e-mail|address|website|hotline|mst|tax\s*code"
        r"|địa\s*chỉ|số\s*điện\s*thoại|mã\s*số\s*thuế|ngày|tháng|năm|cmnd|cccd)\b",
        re.IGNORECASE | re.UNICODE,
    )
    _merge_soft_breaks_mode = str(os.getenv("PDF_DOCX_MERGE_SOFT_BREAKS", "0") or "").strip().lower()

    # Regex for detecting merged single-line paragraphs from pdf2docx (no newlines)
    # Phase-1: "+ text" bullet embedded mid-line
    # Handles BOTH with-space ("word + bullet") and no-space ("word+ bullet") before +
    _pdf_compact_bullet_re = re.compile(
        r'(?<=[a-zA-Z\u00C0-\u1EF9\d.)])[^\S\n]*\+[^\S\n]+(?=\S)', re.UNICODE
    )
    # Pattern for label tokens: "WORD:" or "Word word:" (capital-led, followed by colon)
    # ALL-CAPS branch limited to {0,7} (max 8-char labels like SUBJECT/STUDENT/CHAPTER).
    # Title-case branch: first-letter upper, rest lowercase (Full, Class, Sinh, Lop...).
    # Both prevent false matches on class codes like "DH22TIN07" (9 chars) or
    # concatenated values like "PhucMSSV:" (mixed-case middle).
    _pdf_label_count_re = re.compile(
        r'[A-Z](?:[A-Z\d]{0,7}|[a-z\u00C0-\u1EF9\d]{0,15})'
        r'(?:[^\S\n]+[a-z\u00C0-\u1EF9\d][A-Za-z\u00C0-\u1EF9\d]{0,9}){0,1}:',
        re.UNICODE,
    )

    def _insert_newlines_before_labels(line: str) -> str:
        """Insert \\n before each label token after the first one.
        Works for both space-separated ('Phuc MSSV:') and
        no-space ('PhucMSSV:') boundaries."""
        matches = list(_pdf_label_count_re.finditer(line))
        if len(matches) < 2:
            return line
        parts: list = []
        prev_end = 0
        for i, m in enumerate(matches):
            if i == 0:
                parts.append(line[0:m.end()])
                prev_end = m.end()
            else:
                # Value between prev label end and this label start; strip trailing spaces only
                gap = line[prev_end:m.start()].rstrip(' \t')
                if gap:
                    parts.append(gap)
                parts.append('\n')
                parts.append(line[m.start():m.end()])
                prev_end = m.end()
        if prev_end < len(line):
            parts.append(line[prev_end:])
        return ''.join(parts)

    def _split_compact_pdf_paragraph(text: str) -> str:
        """Inject \\n breaks into compact from-PDF text where labels/bullets
        were glued together by pdf2docx.

        Detects two patterns:
        1. Inline "+" bullets: "...text+ BulletItem..." → split before each bullet.
        2. Chained label-value pairs: "Label1: val1Label2: val2" → split before each label.
        Works regardless of whether there are spaces at the boundaries.
        """
        if not text:
            return text
        norm = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        if len(norm.strip()) < 10:
            return text

        changed = False
        out_lines = []

        for raw_line in norm.split("\n"):
            if not (raw_line or "").strip():
                out_lines.append(raw_line)
                continue

            # Phase 1: split before inline "+ bullet" (no-space or with-space before +)
            line_result = _pdf_compact_bullet_re.sub("\n+ ", raw_line)

            # Phase 2: split chained label-value pairs on each logical line
            line_parts = []
            for part in line_result.split("\n"):
                if not (part or "").strip():
                    line_parts.append(part)
                    continue
                n_labels = len(_pdf_label_count_re.findall(part))
                if n_labels >= 2:
                    line_parts.append(_insert_newlines_before_labels(part))
                else:
                    line_parts.append(part)

            line_result = "\n".join(line_parts)
            if line_result != raw_line:
                changed = True
            out_lines.append(line_result)

        result = "\n".join(out_lines)
        return result if changed else text

    def _from_pdf_merge_soft_breaks_enabled() -> bool:
        if not bool(from_pdf):
            return False
        return _merge_soft_breaks_mode in ("1", "true", "yes", "on", "force", "always")

    def _is_structural_text(text):
        t = (text or "").strip()
        if not t:
            return True
        # Orphan method-call punctuation must stay with content, not as separate structural runs.
        if re.fullmatch(r"[\(\)\.]+", t):
            return False
        return not re.search(r'[\w\u00C0-\u1EF9]', t, flags=re.UNICODE)

    def _is_decorative_leading_run_text(text):
        """Emoji, flags, bullets, icons — keep in place, do not send to translator."""
        raw = text or ""
        stripped = raw.strip()
        if not stripped:
            return True
        if _is_structural_text(raw):
            return True
        if len(stripped) <= 4 and not re.search(r"[\w\u00C0-\u1EF9]", stripped, flags=re.UNICODE):
            return True
        if _EMOJI_ONLY_RUN_RE.fullmatch(stripped):
            return True
        return False

    def _from_pdf_preserve_line_breaks(text: str) -> bool:
        if not bool(from_pdf):
            return False
        raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        if "\n" not in raw:
            return False

        lines = [ln.strip() for ln in raw.split("\n") if ln.strip()]
        if len(lines) < 2:
            return False
        if _paragraph_is_multi_reference_list(raw):
            return True
        if any(_pdf_list_line_re.match(ln) for ln in lines):
            return True
        if leader_re.search(raw):
            return True

        label_like = 0
        short_lines = 0
        upper_short = 0
        punct_dense = 0
        digit_heavy = 0
        form_hint = 0
        for ln in lines:
            words = re.findall(r"\w+", ln, flags=re.UNICODE)
            if len(words) <= 7:
                short_lines += 1
            if len(ln) <= 32 and re.search(r"[A-ZÀ-Ỹ]", ln) and ln.upper() == ln:
                upper_short += 1
            if re.search(r"^[^\n:：]{1,42}\s*[:：]\s*\S*", ln):
                label_like += 1
            elif re.search(r"^[^\n]{1,42}\s*(?:\.{2,}|_{2,}|-{2,})\s*\S*", ln):
                label_like += 1
            if re.search(r"[@:/\\]", ln):
                punct_dense += 1
            if len(re.findall(r"\d", ln)) >= 4:
                digit_heavy += 1
            if _pdf_form_hint_re.search(ln):
                form_hint += 1

        if form_hint >= 1:
            return True
        if digit_heavy >= 2:
            return True
        if label_like >= 2:
            return True
        if label_like >= 1 and short_lines >= 2:
            return True
        if short_lines >= max(2, len(lines) - 1) and (upper_short >= 1 or punct_dense >= 1 or label_like >= 1):
            return True
        if short_lines >= 2 and punct_dense >= 2:
            return True
        return False

    def _from_pdf_should_merge_soft_breaks(lines: list[str]) -> bool:
        if not _from_pdf_merge_soft_breaks_enabled():
            return False
        if len(lines) < 2:
            return False

        word_counts = [len(re.findall(r"\w+", ln, flags=re.UNICODE)) for ln in lines]
        lengths = [len(ln) for ln in lines]
        avg_words = (sum(word_counts) / len(word_counts)) if word_counts else 0.0
        avg_len = (sum(lengths) / len(lengths)) if lengths else 0.0

        short_lines = sum(1 for w in word_counts if w <= 7)
        long_lines = sum(1 for w in word_counts if w >= 10)
        punct_end_lines = sum(1 for ln in lines if re.search(r"[.!?:;…]\s*$", ln))
        digit_heavy = sum(1 for ln in lines if len(re.findall(r"\d", ln)) >= 4)
        form_hint = sum(1 for ln in lines if _pdf_form_hint_re.search(ln))

        if form_hint >= 1:
            return False
        if digit_heavy >= 2:
            return False
        if punct_end_lines >= max(2, len(lines) // 2):
            return False
        if short_lines >= max(2, len(lines) - 1):
            return False
        if avg_words < 8.0 or avg_len < 45.0:
            return False
        if long_lines < max(1, len(lines) // 2):
            return False
        return True

    def _translate_preserve_form_leaders(text):
        raw = text or ""
        if not raw.strip():
            return raw
        if _is_structural_text(raw):
            return raw
        if _guard_enabled and _should_skip_translation(raw):
            return raw
        if _from_pdf_preserve_line_breaks(raw):
            return _translate_preserve_exact_lines(raw)
        if bool(from_pdf):
            normalized_raw = raw.replace("\r\n", "\n").replace("\r", "\n")
            if "\n" in normalized_raw:
                lines = [ln.strip() for ln in normalized_raw.split("\n") if ln.strip()]
                if lines and not _from_pdf_should_merge_soft_breaks(lines):
                    return _translate_preserve_exact_lines(raw)

        # Check batch pre-translation cache for speed
        _cache_key = raw.strip()
        cached = _translation_cache.get(_cache_key)
        if cached is None:
            m = re.match(
                r"^(\s*(?:[^\w\u00C0-\u1EF9\r\n]){1,8}\s*)(.+)$",
                _cache_key,
                flags=re.UNICODE | re.DOTALL,
            )
            if m and m.group(2).strip():
                cached = _translation_cache.get(m.group(2).strip())
        if cached is not None:
            return cached

        if not leader_re.search(raw):
            masked, placeholders = _guard_mask_tokens(raw) if _guard_enabled else (raw, {})
            result = _cleanup_translated_text(
                service._translate_with_retry(masked, target_lang, context='document_docx_line')
            )
            result = _guard_restore_tokens(result, placeholders)
            _translation_cache[_cache_key] = result
            return result

        parts = leader_re.split(raw)
        out_parts = []
        for i, part in enumerate(parts):
            if i % 2 == 1:
                out_parts.append(part)
                continue

            seg = part or ""
            if not seg.strip():
                out_parts.append(seg)
                continue
            if _is_structural_text(seg):
                out_parts.append(seg)
                continue

            try:
                _seg_key = seg.strip()
                if _guard_enabled and _should_skip_translation(seg):
                    out_parts.append(seg)
                elif _seg_key in _translation_cache:
                    out_parts.append(_translation_cache[_seg_key])
                else:
                    masked, placeholders = _guard_mask_tokens(seg) if _guard_enabled else (seg, {})
                    translated = _cleanup_translated_text(
                        service._translate_with_retry(masked, target_lang, context='document_docx_line')
                    )
                    translated = _guard_restore_tokens(translated, placeholders)
                    _translation_cache[_seg_key] = translated
                    out_parts.append(translated)
            except ProviderRateLimitError:
                raise
            except Exception:
                if api_only:
                    raise
                out_parts.append(seg)

        return _cleanup_translated_text("".join(out_parts))

    def _normalize_from_pdf_soft_breaks(text: str) -> str:
        """Collapse pdf2docx soft wraps into normal prose, keep list/reference structures."""
        raw = "" if text is None else str(text)
        if not bool(from_pdf) or not raw.strip():
            return raw

        normalized = raw.replace("\r\n", "\n").replace("\r", "\n")
        if "\n" not in normalized:
            return raw

        if _from_pdf_preserve_line_breaks(normalized):
            return raw

        if _paragraph_is_multi_reference_list(normalized):
            return raw
        if leader_re.search(normalized):
            return raw
        if _looks_like_formula(normalized) or _looks_like_code(normalized):
            return raw

        lines = [ln.strip() for ln in normalized.split("\n") if ln.strip()]
        if len(lines) < 2:
            return raw
        if any(_pdf_list_line_re.match(ln) for ln in lines):
            return raw
        if not _from_pdf_should_merge_soft_breaks(lines):
            return raw

        compact = re.sub(r"(?<=[A-Za-z\u00C0-\u1EF9])-\n(?=[A-Za-z\u00C0-\u1EF9])", "", normalized)
        compact = re.sub(r"(?<!\n)\n(?!\n)", " ", compact)
        compact = re.sub(r"[ \t]+", " ", compact)
        compact = re.sub(r" *\n *", "\n", compact)
        compact = re.sub(r"\n{3,}", "\n\n", compact)
        compact = compact.strip()
        return compact or raw

    if bool(from_pdf):
        try:
            from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                is_pdf_artifact_text,
                is_running_header_text,
            )
        except Exception:
            is_running_header_text = lambda _t: False  # noqa: E731
            is_pdf_artifact_text = lambda _t: False  # noqa: E731

        def _is_pdf_noise_text(text: str) -> bool:
            return is_running_header_text(text) or is_pdf_artifact_text(text)
    else:
        def _is_pdf_noise_text(text: str) -> bool:
            return False

    def _cleanup_from_pdf_marker_artifacts(text: str) -> str:
        """Normalize marker artifacts in from-PDF translated text.

        Common artifacts after pdf2docx + translation:
        - orphan marker paragraphs containing only '+'
        - heading lines prefixed with '+-' (marker + heading dash)
        """
        raw = "" if text is None else str(text)
        if not raw:
            return raw

        normalized = raw.replace("\r\n", "\n").replace("\r", "\n")
        changed = False
        out_lines = []
        for ln in normalized.split("\n"):
            s = ln.strip()
            # Drop orphan marker-only lines (ASCII '+' or private-use symbol markers).
            if re.match(r"^[+\uF000-\uF8FF](?:\s*[+\uF000-\uF8FF])*$", s):
                changed = True
                continue

            fixed = ln
            # Heading artifact: "+- Analyze ..." -> "- Analyze ..."
            fixed = re.sub(r"^(\s*)[+\uF000-\uF8FF]\s*-\s*", r"\1- ", fixed)
            # Remove trailing orphan marker at end of content line: "...\t+" -> "..."
            fixed = re.sub(r"([\w\)\]\.:;!?\u00C0-\u1EF9])\s+[+\uF000-\uF8FF]\s*$", r"\1", fixed)
            if fixed != ln:
                changed = True
            out_lines.append(fixed)

        result = "\n".join(out_lines)
        if changed:
            return result
        return raw

    def _translate_body_paragraph_text(text):
        if _is_pdf_noise_text((text or "").strip()):
            return text or ""
        prepared = _normalize_from_pdf_soft_breaks(text)
        if bool(from_pdf):
            # Recover merged single-line paragraphs (pdf2docx collapsed multiple
            # form/bullet lines into one line without \n)
            prepared = _split_compact_pdf_paragraph(prepared)
            translated = _translate_preserve_form_leaders(prepared)
            return _cleanup_from_pdf_marker_artifacts(translated)
        return _translate_preserve_exact_lines(prepared)

    def _is_toc_paragraph(paragraph):
        try:
            style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
            if "toc" in style_name:
                return True
        except Exception:
            pass

        p_el = paragraph._element
        try:
            for node in p_el.xpath('.//*[local-name()="instrText"]'):
                txt = "".join(node.itertext())
                if "toc" in (txt or "").lower():
                    return True
        except Exception:
            pass

        try:
            for node in p_el.xpath('.//*[local-name()="fldSimple"]'):
                for k, v in (node.attrib or {}).items():
                    if str(k).endswith("}instr") and "toc" in str(v or "").lower():
                        return True
        except Exception:
            pass

        return False

    def _flatten_hyperlinks_in_paragraph(paragraph):
        changed = False
        p_el = paragraph._element
        while True:
            links = list(p_el.xpath('./*[local-name()="hyperlink"]'))
            if not links:
                break
            for link in links:
                parent = link.getparent()
                if parent is None:
                    continue
                idx = parent.index(link)
                for child in list(link):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(link)
                changed = True
        return changed

    def _normalize_toc_run_appearance(paragraph):
        return

    def _normalize_toc_hyperlinks(document):
        touched = 0
        for para in iter_all_paragraphs(document):
            if not _is_toc_paragraph(para):
                continue
            if _flatten_hyperlinks_in_paragraph(para):
                touched += 1
            _normalize_toc_run_appearance(para)
        return touched

    def _strip_all_hyperlinks(document):
        touched = 0
        for para in iter_all_paragraphs(document):
            if _flatten_hyperlinks_in_paragraph(para):
                touched += 1
        return touched

    def _paragraph_has_drawing(paragraph):
        """True if paragraph contains drawings, VML pictures, OMML math, or OLE objects."""
        try:
            return bool(paragraph._element.xpath(
                './/*[local-name()="drawing" or local-name()="pict"'
                ' or local-name()="object" or local-name()="OLEObject"'
                ' or local-name()="oMath" or local-name()="oMathPara"]'
            ))
        except Exception:
            return False

    def _set_paragraph_text_preserve_runs(paragraph, new_text):
        runs = list(paragraph.runs)
        original_texts = [(r.text or "") for r in runs]
        if not runs:
            paragraph.add_run(new_text or "")
            return

        if _paragraph_has_drawing(paragraph):
            non_drawing_runs = []
            non_drawing_indices = []
            for i, r in enumerate(runs):
                try:
                    has_draw = bool(r._element.xpath(
                        './/*[local-name()="drawing" or local-name()="pict"'
                        ' or local-name()="object" or local-name()="OLEObject"'
                        ' or local-name()="oMath" or local-name()="oMathPara"]'
                    ))
                except Exception:
                    has_draw = False
                if not has_draw:
                    non_drawing_runs.append(r)
                    non_drawing_indices.append(i)

            if not non_drawing_runs:
                return

            text_value = new_text or ""
            if "\n" in text_value.replace("\r\n", "\n").replace("\r", "\n"):
                target = _pick_text_target_run(non_drawing_runs) or non_drawing_runs[0]
                for r in non_drawing_runs:
                    if r is not target:
                        r.text = ""
                _write_text_with_line_breaks(paragraph, target, text_value)
            else:
                write_translated_to_run_indices(
                    runs,
                    non_drawing_indices,
                    original_texts,
                    text_value,
                    ensure_font_compat_fn=_ensure_text_font_compat,
                )
            return

        text_value = new_text or ""
        if "\n" in text_value.replace("\r\n", "\n").replace("\r", "\n"):
            target = _pick_text_target_run(runs) or runs[0]
            for r in runs:
                if r is not target:
                    r.text = ""
            _write_text_with_line_breaks(paragraph, target, text_value)
            return

        write_translated_to_run_indices(
            runs,
            list(range(len(runs))),
            original_texts,
            text_value,
            ensure_font_compat_fn=_ensure_text_font_compat,
        )

    def _is_in_table_cell(paragraph):
        try:
            parent = paragraph._element.getparent()
            return bool(parent is not None and (parent.tag or '').endswith('}tc'))
        except Exception:
            return False

    def _is_heading_paragraph(paragraph):
        try:
            style_name = str(getattr(getattr(paragraph, 'style', None), 'name', '') or '').lower()
            style_name = style_name.strip()
            if style_name.startswith('heading') or style_name in ('title', 'subtitle'):
                return True
            if 'title' in style_name or 'tiêu đề' in style_name or 'tieu de' in style_name:
                return True
            return False
        except Exception:
            return False

    def _append_inline_bilingual(paragraph, translated_text, delimiter):
        t = (translated_text or '').strip()
        if not t:
            return False

        d = (delimiter or '|').strip() or '|'
        existing = merged_paragraph_plain(paragraph)
        if d in (existing or ''):
            return False

        apply_docx_paragraph_spacing(paragraph)
        runs = list(paragraph.runs)

        if not runs:
            paragraph.add_run(t)
            return True

        def _run_is_superscript(run):
            try:
                if bool(getattr(getattr(run, 'font', None), 'superscript', False)):
                    return True
            except Exception:
                pass
            try:
                from docx.oxml.ns import qn as _qn
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    va = rpr.find(_qn('w:vertAlign'))
                    if va is not None:
                        v = str(va.get(_qn('w:val'), '') or '').strip().lower()
                        if v in ('superscript', 'subscript'):
                            return True
            except Exception:
                pass
            return False

        def _run_is_bold(run):
            try:
                if run.bold is True:
                    return True
            except Exception:
                pass
            try:
                if getattr(getattr(run, 'font', None), 'bold', None) is True:
                    return True
            except Exception:
                pass
            try:
                from docx.oxml.ns import qn as _qn
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(_qn('w:b'))
                    if b is not None:
                        bv = str(b.get(_qn('w:val'), '') or '').strip().lower()
                        if bv not in ('0', 'false', 'off'):
                            return True
            except Exception:
                pass
            return False

        def _apply_run_style(dst_run, src_run):
            if src_run is None:
                return
            try:
                import copy as _copy
                from docx.oxml.ns import qn as _qn
                src_rpr = src_run._element.find(_qn('w:rPr'))
                if src_rpr is None:
                    return
                new_rpr = _copy.deepcopy(src_rpr)
                sp = new_rpr.find(_qn('w:spacing'))
                if sp is not None:
                    new_rpr.remove(sp)
                old_rpr = dst_run._element.find(_qn('w:rPr'))
                if old_rpr is not None:
                    dst_run._element.remove(old_rpr)
                dst_run._element.insert(0, new_rpr)
            except Exception:
                pass

        src_runs = [r for r in runs if (r.text or '').strip()]
        non_sup_runs = [r for r in src_runs if not _run_is_superscript(r)]

        bold_template = None
        for r in non_sup_runs:
            if _run_is_bold(r):
                bold_template = r
                break

        body_template = None
        for r in non_sup_runs:
            if not _run_is_bold(r):
                body_template = r
                break
        if body_template is None and non_sup_runs:
            body_template = non_sup_runs[0]
        if body_template is None and src_runs:
            body_template = src_runs[0]

        has_mixed_bold_prefix = False
        try:
            prefix = ''
            saw_after = False
            for r in src_runs:
                rt = r.text or ''
                if not rt.strip() or _run_is_superscript(r):
                    continue
                if _run_is_bold(r) and not saw_after:
                    prefix += rt
                    continue
                if prefix.strip():
                    saw_after = True
                break
            if prefix.strip() and saw_after and len(prefix.strip()) <= 48:
                has_mixed_bold_prefix = bool(re.search(r'[:：\.]\s*$', prefix.strip()))
        except Exception:
            pass

        last_run = None
        for r in reversed(runs):
            if (r.text or '').strip():
                last_run = r
                break
        if last_run is None:
            last_run = runs[-1]

        def _append_delimiter_once():
            txt = last_run.text or ''
            tail = txt.rstrip()
            if tail.endswith(d):
                if not txt.endswith(' '):
                    last_run.text = f"{txt} "
            else:
                spacer = "" if txt.endswith((" ", "\t")) else " "
                last_run.text = f"{txt}{spacer}{d} "
            try:
                from docx.oxml.ns import qn as _qn
                for t_el in last_run._element.findall(_qn('w:t')):
                    t_el.set(_qn('xml:space'), 'preserve')
            except Exception:
                pass

        if has_mixed_bold_prefix and bold_template is not None:
            m_label = re.match(r'^(\s*[^\n]{1,64}?(?:[:：\.])\s*)(.*)$', t, flags=re.DOTALL)
            if m_label and (m_label.group(2) or '').strip():
                _append_delimiter_once()
                r1 = paragraph.add_run(m_label.group(1))
                copy_run_rpr_preserve_all(bold_template, r1)
                append_formatted_translation_runs(
                    paragraph,
                    m_label.group(2),
                    non_sup_runs if non_sup_runs else src_runs,
                    ensure_font_compat_fn=_ensure_text_font_compat,
                )
                return True

        _append_delimiter_once()
        append_formatted_translation_runs(
            paragraph,
            t,
            non_sup_runs if non_sup_runs else src_runs,
            ensure_font_compat_fn=_ensure_text_font_compat,
        )
        return True

    def _should_keep_newline_translation_in_same_paragraph(paragraph, paragraph_text):
        """Keep short top-form labels in one paragraph to avoid layout jumps."""
        core = (paragraph_text or '').strip()
        if not core:
            return False
        if re.match(r'^Mẫu\s+GT\s*\d+\s*$', core, flags=re.IGNORECASE):
            return True
        return False

    def _append_translation_as_linebreak(paragraph, translated_text):
        """Append translation after a soft line break inside the same paragraph."""
        t = (translated_text or '').strip()
        if not t:
            return False
        existing = merged_paragraph_plain(paragraph)
        if '\n' in (existing or '') and _paragraph_has_newline_bilingual(paragraph, existing):
            return False

        runs = list(paragraph.runs)
        style_run = None
        for r in reversed(runs):
            if (r.text or '').strip():
                style_run = r
                break
        if style_run is None:
            paragraph.add_run(t)
            return True

        try:
            style_run.add_break()
        except Exception:
            style_run.text = f"{style_run.text or ''}\n"

        trans_run = paragraph.add_run(t)
        try:
            copy_run_rpr_preserve_all(style_run, trans_run)
        except Exception:
            pass
        try:
            _ensure_text_font_compat(trans_run, t)
        except Exception:
            pass
        return True

    def _looks_like_newline_title_already_translated(lines):
        clean = [(ln or "").strip() for ln in (lines or []) if (ln or "").strip()]
        if len(clean) < 4 or len(clean) % 2 != 0:
            return False
        mid = len(clean) // 2
        src_half = clean[:mid]
        trans_half = clean[mid:]
        return (
            _pdf_multiline_block_is_title(src_half)
            and any(_looks_source_for_target(ln) for ln in src_half)
            and trans_half
            and all(not _looks_source_for_target(ln) for ln in trans_half)
        )

    def _normalize_existing_newline_title_block(lines):
        """Normalize already-mixed title bilingual blocks and remove duplicate translations."""
        clean = [(ln or "").strip() for ln in (lines or []) if (ln or "").strip()]
        if len(clean) < 4:
            return None
        if any(_pdf_list_line_re.match(ln) or _ref_entry_line_re.match(ln) for ln in clean):
            return None

        source_lines = [ln for ln in clean if _has_vietnamese_diacritics(ln)]
        target_lines = [ln for ln in clean if not _has_vietnamese_diacritics(ln)]
        if not source_lines or not target_lines:
            return None
        if not _pdf_title_source_lines_acceptable(source_lines):
            return None

        unique_targets = []
        seen_targets = set()
        for ln in target_lines:
            key = _norm_ws(ln).lower()
            if not key or key in seen_targets:
                continue
            seen_targets.add(key)
            unique_targets.append(ln)
        if not unique_targets:
            return None

        normalized = source_lines + unique_targets
        if [_norm_ws(ln) for ln in clean] == [_norm_ws(ln) for ln in normalized]:
            return None
        return normalized

    def _is_newline_title_bilingual_block(lines):
        clean = [(ln or "").strip() for ln in (lines or []) if (ln or "").strip()]
        if len(clean) < 2:
            return False
        if any(_pdf_list_line_re.match(ln) or _ref_entry_line_re.match(ln) for ln in clean):
            return False
        source_lines = [ln for ln in clean if _has_vietnamese_diacritics(ln)]
        target_lines = [ln for ln in clean if not _has_vietnamese_diacritics(ln)]
        if source_lines and _pdf_title_source_lines_acceptable(source_lines):
            return True
        if target_lines and _pdf_multiline_block_is_title(target_lines):
            return True
        return False

    def _normalize_title_lines_from_mixed_blocks(lines):
        clean = [(ln or "").strip() for ln in (lines or []) if (ln or "").strip()]
        if len(clean) < 4:
            return None
        if any(_pdf_list_line_re.match(ln) or _ref_entry_line_re.match(ln) for ln in clean):
            return None

        source_lines = [ln for ln in clean if _has_vietnamese_diacritics(ln)]
        target_lines = [ln for ln in clean if not _has_vietnamese_diacritics(ln)]
        if not source_lines or not target_lines:
            return None
        if not _pdf_title_source_lines_acceptable(source_lines):
            return None

        unique_targets = []
        seen_targets = set()
        for ln in target_lines:
            key = _norm_ws(ln).lower()
            if not key or key in seen_targets:
                continue
            seen_targets.add(key)
            unique_targets.append(ln)
        if not unique_targets:
            return None

        normalized = source_lines + unique_targets
        if [_norm_ws(ln) for ln in clean] == [_norm_ws(ln) for ln in normalized]:
            return None
        return normalized

    def _extract_pdf_title_source_line(line, delimiter):
        core = (line or "").strip()
        if not core:
            return ""
        d = service._normalize_bilingual_delimiter(delimiter)
        if d in core:
            core = core.split(d, 1)[0].strip()
        return core

    def _collect_pdf_title_source_lines(lines, delimiter):
        parts = []
        for ln in (lines or []):
            core = _extract_pdf_title_source_line(ln, delimiter)
            if not core:
                continue
            if _has_vietnamese_diacritics(core):
                parts.append(core)
            elif _looks_source_for_target(core):
                parts.append(core)
        return parts

    def _merge_pdf_title_source_text(lines, delimiter):
        parts = _collect_pdf_title_source_lines(lines, delimiter)
        if not parts:
            return ""
        if len(parts) == 1:
            if not _pdf_title_source_lines_acceptable(parts):
                return ""
        elif not _pdf_title_source_lines_acceptable(parts):
            return ""
        return _norm_ws(" ".join(parts))

    def _apply_pdf_inline_title_bilingual(paragraph, lines, delimiter):
        """Merge wrapped PDF title lines into one logical inline bilingual block."""
        d = service._normalize_bilingual_delimiter(delimiter)
        src_merged = _merge_pdf_title_source_text(lines, d)
        if not src_merged:
            return False
        try:
            tr_merged = (_translate_preserve_form_leaders(src_merged) or "").strip()
        except ProviderRateLimitError:
            raise
        except Exception:
            tr_merged = ""
        if not tr_merged or not _bilingual_translation_usable(src_merged, tr_merged):
            return False
        _set_paragraph_text_preserve_runs(
            paragraph,
            service._join_inline_bilingual(src_merged, tr_merged, d),
        )
        _mark_bilingual_processed(paragraph)
        return True

    def _append_newline_bilingual(paragraph, translated_text, delimiter):
        """Song ngữ xuống dòng = cùng luồng liền kề (format run), rồi tách xuống dòng dưới."""
        t = (translated_text or '').strip()
        if not t:
            return False

        d = service._normalize_bilingual_delimiter(delimiter)
        existing = merged_paragraph_plain(paragraph)
        if _should_keep_newline_translation_in_same_paragraph(paragraph, existing):
            return _append_translation_as_linebreak(paragraph, t)
        if d in (existing or ''):
            if _paragraph_has_newline_bilingual(paragraph, existing):
                return False
            new_p = split_inline_bilingual_to_newline(paragraph, d)
            if new_p is not None:
                try:
                    _seen_para_elems.add(id(new_p))
                except Exception:
                    pass
            return new_p is not None

        if not _append_inline_bilingual(paragraph, t, d):
            return False
        new_p = split_inline_bilingual_to_newline(paragraph, d)
        if new_p is not None:
            try:
                _seen_para_elems.add(id(new_p))
            except Exception:
                pass
        return new_p is not None

    def _paragraph_uses_numbering(paragraph):
        try:
            ppr = paragraph._element.find('.//*[local-name()="pPr"]')
            if ppr is not None:
                num_pr = ppr.find('.//*[local-name()="numPr"]')
                if num_pr is not None:
                    return True
        except Exception:
            pass
        try:
            txt = ''.join((r.text or '') for r in paragraph.runs).strip()
        except Exception:
            txt = ''
        if not txt:
            return False
        return bool(re.match(r'^\s*\d+(?:[\.)]|(?:\.\d+)+(?:[\.)])?)\s+\S', txt))

    def _prefer_linebreak_for_newline(paragraph):
        return _is_heading_paragraph(paragraph) or _paragraph_uses_numbering(paragraph)

    def _normalize_heading_case(document):
        return 0

    def _normalize_table_header_text(text: str) -> str:
        t = (text or '').strip()
        norm = re.sub(r'\s+', ' ', t).lower()
        mapping = {
            'user id': 'User ID',
            'userid': 'User ID',
            'data type': 'Data Types',
            'data types': 'Data Types',
            'description': 'Description',
            'constraints': 'Constraints',
            'constraint': 'Constraints',
            'not nul': 'Not null',
        }
        return mapping.get(norm, t)

    def _normalize_table_layout_and_text(document):
        touched = 0
        term_map = {
            'school': 'Field',
            'mô tả': 'Description',
            'mo ta': 'Description',
            'ràng buộc': 'Constraints',
            'rang buoc': 'Constraints',
            'data types': 'Data Types',
            'data type': 'Data Types',
        }

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        raw = ''.join((rr.text or '') for rr in para.runs)
                        if not raw.strip():
                            continue

                        fixed = _cleanup_translated_text(raw)
                        lowered = fixed.strip().lower()
                        if lowered in term_map:
                            fixed = term_map[lowered]

                        if str(target_lang).strip().lower().startswith('en') and _has_vietnamese_diacritics(fixed):
                            try:
                                fixed = _cleanup_translated_text(_translate_preserve_exact_lines(fixed))
                            except Exception:
                                pass

                        if fixed != raw:
                            _set_paragraph_text_preserve_runs(para, fixed)
                            touched += 1
        return touched

    def _format_profile_field_label(label: str) -> str:
        raw = (label or "").strip()
        low = re.sub(r"\s+", " ", raw.lower())
        mapping = {
            "mssv": "MSSV",
            "student id": "Student ID",
            "full name": "Full name",
            "email": "Email",
            "e-mail": "Email",
            "class": "Class",
            "lớp": "Class",
            "lop": "Class",
            "họ và tên": "Full name",
            "ho ten": "Full name",
        }
        return mapping.get(low, raw[:1].upper() + raw[1:] if raw else raw)

    def _normalize_profile_tab_leaders(document):
        touched = 0
        key_re = re.compile(
            r"^(\s*)(full\s*name|họ\s*và\s*tên|ho\s*(?:va|ten)|mssv|student\s*id|"
            r"email|e-mail|class|lớp|lop)\s*[:\-]?\s*(.*)$",
            flags=re.IGNORECASE,
        )
        for para in iter_all_paragraphs(document):
            raw = "".join((r.text or "") for r in para.runs)
            if not raw.strip():
                continue
            low = raw.lower()
            if "email" in low and "class" in low:
                continue

            m = key_re.match(raw.strip())
            if not m:
                continue

            label = _format_profile_field_label(m.group(2))
            rest = m.group(3) or ""
            value = re.sub(r"^[\.\-_:\s\t]+", "", rest).strip()
            if not value:
                continue

            new_text = f"{label}:\t{value}"
            normalized_raw = re.sub(r"\s+", " ", raw.strip())
            normalized_new = re.sub(r"\s+", " ", new_text)
            if normalized_new != normalized_raw:
                _set_paragraph_text_preserve_runs(para, new_text)
                touched += 1

            try:
                para.paragraph_format.tab_stops.add_tab_stop(
                    Inches(5.6), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS
                )
            except Exception:
                pass
        return touched

    def _merge_orphan_method_punct_runs(document):
        """Re-attach orphan '()' / '.()' runs split from method names during translation."""
        touched = 0
        for para in iter_all_paragraphs(document):
            runs = list(para.runs)
            changed = False
            for i in range(1, len(runs)):
                cur = (runs[i].text or "").strip()
                if not cur or not re.fullmatch(r"[\(\)\.]+", cur):
                    continue
                prev = runs[i - 1].text or ""
                if prev and re.search(r"[\w\)]$", prev.rstrip()):
                    runs[i - 1].text = prev + (runs[i].text or "")
                    runs[i].text = ""
                    changed = True
            if changed:
                raw = "".join((r.text or "") for r in runs)
                fixed = _fix_method_call_tail_artifacts(raw)
                if fixed != raw:
                    _set_paragraph_text_preserve_runs(para, fixed)
                touched += 1
        return touched

    def _fix_list_marker_duplicates(document):
        """Collapse only duplicated markers ('+ +', '- -'); keep single original markers."""
        touched = 0
        for para in iter_all_paragraphs(document):
            raw = "".join((r.text or "") for r in para.runs)
            if not raw.strip():
                continue
            fixed = _fix_duplicate_list_markers_in_text(raw)
            if fixed != raw:
                _set_paragraph_text_preserve_runs(para, fixed)
                touched += 1
        return touched

    def _shrink_table_cell_fonts(document):
        touched = 0
        max_len = int(os.getenv('DOCX_TABLE_SHRINK_LEN', '140'))
        min_pt = float(os.getenv('DOCX_TABLE_SHRINK_MIN_PT', '8'))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = "".join(p.text or "" for p in cell.paragraphs).strip()
                    if not cell_text:
                        continue
                    if len(cell_text) < max_len:
                        continue
                    scale = max(0.7, min(1.0, float(max_len) / float(max(1, len(cell_text)))))
                    adjusted = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            try:
                                if run.font.size is None:
                                    continue
                                new_pt = max(min_pt, float(run.font.size.pt) * scale)
                                run.font.size = Pt(new_pt)
                                adjusted = True
                            except Exception:
                                continue
                    if adjusted:
                        touched += 1
        return touched

    def _relax_table_row_heights(document):
        """Allow rows to grow for bilingual text instead of clipping fixed-height cells."""
        relaxed = 0
        for table in document.tables:
            try:
                table.autofit = True
            except Exception:
                pass
            for row in table.rows:
                changed = False
                try:
                    if row.height is not None:
                        row.height = None
                        changed = True
                except Exception:
                    pass
                try:
                    if row.height_rule != WD_ROW_HEIGHT_RULE.AT_LEAST:
                        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                        changed = True
                except Exception:
                    pass
                try:
                    tr_pr = row._tr.get_or_add_trPr()
                    for tr_h in list(tr_pr.xpath('./*[local-name()="trHeight"]')):
                        tr_pr.remove(tr_h)
                        changed = True
                except Exception:
                    pass
                if changed:
                    relaxed += 1
        return relaxed

    def _center_inline_images(document):
        centered = 0
        for para in iter_all_paragraphs(document):
            has_drawing = False
            try:
                for run in para.runs:
                    dr = run._element.xpath('.//*[local-name()="drawing"]')
                    if dr:
                        has_drawing = True
                        break
            except Exception:
                has_drawing = False
            if not has_drawing:
                continue

            # Only center image-only paragraphs; avoid centering mixed text paragraphs.
            try:
                para_text = merged_paragraph_plain(para).strip()
            except Exception:
                para_text = ""
            if para_text:
                continue

            try:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                centered += 1
            except Exception:
                pass
        return centered

    def _force_remaining_phrase_fixes(document):
        touched = 0
        replacements = [
            (r"\bKHOA\s+CÔNG\s+NGHỆ\s+THÔNG\s+TIN\b", "FACULTY OF INFORMATION TECHNOLOGY"),
            (r"\bKHOA\s+CONG\s+NGHE\s+THONG\s+TIN\b", "FACULTY OF INFORMATION TECHNOLOGY"),
            (r"Gửi\s+lại\s+phiếu\s+đăng\s+ký\s+qua\s+Email\s*:", "Resubmit the registration form via Email:"),
            (r"Gui\s+lai\s+phieu\s+dang\s+ky\s+qua\s+Email\s*:", "Resubmit the registration form via Email:"),
        ]

        for para in iter_all_paragraphs(document):
            if _paragraph_has_drawing(para):
                continue
            raw = ''.join((r.text or '') for r in para.runs)
            if not raw.strip():
                continue
            fixed = raw
            for pat, rep in replacements:
                fixed = re.sub(pat, rep, fixed, flags=re.IGNORECASE)
            fixed = _cleanup_translated_text(fixed)
            if fixed != raw:
                _set_paragraph_text_preserve_runs(para, fixed)
                touched += 1
        return touched

    def translate_paragraph_runs(paragraph, idx=None, total=None):
        runs = list(paragraph.runs)
        if not runs:
            return

        if bool(from_pdf) or bi_mode in ('inline', 'newline'):
            apply_docx_paragraph_spacing(paragraph)
            runs = list(paragraph.runs)
        paragraph_text = merged_paragraph_plain(paragraph)
        if not paragraph_text.strip():
            return
        if _should_preserve_author_metadata(paragraph, paragraph_text):
            return
        translate_fn = _translate_body_paragraph_text if bool(from_pdf) else _translate_preserve_form_leaders

        if bi_mode in ("inline", "newline") and _paragraph_is_multi_reference_list(paragraph_text):
            _handle_multi_reference_bilingual(paragraph, paragraph_text)
            return

        if bi_mode == 'newline':
            if _is_structural_text(paragraph_text):
                return
            if _is_bilingual_processed(paragraph):
                return
            try:
                translated_para = translate_fn(paragraph_text)
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Translator failed for paragraph: {e}")
                if api_only:
                    raise
                translated_para = paragraph_text
            d = service._normalize_bilingual_delimiter(bilingual_delimiter)
            _try_apply_bilingual(paragraph, paragraph_text, translated_para, d)
        elif bi_mode == 'inline':
            if _is_bilingual_processed(paragraph):
                return
            try:
                translated_para = translate_fn(paragraph_text)
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Translator failed for paragraph: {e}")
                if api_only:
                    raise
                translated_para = paragraph_text
            d = service._normalize_bilingual_delimiter(bilingual_delimiter)
            _try_apply_bilingual(paragraph, paragraph_text, translated_para, d)
        else:
            _translate_format_groups(paragraph, translate_fn)

        if progress_callback and idx is not None and total is not None:
            progress_callback(10 + int((idx / total) * 70), f"Translating paragraph {idx+1}/{total}")

    _seen_para_elems = set()
    _bilingual_processed = set()

    def _seen_or_mark(paragraph):
        try:
            key = id(paragraph._element)
        except Exception:
            key = id(paragraph)
        if key in _seen_para_elems:
            return True
        _seen_para_elems.add(key)
        return False

    def _paragraph_text(paragraph):
        try:
            return merged_paragraph_plain(paragraph)
        except Exception:
            return ''

    def _norm_ws(s):
        return re.sub(r'\s+', ' ', (s or '').strip())

    def _looks_source_for_target(raw_text):
        t_root = str(target_lang).strip().lower()
        txt = (raw_text or '').strip()
        if not txt:
            return False
        if t_root.startswith('en'):
            try:
                return bool(service._looks_vietnamese_like_text(txt))
            except Exception:
                return _has_vietnamese_diacritics(txt)
        if t_root.startswith('vi'):
            if _has_vietnamese_diacritics(txt):
                return False
            try:
                if service._looks_vietnamese_like_text(txt):
                    return False
            except Exception:
                pass
            return bool(re.search(r'[A-Za-z]{3,}', txt))
        return bool(re.search(r'[\w\u00C0-\u1EF9]', txt, flags=re.UNICODE))

    def _bilingual_translation_usable(source, translated):
        """Reject echo translations (e.g. English / English when target is Vietnamese)."""
        src = _norm_ws(source)
        tr = _norm_ws(translated)
        if not tr or tr == src:
            return False
        t_root = str(target_lang).strip().lower()
        if t_root.startswith('vi'):
            if _has_vietnamese_diacritics(tr):
                return True
            src_words = {w.lower() for w in re.findall(r'[A-Za-z]{3,}', src)}
            tr_words = {w.lower() for w in re.findall(r'[A-Za-z]{3,}', tr)}
            if tr_words and tr_words == src_words:
                return False
            if len(tr_words) >= 2 and not _has_vietnamese_diacritics(tr):
                return False
            return True
        if t_root.startswith('en'):
            if _has_vietnamese_diacritics(tr):
                return True
            try:
                return bool(service._looks_vietnamese_like_text(tr))
            except Exception:
                return tr != src
        return tr != src

    def _mark_bilingual_processed(paragraph):
        try:
            _bilingual_processed.add(id(paragraph._element))
        except Exception:
            _bilingual_processed.add(id(paragraph))

    def _is_bilingual_processed(paragraph):
        try:
            key = id(paragraph._element)
        except Exception:
            key = id(paragraph)
        return key in _bilingual_processed

    def _paragraph_bilingual_complete(paragraph, raw_text, delimiter):
        d = (delimiter or '|').strip() or '|'
        raw = raw_text or ''
        if bi_mode == 'inline':
            if d not in raw:
                return False
            trans_side = raw.split(d, 1)[1].strip()
            return bool(trans_side) and not _inline_side_still_untranslated(trans_side)
        if bi_mode == 'newline':
            return _paragraph_has_newline_bilingual(paragraph, raw)
        return False

    def _pdf_multiline_block_is_title(lines):
        """Detect a PDF title/heading split across physical wrapped lines."""
        clean = [(ln or "").strip() for ln in (lines or []) if (ln or "").strip()]
        if len(clean) < 2 or len(clean) > 8:
            return False
        if any(_pdf_list_line_re.match(ln) or _ref_entry_line_re.match(ln) for ln in clean):
            return False
        joined = " ".join(clean)
        if len(joined) > 280:
            return False
        # Body/abstract prose should not be treated as title.
        if any(re.search(r"[.!?]\s*$", ln) for ln in clean):
            return False
        letter_count = len(re.findall(r"[A-Za-zÀ-ỹ]", joined))
        if letter_count < 12:
            return False
        lower_count = len(re.findall(r"[a-zà-ỹ]", joined))
        upper_count = len(re.findall(r"[A-ZÀ-Ỹ]", joined))
        uppercase_ratio = upper_count / max(1, upper_count + lower_count)
        avg_words = sum(len(re.findall(r"[\wÀ-ỹ]+", ln, flags=re.UNICODE)) for ln in clean) / max(1, len(clean))
        return uppercase_ratio >= 0.72 and avg_words <= 9

    def _pdf_title_source_lines_acceptable(lines):
        """Accept PDF title source blocks even when pdf2docx merged wrapped lines."""
        if _pdf_multiline_block_is_title(lines):
            return True
        clean = [(ln or "").strip() for ln in (lines or []) if (ln or "").strip()]
        if len(clean) < 2 or len(clean) > 8:
            return False
        if any(_pdf_list_line_re.match(ln) or _ref_entry_line_re.match(ln) for ln in clean):
            return False
        joined = " ".join(clean)
        if len(joined) > 320 or any(re.search(r"[.!?]\s*$", ln) for ln in clean):
            return False
        if not any(_has_vietnamese_diacritics(ln) for ln in clean):
            return False
        letter_count = len(re.findall(r"[A-Za-zÀ-ỹ]", joined))
        if letter_count < 12:
            return False
        lower_count = len(re.findall(r"[a-zà-ỹ]", joined))
        upper_count = len(re.findall(r"[A-ZÀ-Ỹ]", joined))
        uppercase_ratio = upper_count / max(1, upper_count + lower_count)
        avg_words = sum(len(re.findall(r"[\wÀ-ỹ]+", ln, flags=re.UNICODE)) for ln in clean) / max(1, len(clean))
        return uppercase_ratio >= 0.65 and avg_words <= 14

    def _try_apply_pdf_linewise_bilingual(paragraph, source_text, translated_text, delimiter):
        """For PDF-converted multi-line paragraphs, pair each line in-place.

        pdf2docx can merge a heading plus following references into one paragraph.
        Appending one big translation block at the end makes the translation appear
        beside/below the wrong source line. Pairing line-by-line keeps translations
        anchored to their source lines and leaves already-target-language lines
        untouched.
        """
        if not bool(from_pdf) or bi_mode not in ('inline', 'newline'):
            return None
        src_raw = (source_text or '').replace("\r\n", "\n").replace("\r", "\n")
        tr_raw = (translated_text or '').replace("\r\n", "\n").replace("\r", "\n")
        if "\n" not in src_raw or "\n" not in tr_raw:
            return None

        src_lines = src_raw.split("\n")
        tr_lines = tr_raw.split("\n")
        if len(src_lines) != len(tr_lines):
            return None

        d = service._normalize_bilingual_delimiter(delimiter)
        out_lines = []
        changed = False
        if bi_mode == "inline":
            src_only = _collect_pdf_title_source_lines(src_lines, d)
            if _pdf_title_source_lines_acceptable(src_only):
                if _apply_pdf_inline_title_bilingual(paragraph, src_lines, d):
                    return True
        if bi_mode == 'newline' and _pdf_multiline_block_is_title(src_lines):
            trans_lines = []
            for src_line, tr_line in zip(src_lines, tr_lines):
                src_core = (src_line or '').strip()
                tr_core = (tr_line or '').strip()
                if not src_core:
                    continue
                if (not tr_core or _norm_ws(src_core) == _norm_ws(tr_core)) and _looks_source_for_target(src_core):
                    try:
                        retry_line = (_translate_preserve_form_leaders(src_core) or '').strip()
                        if retry_line and _norm_ws(retry_line) != _norm_ws(src_core):
                            tr_core = retry_line
                    except ProviderRateLimitError:
                        raise
                    except Exception:
                        pass
                if tr_core and _norm_ws(src_core) != _norm_ws(tr_core) and _bilingual_translation_usable(src_core, tr_core):
                    trans_lines.append(tr_core)
            if trans_lines:
                _set_paragraph_text_preserve_runs(paragraph, "\n".join(src_lines + trans_lines))
                _mark_bilingual_processed(paragraph)
                return True

        for src_line, tr_line in zip(src_lines, tr_lines):
            src_core = (src_line or '').strip()
            tr_core = (tr_line or '').strip()
            if not src_core:
                out_lines.append(src_line)
                continue
            if (not tr_core or _norm_ws(src_core) == _norm_ws(tr_core)) and _looks_source_for_target(src_core):
                try:
                    retry_line = (_translate_preserve_form_leaders(src_core) or '').strip()
                    if retry_line and _norm_ws(retry_line) != _norm_ws(src_core):
                        tr_core = retry_line
                except ProviderRateLimitError:
                    raise
                except Exception:
                    pass
            if not tr_core or _norm_ws(src_core) == _norm_ws(tr_core):
                out_lines.append(src_line)
                continue
            if not _bilingual_translation_usable(src_core, tr_core):
                out_lines.append(src_line)
                continue
            if bi_mode == 'inline':
                out_lines.append(service._join_inline_bilingual(src_line.rstrip(), tr_core, d))
            else:
                out_lines.append(src_line.rstrip())
                out_lines.append(tr_core)
            changed = True

        if not changed:
            return False

        _set_paragraph_text_preserve_runs(paragraph, "\n".join(out_lines))
        _mark_bilingual_processed(paragraph)
        return True

    def _try_apply_bilingual(paragraph, source_text, translated_text, delimiter):
        src = (source_text or '').strip()
        t = (translated_text or '').strip()
        if not src:
            return False
        d = service._normalize_bilingual_delimiter(delimiter)
        raw = _paragraph_text(paragraph)
        if _paragraph_bilingual_complete(paragraph, raw, d):
            _mark_bilingual_processed(paragraph)
            return False
        linewise = _try_apply_pdf_linewise_bilingual(paragraph, src, t, d)
        if linewise is not None:
            return bool(linewise)
        if not t or not _bilingual_translation_usable(src, t):
            return False
        if bi_mode == 'inline':
            if d in raw:
                src_side = raw.split(d, 1)[0].strip()
                if _bilingual_translation_usable(src_side or src, t):
                    _set_paragraph_text_preserve_runs(
                        paragraph,
                        service._join_inline_bilingual(src_side or src, t, d),
                    )
                    _mark_bilingual_processed(paragraph)
                    return True
                return False
            ok = _append_inline_bilingual(paragraph, t, d)
        else:
            ok = _append_newline_bilingual(paragraph, t, d)
        if ok:
            _mark_bilingual_processed(paragraph)
        return ok

    def _rescue_table_cells_all_modes():
        touched = 0
        d = service._normalize_bilingual_delimiter(bilingual_delimiter)
        target_root = str(target_lang).strip().lower()

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paras = list(cell.paragraphs)
                    i = 0
                    while i < len(paras):
                        p = paras[i]
                        raw = _paragraph_text(p)
                        core = (raw or '').strip()
                        if not core or _paragraph_has_drawing(p) or _is_structural_text(core):
                            i += 1
                            continue
                        # Skip cells already given a usable bilingual translation to avoid
                        # re-translating (wasteful API calls) and re-appending (duplication).
                        if bi_mode in ('inline', 'newline'):
                            if _is_bilingual_processed(p) or _paragraph_bilingual_complete(p, raw, d):
                                _mark_bilingual_processed(p)
                                i += 1
                                continue
                        should_retry = _looks_source_for_target(core)
                        # Fallback detector for EN target: keep retrying table text that still looks Vietnamese.
                        if (not should_retry) and target_root.startswith('en'):
                            should_retry = _has_vietnamese_diacritics(core)
                        if not should_retry:
                            i += 1
                            continue

                        try:
                            translated = _translate_preserve_exact_lines(core)
                        except ProviderRateLimitError:
                            raise
                        except Exception:
                            i += 1
                            continue

                        t = (translated or '').strip()
                        if not t or not _bilingual_translation_usable(core, t):
                            i += 1
                            continue

                        if bi_mode == 'inline':
                            if _try_apply_bilingual(p, core, t, d):
                                touched += 1
                            i += 1
                            continue

                        if bi_mode == 'newline':
                            has_next_same = False
                            if i + 1 < len(paras):
                                nxt = (_paragraph_text(paras[i + 1]) or '').strip()
                                if _norm_ws(nxt) == _norm_ws(t):
                                    has_next_same = True
                                elif nxt and (not _is_structural_text(nxt)) and (not _looks_source_for_target(nxt)):
                                    has_next_same = True

                            if not has_next_same and _try_apply_bilingual(p, core, t, d):
                                touched += 1
                                paras = list(cell.paragraphs)
                                i += 2
                                continue

                            i += 1
                            continue

                        # normal mode: force-replace untranslated source text inside table cells.
                        if bool(from_pdf):
                            _apply_translation_to_runs(p, translated)
                        else:
                            _set_paragraph_text_preserve_runs(p, translated)
                        touched += 1
                        i += 1

        return touched

    def _rescue_untranslated_paragraphs(paragraphs):
        """Second pass: re-translate paragraphs still in source language."""
        if bi_mode in ("inline", "newline"):
            return 0
        touched = 0
        for para in paragraphs or []:
            if _paragraph_has_drawing(para):
                continue
            raw = _paragraph_text(para)
            core = (raw or "").strip()
            if not core or _is_structural_text(core):
                continue
            source_text, prefix = _paragraph_translation_source_text(para, raw)
            translate_input = (source_text or core).strip()
            if not translate_input:
                continue
            should_retry = _looks_source_for_target(translate_input)
            if not should_retry and str(target_lang).strip().lower().startswith("en"):
                should_retry = _has_vietnamese_diacritics(translate_input)
                if not should_retry:
                    for run in para.runs:
                        rt = (run.text or "").strip()
                        if not rt or _is_structural_text(rt) or _run_has_drawing_xml(run):
                            continue
                        if _has_vietnamese_diacritics(rt):
                            should_retry = True
                            break
            if not should_retry:
                continue
            try:
                translated = _translate_preserve_exact_lines(translate_input)
            except ProviderRateLimitError:
                raise
            except Exception:
                continue
            t = (translated or "").strip()
            if not t or _norm_ws(t) == _norm_ws(translate_input):
                continue
            if prefix or any(
                _is_decorative_leading_run_text(original or "")
                for original in [(r.text or "") for r in para.runs]
            ):
                _apply_translation_to_runs(para, translated)
            else:
                _translate_format_groups(para, _translate_body_paragraph_text)
            touched += 1
        return touched

    def _paragraph_has_inline_bilingual(text):
        d = service._normalize_bilingual_delimiter(bilingual_delimiter)
        return d in (text or "")

    def _paragraph_has_newline_bilingual(paragraph, paragraph_text):
        core = (paragraph_text or "").strip()
        if not core:
            return False
        normalized = (paragraph_text or "").replace("\r\n", "\n").replace("\r", "\n")
        if "\n" in normalized:
            parts = [p.strip() for p in normalized.split("\n") if p.strip()]
            if len(parts) >= 2 and parts[0] == core.split("\n")[0].strip():
                tail = "\n".join(parts[1:]).strip()
                if tail and not _looks_source_for_target(tail):
                    return True
        try:
            parent = paragraph._element.getparent()
            if parent is None:
                return False
            siblings = list(parent)
            idx = siblings.index(paragraph._element)
            if idx + 1 >= len(siblings):
                return False
            from docx.text.paragraph import Paragraph as _Paragraph

            nxt = _Paragraph(siblings[idx + 1], paragraph._parent)
            nxt_text = (_paragraph_text(nxt) or "").strip()
            if not nxt_text or _is_structural_text(nxt_text):
                return False
            if _norm_ws(nxt_text) == _norm_ws(core):
                return False
            if not _looks_source_for_target(nxt_text):
                return True
        except Exception:
            return False
        return False

    def _inline_side_still_untranslated(trans_side: str) -> bool:
        side = (trans_side or "").strip()
        if not side:
            return True
        t_root = str(target_lang).strip().lower()
        if t_root.startswith("en"):
            try:
                return bool(service._looks_vietnamese_like_text(side))
            except Exception:
                return _has_vietnamese_diacritics(side)
        if t_root.startswith("vi"):
            if _has_vietnamese_diacritics(side):
                return False
            latin_words = re.findall(r"[A-Za-z]{3,}", side)
            return len(latin_words) >= 2
        return False

    def _rescue_untranslated_bilingual_paragraphs(paragraphs):
        """Second pass for inline/newline: only fix paragraphs still missing a usable translation."""
        if bi_mode not in ("inline", "newline"):
            return 0
        touched = 0
        d = service._normalize_bilingual_delimiter(bilingual_delimiter)
        paras = list(paragraphs or [])
        i = 0
        while i < len(paras):
            p = paras[i]
            if _paragraph_has_drawing(p) or _is_bilingual_processed(p):
                i += 1
                continue
            raw = _paragraph_text(p)
            core = (raw or "").strip()
            if not core or _is_structural_text(core):
                i += 1
                continue
            if _should_preserve_author_metadata(p, core):
                i += 1
                continue
            if _paragraph_bilingual_complete(p, raw, d):
                _mark_bilingual_processed(p)
                i += 1
                continue
            if not _looks_source_for_target(core):
                i += 1
                continue
            src_core = core.split(d, 1)[0].strip() if (bi_mode == "inline" and d in raw) else core
            try:
                translated = _translate_preserve_exact_lines(src_core)
            except ProviderRateLimitError:
                raise
            except Exception:
                i += 1
                continue
            if _try_apply_bilingual(p, src_core, translated, d):
                touched += 1
                if bi_mode == "newline":
                    paras = list(getattr(p, "_parent", None).paragraphs) if hasattr(p, "_parent") else paras
                    i += 2
                    continue
            i += 1
        return touched

    def _repair_pdf_multiline_bilingual_positions():
        """Final repair for pdf2docx paragraphs that contain multiple logical lines.

        Some PDF conversions keep headings, abstracts, and references inside one
        Word paragraph. In bilingual modes, the safest placement is per physical
        line: translate immediately beside/below that line, and leave lines that
        are already in the target language unchanged.
        """
        if not bool(from_pdf) or bi_mode not in ("inline", "newline"):
            return 0
        touched = 0
        d = service._normalize_bilingual_delimiter(bilingual_delimiter)

        def _looks_like_existing_translation_for(src_line: str, next_line: str) -> bool:
            nxt = (next_line or "").strip()
            if not nxt:
                return False
            if re.match(r"^\s*(?:\[\d+\]|\d{1,3}[.)])\s+", nxt):
                return False
            if _looks_source_for_target(nxt):
                return False
            return True

        def _normalize_consecutive_title_paragraphs():
            if bi_mode not in ("inline", "newline"):
                return 0
            fixed = 0
            paras = list(doc.paragraphs)
            d_local = service._normalize_bilingual_delimiter(bilingual_delimiter)

            def _is_title_line_candidate(text):
                line = (text or "").strip()
                if not line:
                    return False
                norm = _norm_ws(line).lower()
                if norm in {"tóm tắt", "tom tat", "summary", "abstract", "từ khóa", "tu khoa", "keywords"}:
                    return False
                if "@" in line or re.search(r"\bemail\b|\bcontact\b", line, flags=re.IGNORECASE):
                    return False
                if _pdf_list_line_re.match(line) or _ref_entry_line_re.match(line):
                    return False
                if len(line) > 150:
                    return False
                words = re.findall(r"[A-Za-zÀ-Ỹà-ỹĐđ0-9]+", line)
                if not words or len(words) > 14:
                    return False
                letters = [ch for ch in line if ch.isalpha()]
                if not letters:
                    return bool(re.search(r"\b20\d{2}\b", line))
                upper_ratio = sum(1 for ch in letters if ch.upper() == ch and ch.lower() != ch) / max(1, len(letters))
                return upper_ratio >= 0.65

            def _paragraph_title_candidate_lines(paragraph):
                raw = _paragraph_text(paragraph)
                if not raw:
                    return []
                lines = [(ln or "").strip() for ln in raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
                lines = [ln for ln in lines if ln]
                if not lines:
                    return []
                if all(_is_title_line_candidate(ln) for ln in lines):
                    return lines
                return []

            def _normalize_single_line_title_paragraphs():
                scan_limit = min(len(paras), 40)
                i = 0
                while i < scan_limit:
                    p = paras[i]
                    if _paragraph_has_drawing(p):
                        i += 1
                        continue
                    first_lines = _paragraph_title_candidate_lines(p)
                    if not first_lines:
                        i += 1
                        continue

                    group = [p]
                    group_lines = list(first_lines)
                    j = i + 1
                    while j < scan_limit and len(group_lines) <= 24:
                        q = paras[j]
                        if _paragraph_has_drawing(q):
                            break
                        q_lines = _paragraph_title_candidate_lines(q)
                        if not q_lines:
                            break
                        group.append(q)
                        group_lines.extend(q_lines)
                        j += 1

                    if bi_mode == "inline":
                        src_merged = _merge_pdf_title_source_text(group_lines, d_local)
                        if not src_merged:
                            i += 1
                            continue
                        try:
                            tr_merged = (_translate_preserve_form_leaders(src_merged) or "").strip()
                        except ProviderRateLimitError:
                            raise
                        except Exception:
                            tr_merged = ""
                        if not tr_merged:
                            i += 1
                            continue
                        _set_paragraph_text_preserve_runs(
                            group[0],
                            service._join_inline_bilingual(src_merged, tr_merged, d_local),
                        )
                        _mark_bilingual_processed(group[0])
                        for extra in group[1:]:
                            try:
                                parent = extra._element.getparent()
                                if parent is not None:
                                    parent.remove(extra._element)
                            except Exception:
                                pass
                        fixed += 1
                        paras = list(doc.paragraphs)
                        i += 1
                        continue

                    if len(group_lines) < 4:
                        i += 1
                        continue
                    normalized = _normalize_title_lines_from_mixed_blocks(group_lines)
                    if normalized is None:
                        i = j
                        continue

                    _set_paragraph_text_preserve_runs(group[0], "\n".join(normalized))
                    _mark_bilingual_processed(group[0])
                    for extra in group[1:]:
                        try:
                            parent = extra._element.getparent()
                            if parent is not None:
                                parent.remove(extra._element)
                        except Exception:
                            pass
                    return 1
                return 0

            fixed += _normalize_single_line_title_paragraphs()
            paras = list(doc.paragraphs)
            i = 0
            while i < len(paras):
                p = paras[i]
                if _paragraph_has_drawing(p):
                    i += 1
                    continue
                raw = _paragraph_text(p)
                lines = raw.replace("\r\n", "\n").replace("\r", "\n").split("\n") if raw else []
                if not _is_newline_title_bilingual_block(lines):
                    i += 1
                    continue

                group = [p]
                group_lines = [(ln or "").strip() for ln in lines if (ln or "").strip()]
                j = i + 1
                while j < len(paras) and len(group_lines) <= 18:
                    q = paras[j]
                    if _paragraph_has_drawing(q):
                        break
                    q_raw = _paragraph_text(q)
                    q_lines = q_raw.replace("\r\n", "\n").replace("\r", "\n").split("\n") if q_raw else []
                    if not _is_newline_title_bilingual_block(q_lines):
                        break
                    group.append(q)
                    group_lines.extend((ln or "").strip() for ln in q_lines if (ln or "").strip())
                    j += 1

                if len(group) <= 1:
                    i += 1
                    continue

                if bi_mode == "inline":
                    src_merged = _merge_pdf_title_source_text(group_lines, d_local)
                    if not src_merged:
                        i = j
                        continue
                    try:
                        tr_merged = (_translate_preserve_form_leaders(src_merged) or "").strip()
                    except ProviderRateLimitError:
                        raise
                    except Exception:
                        tr_merged = ""
                    if not tr_merged:
                        i = j
                        continue
                    _set_paragraph_text_preserve_runs(
                        group[0],
                        service._join_inline_bilingual(src_merged, tr_merged, d_local),
                    )
                else:
                    normalized = _normalize_title_lines_from_mixed_blocks(group_lines)
                    if normalized is None:
                        i = j
                        continue
                    _set_paragraph_text_preserve_runs(group[0], "\n".join(normalized))
                _mark_bilingual_processed(group[0])
                for extra in group[1:]:
                    try:
                        parent = extra._element.getparent()
                        if parent is not None:
                            parent.remove(extra._element)
                    except Exception:
                        pass
                fixed += 1
                paras = list(doc.paragraphs)
                i += 1
            return fixed

        _title_fixed = _normalize_consecutive_title_paragraphs()
        touched += _title_fixed

        for para in iter_all_paragraphs(doc):
            if _paragraph_has_drawing(para):
                continue
            raw = _paragraph_text(para)
            if not raw or "\n" not in raw:
                continue
            if bi_mode == "inline" and d in raw:
                continue
            lines = raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")
            out_lines = []
            changed = False
            if bi_mode == "newline":
                normalized_title = _normalize_existing_newline_title_block(lines)
                if normalized_title is not None:
                    _set_paragraph_text_preserve_runs(para, "\n".join(normalized_title))
                    _mark_bilingual_processed(para)
                    touched += 1
                    continue
            if bi_mode == "newline" and _looks_like_newline_title_already_translated(lines):
                _mark_bilingual_processed(para)
                continue
            if bi_mode == "inline":
                src_only = _collect_pdf_title_source_lines(lines, d)
                if _pdf_title_source_lines_acceptable(src_only):
                    if _apply_pdf_inline_title_bilingual(para, lines, d):
                        touched += 1
                        continue
            if bi_mode == "newline" and _pdf_multiline_block_is_title(lines):
                trans_lines = []
                for line in lines:
                    core = (line or "").strip()
                    if not core or not _looks_source_for_target(core):
                        continue
                    try:
                        translated = (_translate_preserve_form_leaders(core) or "").strip()
                    except ProviderRateLimitError:
                        raise
                    except Exception:
                        continue
                    if translated and _bilingual_translation_usable(core, translated):
                        trans_lines.append(translated)
                if trans_lines:
                    _set_paragraph_text_preserve_runs(para, "\n".join(lines + trans_lines))
                    _mark_bilingual_processed(para)
                    touched += 1
                continue

            for idx, line in enumerate(lines):
                core = (line or "").strip()
                out_lines.append(line)
                if not core or not _looks_source_for_target(core):
                    continue
                if idx + 1 < len(lines) and _looks_like_existing_translation_for(core, lines[idx + 1]):
                    continue
                try:
                    translated = (_translate_preserve_form_leaders(core) or "").strip()
                except ProviderRateLimitError:
                    raise
                except Exception:
                    continue
                if not translated or not _bilingual_translation_usable(core, translated):
                    continue
                if bi_mode == "inline":
                    out_lines[-1] = service._join_inline_bilingual(line.rstrip(), translated, d)
                else:
                    out_lines.append(translated)
                changed = True
            if changed:
                _set_paragraph_text_preserve_runs(para, "\n".join(out_lines))
                _mark_bilingual_processed(para)
                touched += 1
        return touched

    def _rescue_untranslated_body_paragraphs():
        return _rescue_untranslated_paragraphs(body_paras)

    def _sanitize_from_pdf_paragraphs(document):
        """Normalize selected pdf2docx spacing artifacts with conservative defaults."""
        fixed_count = 0
        tab_paragraphs = 0
        sanitize_replace_nbsp = str(os.getenv("PDF_DOCX_SANITIZE_REPLACE_NBSP", "1")).strip().lower() in (
            "1",
            "true",
            "yes",
            "on",
        )
        sanitize_replace_tabs = str(os.getenv("PDF_DOCX_SANITIZE_REPLACE_TABS", "0")).strip().lower() in (
            "1",
            "true",
            "yes",
            "on",
        )
        sanitize_strip_soft_hyphen = str(os.getenv("PDF_DOCX_SANITIZE_STRIP_SOFT_HYPHEN", "0")).strip().lower() in (
            "1",
            "true",
            "yes",
            "on",
        )

        try:
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WPA
            justify_values = {
                int(v)
                for v in (
                    getattr(_WPA, 'JUSTIFY', None),
                    getattr(_WPA, 'DISTRIBUTE', None),
                    getattr(_WPA, 'THAI_JUSTIFY', None),
                )
                if v is not None
            }
            center_like_values = {
                int(v)
                for v in (
                    getattr(_WPA, 'CENTER', None),
                    getattr(_WPA, 'RIGHT', None),
                )
                if v is not None
            }
            left_value = getattr(_WPA, 'LEFT', 0)
        except Exception:
            justify_values = {3, 4, 7, 8, 9}
            center_like_values = {1, 2}
            left_value = 0

        try:
            from docx.oxml.ns import qn as _qn
        except Exception:
            _qn = None

        def _normalize_symbol_marker_run_text(run, text: str) -> str:
            """Replace private-use Symbol/Wingdings marker runs with plain '+' marker.
            pdf2docx often emits bullets as PUA chars (e.g. U+F02B in SymbolMT),
            which can render as tofu boxes after DOCX->PDF export."""
            raw_text = text or ""
            if not raw_text:
                return raw_text
            try:
                if not _run_uses_symbol_font(run):
                    return raw_text
            except Exception:
                return raw_text

            stripped = raw_text.strip()
            if not stripped:
                return raw_text
            # Only normalize short marker-like runs fully composed of PUA symbols.
            if len(stripped) > 4:
                return raw_text
            if not all('\uf000' <= ch <= '\uf8ff' for ch in stripped):
                return raw_text

            lead = raw_text[: len(raw_text) - len(raw_text.lstrip())]
            tail = raw_text[len(raw_text.rstrip()) :]
            return f"{lead}+{tail}"

        for para in iter_all_paragraphs(document):
            try:
                runs = list(para.runs)
            except Exception:
                continue
            if not runs:
                continue

            # Never mutate run text in paragraphs that contain drawings or OMML math:
            # setting run.text calls clear_content() which destroys embedded objects.
            if _para_has_protected_xml(para):
                continue

            local_changed = False
            local_has_tab = False

            for run in runs:
                # Skip individual runs that carry drawing/OLE children.
                if _run_has_drawing_xml(run):
                    continue
                try:
                    raw = run.text or ''
                except Exception:
                    continue
                if not raw:
                    continue
                cleaned = raw
                if sanitize_replace_nbsp:
                    cleaned = cleaned.replace('\u00A0', ' ')
                if sanitize_replace_tabs:
                    cleaned = cleaned.replace('\t', ' ')
                if sanitize_strip_soft_hyphen:
                    cleaned = cleaned.replace('\u00AD', '')
                cleaned = _normalize_symbol_marker_run_text(run, cleaned)
                if sanitize_replace_tabs and '\t' in raw:
                    local_has_tab = True
                if cleaned != raw:
                    run.text = cleaned
                    _ensure_text_font_compat(run, cleaned)
                    local_changed = True

            if local_changed:
                apply_docx_paragraph_spacing(para)
                fixed_count += 1

            if local_has_tab:
                tab_paragraphs += 1
                if _qn is not None:
                    try:
                        ppr = para._element.find(_qn('w:pPr'))
                        if ppr is not None:
                            tabs = ppr.find(_qn('w:tabs'))
                            if tabs is not None:
                                ppr.remove(tabs)
                    except Exception:
                        pass

            if str(os.getenv("PDF_DOCX_SANITIZE_FORCE_LEFT", "0")).strip().lower() in (
                "1",
                "true",
                "yes",
                "on",
            ):
                try:
                    a = para.alignment
                    a_val = int(a) if a is not None else None
                    para_text = merged_paragraph_plain(para).strip()
                    word_count = len(re.findall(r'\w+', para_text, flags=re.UNICODE))
                    looks_center_artifact = bool(para_text) and (
                        len(para_text) >= 60
                        or word_count >= 10
                        or (
                            len(para_text) >= 30
                            and para_text.endswith(('.', '!', '?', ':', ';'))
                        )
                    )

                    should_force_left = (
                        a_val in justify_values
                        or (
                            a_val in center_like_values
                            and looks_center_artifact
                            and not _is_heading_paragraph(para)
                            and not _is_in_table_cell(para)
                            and not _paragraph_has_drawing(para)
                        )
                    )

                    if should_force_left:
                        para.alignment = left_value
                        fixed_count += 1
                except Exception:
                    pass

        return fixed_count, tab_paragraphs

    if bool(from_pdf):
        try:
            fixed_count, tab_paragraphs = _sanitize_from_pdf_paragraphs(doc)
            if progress_callback and (fixed_count > 0 or tab_paragraphs > 0):
                progress_callback(
                    9,
                    f"DOCX from_pdf cleanup: fixed {fixed_count} paragraphs, tab-paragraphs {tab_paragraphs}",
                )
        except Exception:
            pass
        try:
            from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                merge_pdf_line_fragments_in_doc,
            )

            merge_stats = merge_pdf_line_fragments_in_doc(doc)
            merged_total = int(merge_stats.get("fragments_merged", 0)) + int(
                merge_stats.get("body_fragments_merged", 0)
            )
            if progress_callback and merged_total > 0:
                progress_callback(
                    9,
                    f"DOCX line-wrap merge: collapsed {merged_total} fragment paragraph(s)",
                )
        except Exception as exc:
            try:
                print(f"PDF line fragment merge skipped: {exc}")
            except Exception:
                pass

    try:
        _normalize_toc_hyperlinks(doc)
    except Exception:
        pass

    if bool(from_pdf) and bi_mode in ("inline", "newline"):
        try:
            dup_removed = dedupe_consecutive_identical_paragraphs(doc)
            if progress_callback and dup_removed > 0:
                progress_callback(9, f"DOCX dedupe: removed {dup_removed} duplicate paragraphs")
        except Exception:
            pass

    paragraphs = [p for p in doc.paragraphs]
    from concurrent.futures import as_completed

    total_work = 0
    completed = 0

    body_paras = []
    for para in paragraphs:
        if _seen_or_mark(para):
            continue
        paragraph_text = merged_paragraph_plain(para)
        if not paragraph_text.strip():
            continue
        body_paras.append(para)
    total_work = max(1, len(body_paras))

    # === BATCH PRE-TRANSLATION for speed optimization ===
    # Collect all unique translatable texts from body + tables + headers/footers,
    # batch them with <<<S>>> separator into fewer API calls, cache results.
    _translation_cache = {}
    _all_texts = []
    _seen_texts = set()

    def _queue_text(para):
        """Collect paragraph text for batch pre-translation."""
        text = _normalize_from_pdf_soft_breaks(merged_paragraph_plain(para)).strip()
        if bool(from_pdf) and _from_pdf_preserve_line_breaks(text):
            return
        if bool(from_pdf):
            normalized = text.replace("\r\n", "\n").replace("\r", "\n")
            if "\n" in normalized:
                lines = [ln.strip() for ln in normalized.split("\n") if ln.strip()]
                if lines and not _from_pdf_should_merge_soft_breaks(lines):
                    return
            # Exclude paragraphs that _split_compact_pdf_paragraph would rewrite.
            # Caching pre-split text is useless and can bypass line-structure recovery.
            if _split_compact_pdf_paragraph(text) != text:
                return
        if text and not _is_structural_text(text) and text not in _seen_texts:
            if _should_skip_translation(text) or _is_pdf_noise_text(text):
                return
            source_text, _prefix = _paragraph_translation_source_text(para, text)
            queue_key = (source_text or text).strip()
            if not queue_key or queue_key in _seen_texts:
                return
            if _should_skip_translation(queue_key) or _is_pdf_noise_text(queue_key):
                return
            _seen_texts.add(queue_key)
            masked, placeholders = _guard_mask_tokens(queue_key) if _guard_enabled else (queue_key, {})
            _all_texts.append({"raw": queue_key, "masked": masked, "placeholders": placeholders})

    for para in body_paras:
        _queue_text(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _queue_text(p)
    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                _queue_text(p)
            for p in section.footer.paragraphs:
                _queue_text(p)
    except Exception:
        pass
    try:
        for p in _iter_docx_extra_paragraphs(doc):
            _queue_text(p)
    except Exception:
        pass

    _special_text_segments = collect_docx_special_text_segments(doc)
    for seg in _special_text_segments:
        core = seg.get("core") or ""
        if not core or core in _seen_texts:
            continue
        if _should_skip_translation(core) or _is_pdf_noise_text(core):
            continue
        _seen_texts.add(core)
        masked, placeholders = _guard_mask_tokens(core) if _guard_enabled else (core, {})
        _all_texts.append({"raw": core, "masked": masked, "placeholders": placeholders})

    if _all_texts:
        _BATCH_CHARS = int(os.getenv('DOCX_BATCH_CHAR_LIMIT', '1200'))
        _BATCH_COUNT = int(os.getenv('DOCX_BATCH_COUNT_LIMIT', '8'))
        _SEP = '\n<<<S>>>\n'
        _i = 0
        _total_texts = len(_all_texts)
        while _i < _total_texts:
            _batch_entries = []
            _batch_masked = []
            _chars = 0
            while _i < _total_texts and len(_batch_entries) < _BATCH_COUNT and _chars < _BATCH_CHARS:
                entry = _all_texts[_i]
                _batch_entries.append(entry)
                _batch_masked.append(entry.get("masked") or "")
                _chars += len(entry.get("masked") or "")
                _i += 1
            if not _batch_entries:
                break
            try:
                if len(_batch_entries) == 1:
                    entry = _batch_entries[0]
                    _r = service._translate_with_retry(entry.get("masked") or "", target_lang, context='document_docx_line')
                    restored = _guard_restore_tokens(_cleanup_translated_text(_r), entry.get("placeholders") or {})
                    _translation_cache[entry.get("raw")] = restored
                else:
                    _joined = _SEP.join(_batch_masked)
                    _r = service._translate_with_retry(_joined, target_lang, context='document_docx_batch')
                    _parts = [p.strip() for p in _r.split('<<<S>>>')]
                    if len(_parts) == len(_batch_entries):
                        for entry, _dst in zip(_batch_entries, _parts):
                            restored = _guard_restore_tokens(_cleanup_translated_text(_dst), entry.get("placeholders") or {})
                            _translation_cache[entry.get("raw")] = restored
                    else:
                        # Separator count mismatch — fallback to individual translation
                        for entry in _batch_entries:
                            try:
                                _r2 = service._translate_with_retry(entry.get("masked") or "", target_lang, context='document_docx_line')
                                restored = _guard_restore_tokens(
                                    _cleanup_translated_text(_r2),
                                    entry.get("placeholders") or {},
                                )
                                _translation_cache[entry.get("raw")] = restored
                            except ProviderRateLimitError:
                                raise
                            except Exception:
                                pass
            except ProviderRateLimitError:
                raise
            except Exception:
                # Batch failed — fallback to individual translation
                for entry in _batch_entries:
                    try:
                        _r2 = service._translate_with_retry(entry.get("masked") or "", target_lang, context='document_docx_line')
                        restored = _guard_restore_tokens(
                            _cleanup_translated_text(_r2),
                            entry.get("placeholders") or {},
                        )
                        _translation_cache[entry.get("raw")] = restored
                    except ProviderRateLimitError:
                        raise
                    except Exception:
                        pass
            if progress_callback:
                progress_callback(
                    10 + int((min(_i, _total_texts) / max(1, _total_texts)) * 70),
                    f"Batch translating {min(_i, _total_texts)}/{_total_texts}",
                )

    for para in body_paras:
        try:
            if bool(from_pdf) or bi_mode in ('inline', 'newline'):
                apply_docx_paragraph_spacing(para)
            paragraph_text_src = merged_paragraph_plain(para)

            if _is_pdf_noise_text(paragraph_text_src.strip()):
                completed += 1
                if progress_callback:
                    progress_callback(
                        10 + int((completed / total_work) * 70),
                        f"Translating paragraph {completed}/{total_work}",
                    )
                continue

            if _should_preserve_author_metadata(para, paragraph_text_src):
                completed += 1
                if progress_callback:
                    progress_callback(
                        10 + int((completed / total_work) * 70),
                        f"Translating paragraph {completed}/{total_work}",
                    )
                continue

            if bi_mode in ("inline", "newline") and _paragraph_is_multi_reference_list(paragraph_text_src):
                _handle_multi_reference_bilingual(para, paragraph_text_src)
                completed += 1
                if progress_callback:
                    progress_callback(
                        10 + int((completed / total_work) * 70),
                        f"Translating paragraph {completed}/{total_work}",
                    )
                continue

            if bi_mode == 'inline':
                d = service._normalize_bilingual_delimiter(bilingual_delimiter)
                if not _is_bilingual_processed(para):
                    translated = _translate_body_paragraph_text(paragraph_text_src)
                    _try_apply_bilingual(para, paragraph_text_src, translated, d)
            elif bi_mode == 'newline':
                if _is_structural_text(paragraph_text_src):
                    pass
                elif not _is_bilingual_processed(para):
                    translated = _translate_body_paragraph_text(paragraph_text_src)
                    d = service._normalize_bilingual_delimiter(bilingual_delimiter)
                    _try_apply_bilingual(para, paragraph_text_src, translated, d)
            else:
                _translate_format_groups(para, _translate_body_paragraph_text)
        except ProviderRateLimitError:
            print("Provider rate limit detected during paragraph processing, aborting job.")
            raise
        except Exception as e:
            print(f"Paragraph translation failed: {e}")
            if api_only:
                raise
        completed += 1
        if progress_callback:
            progress_callback(
                10 + int((completed / total_work) * 70),
                f"Translating paragraph {completed}/{total_work}",
            )

    for table in doc.tables:
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                cell = table.rows[r].cells[c]
                for p_idx, p in enumerate(cell.paragraphs):
                    if _seen_or_mark(p):
                        continue
                    translate_paragraph_runs(p, p_idx, len(cell.paragraphs))

    try:
        for section in doc.sections:
            header = section.header
            for p_idx, p in enumerate(header.paragraphs):
                if _seen_or_mark(p):
                    continue
                translate_paragraph_runs(p, p_idx, len(header.paragraphs))
            footer = section.footer
            for p_idx, p in enumerate(footer.paragraphs):
                if _seen_or_mark(p):
                    continue
                translate_paragraph_runs(p, p_idx, len(footer.paragraphs))
    except Exception:
        pass

    try:
        extra_paras = list(_iter_docx_extra_paragraphs(doc))
        for p_idx, p in enumerate(extra_paras):
            if _seen_or_mark(p):
                continue
            translate_paragraph_runs(p, p_idx, len(extra_paras) or 1)
    except Exception:
        pass

    try:
        rescued_body = _rescue_untranslated_body_paragraphs()
        extra_rescue_paras = []
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extra_rescue_paras.extend(list(cell.paragraphs))
        except Exception:
            pass
        try:
            for section in doc.sections:
                extra_rescue_paras.extend(list(section.header.paragraphs))
                extra_rescue_paras.extend(list(section.footer.paragraphs))
        except Exception:
            pass
        try:
            extra_rescue_paras.extend(list(_iter_docx_extra_paragraphs(doc)))
        except Exception:
            pass
        rescued_extra = _rescue_untranslated_paragraphs(extra_rescue_paras)
        bilingual_rescue_paras = list(body_paras)
        bilingual_rescue_paras.extend(extra_rescue_paras)
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        bilingual_rescue_paras.extend(list(cell.paragraphs))
        except Exception:
            pass
        rescued_bilingual = _rescue_untranslated_bilingual_paragraphs(bilingual_rescue_paras)
        rescued_total = rescued_body + rescued_extra + rescued_bilingual
        if progress_callback and rescued_total > 0:
            progress_callback(81, f"DOCX rescue pass: {rescued_total} paragraphs")
    except ProviderRateLimitError:
        raise
    except Exception:
        pass

    def _resolve_cached_translation(text):
        key = (text or "").strip()
        if key in _translation_cache:
            return _translation_cache[key]
        return _translate_body_paragraph_text(text)

    def _resolve_special_text_translation(text):
        key = (text or "").strip()
        if not key:
            return text
        cached = _translation_cache.get(key)
        if cached and str(cached).strip() and str(cached).strip() != key:
            return cached
        translated = _translate_body_paragraph_text(text)
        if translated and str(translated).strip() and str(translated).strip() != key:
            _translation_cache[key] = translated
            return translated
        try:
            masked, placeholders = _guard_mask_tokens(text) if _guard_enabled else (text, {})
            translated = _cleanup_translated_text(
                service._translate_with_retry(masked, target_lang, context="document_docx_line")
            )
            translated = _guard_restore_tokens(translated, placeholders)
            if translated and str(translated).strip():
                _translation_cache[key] = translated
            return translated
        except ProviderRateLimitError:
            raise
        except Exception:
            return text

    try:
        txbx_count = apply_docx_special_text_translations(
            doc,
            collect_docx_special_text_segments(doc),
            _resolve_special_text_translation,
            bilingual_mode=bi_mode,
            bilingual_delimiter=bilingual_delimiter,
            join_inline_fn=service._join_inline_bilingual,
        )
        if progress_callback and txbx_count > 0:
            progress_callback(82, f"DOCX textbox/WordArt: translated {txbx_count} text(s)")
    except ProviderRateLimitError:
        raise
    except Exception as e:
        print(f"DOCX textbox translation skipped: {e}")

    try:
        rescued_table_cells = _rescue_table_cells_all_modes()
        if progress_callback and rescued_table_cells > 0:
            progress_callback(81, f"DOCX table rescue: {rescued_table_cells} cells")
    except ProviderRateLimitError:
        raise
    except Exception:
        pass

    try:
        repaired_pdf_lines = _repair_pdf_multiline_bilingual_positions()
        if progress_callback and repaired_pdf_lines > 0:
            progress_callback(82, f"DOCX PDF bilingual line repair: {repaired_pdf_lines} paragraphs")
    except ProviderRateLimitError:
        raise
    except Exception:
        pass

    try:
        if bool(from_pdf) and bi_mode in ("inline", "newline"):
            d_cleanup = service._normalize_bilingual_delimiter(bilingual_delimiter)

            def _line_is_pdf_title_candidate(line):
                core = (line or "").strip()
                if not core:
                    return False
                norm = _norm_ws(core).lower()
                if norm in {"tóm tắt", "tom tat", "summary", "abstract", "từ khóa", "tu khoa", "keywords"}:
                    return False
                if "@" in core or re.search(r"\bemail\b|\bcontact\b", core, flags=re.IGNORECASE):
                    return False
                words = re.findall(r"[A-Za-zÀ-Ỹà-ỹĐđ0-9]+", core)
                if not words or len(words) > 14:
                    return False
                letters = [ch for ch in core if ch.isalpha()]
                if not letters:
                    return False
                upper_ratio = sum(1 for ch in letters if ch.upper() == ch and ch.lower() != ch) / max(1, len(letters))
                return upper_ratio >= 0.65

            def _paragraph_plain_lines(paragraph):
                raw = (getattr(paragraph, "text", None) or _paragraph_text(paragraph) or "")
                return [(ln or "").strip() for ln in raw.replace("\r\n", "\n").replace("\r", "\n").split("\n") if (ln or "").strip()]

            def _cleanup_pdf_title_block(group, group_lines):
                source_lines = _collect_pdf_title_source_lines(group_lines, d_cleanup)
                if len(source_lines) < 2 or not _pdf_title_source_lines_acceptable(source_lines):
                    return False
                if bi_mode == "inline":
                    src_merged = _merge_pdf_title_source_text(group_lines, d_cleanup)
                    if not src_merged:
                        return False
                    try:
                        tr_merged = (_translate_preserve_form_leaders(src_merged) or "").strip()
                    except ProviderRateLimitError:
                        raise
                    except Exception:
                        return False
                    if not tr_merged:
                        return False
                    _set_paragraph_text_preserve_runs(
                        group[0],
                        service._join_inline_bilingual(src_merged, tr_merged, d_cleanup),
                    )
                else:
                    target_lines = [
                        ln for ln in group_lines
                        if not _has_vietnamese_diacritics(_extract_pdf_title_source_line(ln, d_cleanup) or ln)
                        and not _has_vietnamese_diacritics(ln)
                    ]
                    unique_targets = []
                    seen_targets = set()
                    for ln in target_lines:
                        key = _norm_ws(ln).lower()
                        if not key or key in seen_targets:
                            continue
                        seen_targets.add(key)
                        unique_targets.append(ln)
                    if not unique_targets:
                        return False
                    _set_paragraph_text_preserve_runs(group[0], "\n".join(source_lines + unique_targets))
                _mark_bilingual_processed(group[0])
                for extra in group[1:]:
                    try:
                        parent = extra._element.getparent()
                        if parent is not None:
                            parent.remove(extra._element)
                    except Exception:
                        pass
                return True

            paras = list(doc.paragraphs)
            cleanup_count = 0
            scan_limit = min(len(paras), 30)
            i = 0
            while i < scan_limit:
                p = paras[i]
                p_lines = _paragraph_plain_lines(p)
                if not p_lines:
                    i += 1
                    continue

                if bi_mode == "inline":
                    if not all(_line_is_pdf_title_candidate(ln) for ln in p_lines):
                        i += 1
                        continue
                    group = [p]
                    group_lines = list(p_lines)
                    j = i + 1
                    while j < scan_limit and len(group) < 8:
                        q = paras[j]
                        q_lines = _paragraph_plain_lines(q)
                        if not q_lines or not all(_line_is_pdf_title_candidate(ln) for ln in q_lines):
                            break
                        group.append(q)
                        group_lines.extend(q_lines)
                        j += 1
                    if _cleanup_pdf_title_block(group, group_lines):
                        cleanup_count += 1
                        paras = list(doc.paragraphs)
                        scan_limit = min(len(paras), 30)
                    i += 1
                    continue

                if not any(_has_vietnamese_diacritics(ln) for ln in p_lines):
                    i += 1
                    continue
                if not all(_line_is_pdf_title_candidate(ln) for ln in p_lines):
                    i += 1
                    continue

                group = [p]
                group_lines = list(p_lines)
                j = i + 1
                while j < scan_limit and len(group) < 6:
                    q = paras[j]
                    q_lines = _paragraph_plain_lines(q)
                    if not q_lines:
                        break
                    if any(_has_vietnamese_diacritics(ln) for ln in q_lines):
                        break
                    if not all(_line_is_pdf_title_candidate(ln) for ln in q_lines):
                        break
                    group.append(q)
                    group_lines.extend(q_lines)
                    j += 1

                if _cleanup_pdf_title_block(group, group_lines):
                    cleanup_count += 1
                    paras = list(doc.paragraphs)
                    scan_limit = min(len(paras), 30)
                i += 1

            if progress_callback and cleanup_count > 0:
                progress_callback(82, f"DOCX PDF title cleanup: {cleanup_count} block(s)")
    except Exception as e:
        print(f"DOCX PDF title cleanup skipped: {e}")

    try:
        if bool(from_pdf) or bi_mode in ('inline', 'newline'):
            relaxed_rows = _relax_table_row_heights(doc)
            if progress_callback and relaxed_rows > 0:
                progress_callback(82, f"DOCX table layout: relaxed {relaxed_rows} row heights")
    except Exception:
        pass

    if ocr_images and service.ocr_translate_overlay:
        if progress_callback:
            progress_callback(82, "OCR images in DOCX...")

        protected_image_partnames = _collect_header_footer_image_partnames(doc)

        paras_to_scan = iter_all_paragraphs(doc)
        total_paras = len(paras_to_scan) or 1
        images_found = 0
        ocr_attempted = 0
        ocr_success = 0
        ocr_disabled = False

        image_replacements = {}
        ocr_export_entries = []
        text_insert_entries = []
        text_replace_entries = []

        def _is_probably_logo_or_nontext(ocr_text: str) -> bool:
            raw = (ocr_text or '').strip()
            if not raw:
                return True
            words = re.findall(r'\w+', raw, flags=re.UNICODE)
            if len(words) <= 2 and len(raw) < 24:
                return True
            return False

        for idx, para in enumerate(paras_to_scan):
            if ocr_disabled:
                break
            rids = paragraph_image_rids(para)
            if not rids:
                continue
            for rid in rids:
                img_part = rid_to_image_part(para, rid)
                if not img_part:
                    continue
                partname = str(getattr(img_part, 'partname', '') or '').lstrip('/')
                if partname and partname in protected_image_partnames:
                    continue
                try:
                    blob = getattr(img_part, 'blob', None)
                    if not blob:
                        continue

                    images_found += 1

                    ext = image_part_ext(img_part)
                    tmp_name = f"docx_img_{uuid.uuid4().hex}{ext}"
                    tmp_path = os.path.join(service.upload_folder, tmp_name)
                    with open(tmp_path, 'wb') as f:
                        f.write(blob)

                    ocr_attempted += 1
                    try:
                        ocr_text, translated_text, png_bytes, ai_recommended_mode = service.ocr_translate_overlay(
                            tmp_path,
                            'auto',
                            target_lang,
                            ocr_langs,
                        )
                    finally:
                        try:
                            os.remove(tmp_path)
                        except Exception:
                            pass

                    if not ocr_text or not str(ocr_text).strip():
                        continue

                    if _is_probably_logo_or_nontext(ocr_text):
                        continue

                    per_mode = mode
                    if mode == 'auto':
                        per_mode = _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode)

                    print(f"  [IMAGE #{images_found}] OCR={len(ocr_text)}chars, AI_class={ai_recommended_mode}, per_mode={per_mode}")

                    try:
                        if per_mode in ('text', 'both'):
                            partname = str(getattr(img_part, 'partname', '') or '').lstrip('/')
                            ocr_export_entries.append({
                                'image': partname or '(embedded image)',
                                'ocr_text': (ocr_text or '').strip(),
                                'translated_text': (translated_text or '').strip(),
                            })
                            normalized_translated = _normalize_ocr_text_for_docx((translated_text or '').strip())

                            if per_mode == 'text' and mode == 'text':
                                text_replace_entries.append((para, rid, normalized_translated))
                            else:
                                text_insert_entries.append((para, normalized_translated))
                    except Exception:
                        pass

                    if per_mode in ('image', 'both') and png_bytes and len(png_bytes) > 100:
                        try:
                            if partname:
                                new_bytes = _overlay_bytes_to_original_format(png_bytes, ext)
                                image_replacements[partname] = new_bytes
                                ocr_success += 1
                        except Exception:
                            continue
                except Exception as e:
                    msg = str(e).lower()
                    if ('tesseract' in msg and ('not installed' in msg or 'path' in msg)) or \
                       ('ocr unavailable' in msg):
                        ocr_disabled = True
                        if progress_callback:
                            progress_callback(85, "Skipping DOCX image OCR (OCR not available)")
                        break
                    if 'ai provider' in msg and ('not configured' in msg or 'rate' in msg):
                        ocr_disabled = True
                        if progress_callback:
                            progress_callback(85, f"Skipping DOCX image OCR: {e}")
                        break
                    print(f"DOCX image OCR error (continuing): {e}")
                    continue

            if progress_callback and (idx % 10 == 0):
                progress_callback(82 + int((idx / total_paras) * 10), f"OCR scanning {idx+1}/{total_paras}")

        try:
            pkg_scan_enabled = str(os.getenv('DOCX_OCR_PACKAGE_SCAN', '0')).strip().lower() in ('1', 'true', 'yes', 'on')
            if pkg_scan_enabled and (not ocr_disabled) and mode in ('image', 'both', 'auto'):
                pkg = getattr(getattr(doc, 'part', None), 'package', None)
                parts = list(getattr(pkg, 'parts', []) or [])

                extra_attempted = 0
                extra_replaced = 0
                for part in parts:
                    try:
                        ct = str(getattr(part, 'content_type', '') or '').lower()
                        if not ct.startswith('image/'):
                            continue
                        partname = str(getattr(part, 'partname', '') or '').lstrip('/')
                        if not partname:
                            continue
                        if partname in protected_image_partnames:
                            continue
                        if partname in image_replacements:
                            continue
                        blob = getattr(part, 'blob', None)
                        if not blob:
                            continue

                        ext = image_part_ext(part)
                        tmp_name = f"docx_img_pkg_{uuid.uuid4().hex}{ext}"
                        tmp_path = os.path.join(service.upload_folder, tmp_name)
                        with open(tmp_path, 'wb') as f:
                            f.write(blob)

                        extra_attempted += 1
                        try:
                            ocr_text, translated_text, png_bytes, ai_recommended_mode = service.ocr_translate_overlay(
                                tmp_path,
                                'auto',
                                target_lang,
                                ocr_langs,
                            )
                        finally:
                            try:
                                os.remove(tmp_path)
                            except Exception:
                                pass

                        if not ocr_text or not str(ocr_text).strip():
                            continue

                        if _is_probably_logo_or_nontext(ocr_text):
                            continue

                        per_mode = mode
                        if mode == 'auto':
                            per_mode = _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode)

                        if per_mode in ('image', 'both') and png_bytes and len(png_bytes) > 100:
                            try:
                                new_bytes = _overlay_bytes_to_original_format(png_bytes, ext)
                                image_replacements[partname] = new_bytes
                                extra_replaced += 1
                            except Exception:
                                continue
                    except Exception as e:
                        msg = str(e).lower()
                        if ('tesseract' in msg and ('not installed' in msg or 'path' in msg)) or \
                           ('ocr unavailable' in msg):
                            ocr_disabled = True
                            if progress_callback:
                                progress_callback(85, "Skipping DOCX image OCR (OCR not available)")
                            break
                        if 'ai provider' in msg and ('not configured' in msg or 'rate' in msg):
                            ocr_disabled = True
                            if progress_callback:
                                progress_callback(85, f"Skipping DOCX image OCR: {e}")
                            break
                        continue

                if progress_callback and (extra_attempted or extra_replaced):
                    progress_callback(
                        92,
                        f"DOCX OCR (package scan): attempted={extra_attempted}, replaced={extra_replaced}",
                    )
        except Exception:
            pass

        if progress_callback:
            if images_found <= 0:
                progress_callback(92, "DOCX OCR: no embedded images found")
            else:
                progress_callback(
                    92,
                    f"DOCX OCR: found={images_found}, attempted={ocr_attempted}, replaced={ocr_success}",
                )

        try:
            if text_replace_entries:
                for para, rid, trans_text in text_replace_entries:
                    if not trans_text or not trans_text.strip():
                        continue
                    try:
                        replace_image_with_text(para, rid, trans_text)
                    except Exception:
                        continue
        except Exception:
            pass

    try:
        if ocr_images and 'text_insert_entries' in locals() and text_insert_entries:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn

            for para, trans_text in reversed(text_insert_entries):
                if not trans_text or not trans_text.strip():
                    continue
                try:
                    text_paragraphs = [p.strip() for p in trans_text.split('\n') if p.strip()]
                    if not text_paragraphs:
                        text_paragraphs = [trans_text.strip()]

                    for t_idx, t_para in enumerate(reversed(text_paragraphs)):
                        new_p = OxmlElement('w:p')

                        run = OxmlElement('w:r')

                        t_el = OxmlElement('w:t')
                        t_el.set(_qn('xml:space'), 'preserve')
                        t_el.text = t_para
                        run.append(t_el)
                        new_p.append(run)

                        para._element.addnext(new_p)
                except Exception:
                    continue
    except Exception:
        pass

    try:
        targeted_default = "0" if bool(from_pdf) else "1"
        targeted_fixes_enabled = _env_on("DOCX_TARGETED_FIXES", targeted_default)
        if bi_mode in ('inline', 'newline'):
            link_count = 0
            leader_count = 0
            table_count = 0
            img_count = 0
            forced_count = 0
            if progress_callback:
                progress_callback(96, "DOCX targeted fixes skipped in bilingual mode")
        elif not targeted_fixes_enabled:
            link_count = 0
            leader_count = 0
            table_count = 0
            img_count = 0
            forced_count = 0
            if progress_callback:
                progress_callback(96, "DOCX targeted fixes disabled (DOCX_TARGETED_FIXES=0)")
        else:
            link_count = _strip_all_hyperlinks(doc)
            leader_count = _normalize_profile_tab_leaders(doc)
            list_marker_count = _fix_list_marker_duplicates(doc)
            method_punct_count = _merge_orphan_method_punct_runs(doc)
            table_count = _normalize_table_layout_and_text(doc)
            img_count = _center_inline_images(doc)
            forced_count = _force_remaining_phrase_fixes(doc)
            if progress_callback:
                progress_callback(
                    96,
                    (
                        f"DOCX targeted fixes: links={link_count}, "
                        f"leaders={leader_count}, list_markers={list_marker_count}, "
                        f"method_punct={method_punct_count}, "
                        f"table={table_count}, images={img_count}, forced={forced_count}"
                    ),
                )
    except Exception as e:
        if progress_callback:
            progress_callback(96, f"DOCX targeted fixes skipped: {e}")

    try:
        if str(os.getenv('DOCX_TABLE_SHRINK', '0')).strip().lower() in ('1', 'true', 'yes', 'on'):
            shrink_count = _shrink_table_cell_fonts(doc)
            if progress_callback:
                progress_callback(97, f"DOCX layout recovery: shrunk {shrink_count} table cells")
    except Exception:
        pass

    try:
        subset_fixed = _normalize_subset_fonts_in_document(doc)
        if bool(from_pdf):
            try:
                from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                    sanitize_document_text_fonts,
                )

                subset_fixed += sanitize_document_text_fonts(doc)
            except Exception:
                pass
        if progress_callback and subset_fixed:
            progress_callback(97, f"DOCX font normalization: fixed {subset_fixed} font entries")
    except Exception:
        pass

    def _fix_translated_paragraph_artifacts(document):
        """Final pass: fix label punctuation and rebuild bold label + body runs."""
        touched = 0
        for para in iter_all_paragraphs(document):
            if _paragraph_has_drawing(para):
                continue
            raw = "".join((r.text or "") for r in para.runs)
            if not raw.strip():
                continue
            fixed = _fix_duplicate_label_punctuation(raw)
            fixed = _fix_method_call_tail_artifacts(fixed)
            if _paragraph_should_use_bold_label_layout(para, fixed):
                if _apply_translation_with_bold_label(para, fixed):
                    touched += 1
                    continue
            if fixed != raw:
                _set_paragraph_text_preserve_runs(para, fixed)
                touched += 1
        return touched

    try:
        artifact_fixes = _fix_translated_paragraph_artifacts(doc)
        if progress_callback and artifact_fixes > 0:
            progress_callback(98, f"DOCX artifact cleanup: fixed {artifact_fixes} paragraphs")
    except Exception:
        pass

    try:
        marker_cleaned = 0
        for para in iter_all_paragraphs(doc):
            for run in para.runs:
                raw = run.text or ""
                cleaned = strip_run_boundary_markers(raw)
                if cleaned != raw:
                    run.text = cleaned
                    marker_cleaned += 1
        if progress_callback and marker_cleaned > 0:
            progress_callback(98, f"DOCX marker cleanup: {marker_cleaned} runs")
    except Exception:
        pass

    try:
        comma_deduped = 0
        for para in iter_all_paragraphs(doc):
            runs = list(para.runs)
            for i in range(1, len(runs)):
                prev = runs[i - 1].text or ""
                cur = runs[i].text or ""
                if prev.rstrip().endswith(",") and cur.strip() in (",", ", "):
                    runs[i].text = ""
                    comma_deduped += 1
        if progress_callback and comma_deduped > 0:
            progress_callback(98, f"DOCX punctuation cleanup: {comma_deduped} runs")
    except Exception:
        pass

    try:
        if bi_mode == "inline":
            d_inline = service._normalize_bilingual_delimiter(bilingual_delimiter)
            d_esc = re.escape(d_inline)
            pipe_deduped = 0

            def _collapse_duplicate_inline_delimiters(text):
                out = text or ""
                collapsed = re.sub(rf"(?:\s*{d_esc}\s*){{2,}}", f" {d_inline} ", out)
                return collapsed

            for para in iter_all_paragraphs(doc):
                raw = _paragraph_text(para)
                if not raw or d_inline not in raw:
                    continue
                cleaned = _collapse_duplicate_inline_delimiters(raw)
                if cleaned != raw:
                    _set_paragraph_text_preserve_runs(para, cleaned)
                    pipe_deduped += 1
            if progress_callback and pipe_deduped > 0:
                progress_callback(98, f"DOCX inline delimiter cleanup: {pipe_deduped} paragraphs")
    except Exception:
        pass

    mode_prefix = ""
    if bi_mode == "inline":
        mode_prefix = "inline_"
    elif bi_mode == "newline":
        mode_prefix = "newline_"
    output_filename = f"translated_{mode_prefix}{os.path.basename(file_path)}"
    if not output_filename.lower().endswith('.docx'):
        output_filename += '.docx'
    output_path = os.path.join(service.download_folder, output_filename)

    doc.save(output_path)

    try:
        if ocr_images and mode in ('image', 'both', 'auto') and 'image_replacements' in locals() and image_replacements:
            if progress_callback:
                progress_callback(93, "Applying translated overlays to DOCX images...")
            tmp_out = output_path + ".tmp"
            with zipfile.ZipFile(output_path, 'r') as zin, zipfile.ZipFile(tmp_out, 'w') as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    repl = image_replacements.get(item.filename)
                    if repl is not None:
                        data = repl
                    zout.writestr(item, data)
            try:
                os.replace(tmp_out, output_path)
            except Exception:
                try:
                    os.remove(output_path)
                except Exception:
                    pass
                os.rename(tmp_out, output_path)
    except Exception as e:
        if progress_callback:
            progress_callback(94, f"DOCX image overlay patch failed: {e}")

    try:
        docx.Document(output_path)
    except Exception as e:
        if progress_callback:
            progress_callback(95, "DOCX validation failed, writing fallback text file")
        fallback_filename = output_filename
        if not fallback_filename.lower().endswith('.txt'):
            fallback_filename += '.txt'
        fallback_path = os.path.join(service.download_folder, fallback_filename)
        lines = []
        for p in doc.paragraphs:
            lines.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    lines.append(cell.text)
        with open(fallback_path, 'w', encoding='utf-8') as f:
            f.write("NOTE: DOCX creation failed on server. Showing plain text fallback below.\n\n")
            f.write('\n'.join(lines))
        output_path = fallback_path

    if progress_callback:
        progress_callback(100, "Completed")
    return output_path
